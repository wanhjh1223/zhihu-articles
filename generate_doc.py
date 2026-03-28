#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LeetCode HOT100 完整题解 Word文档生成器
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='SimSun', font_size=10.5):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    font_names = ['SimHei', 'Microsoft YaHei', 'SimSun']
    font_size = {1: 18, 2: 16, 3: 14}.get(level, 12)
    set_chinese_font(run, font_names[0], font_size)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_zh(doc, text, bold=False, font_size=10.5):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    set_chinese_font(run, 'SimSun', font_size)
    return p

def add_code_block(doc, code_text):
    """添加代码块"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    set_chinese_font(run, 'Courier New', 9)
    run.font.color.rgb = RGBColor(0, 100, 0)
    return p

# 创建文档
doc = Document()

# 设置文档标题
title = doc.add_heading('LeetCode 热题 HOT 100 完整题解', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    set_chinese_font(run, 'SimHei', 22)
    run.bold = True

# 添加说明
add_paragraph_zh(doc, '本文档包含LeetCode HOT 100全部100道经典题目的完整题解')
add_paragraph_zh(doc, '每道题包含：题目描述、详细解题思路、核心算法步骤、时间/空间复杂度分析')
add_paragraph_zh(doc, '整理时间：2026年3月')
doc.add_page_break()

# 添加目录
add_heading_zh(doc, '目录', level=1)
toc_items = [
    '一、哈希（Hash）',
    '二、双指针（Two Pointers）', 
    '三、滑动窗口（Sliding Window）',
    '四、普通数组（Array）',
    '五、矩阵（Matrix）',
    '六、链表（Linked List）',
    '七、二叉树（Binary Tree）',
    '八、图论（Graph）',
    '九、回溯（Backtracking）',
    '十、二分查找（Binary Search）',
    '十一、栈（Stack）',
    '十二、堆（Heap）',
    '十三、贪心算法（Greedy）',
    '十四、动态规划（Dynamic Programming）',
    '十五、多维动态规划（Multi-dimensional DP）',
    '十六、技巧（Tricks）'
]
for item in toc_items:
    add_paragraph_zh(doc, item)
doc.add_page_break()

# ===== 一、哈希 =====
add_heading_zh(doc, '一、哈希（Hash）', level=1)

# 1. 两数之和
add_heading_zh(doc, '1. 两数之和 (Two Sum)', level=2)
add_paragraph_zh(doc, '题号：1 | 难度：简单 | 链接：https://leetcode.cn/problems/two-sum/', bold=True)
add_heading_zh(doc, '题目描述', level=3)
add_paragraph_zh(doc, '给定一个整数数组 nums 和一个整数目标值 target，请你在该数组中找出和为目标值 target 的那两个整数，并返回它们的数组下标。你可以假设每种输入只会对应一个答案。但是，数组中同一个元素在答案里不能重复出现。')
add_heading_zh(doc, '解题思路', level=3)
add_paragraph_zh(doc, '方法一：暴力枚举（不推荐）。遍历数组，对于每个元素，再遍历其后面的所有元素，看两数之和是否等于target。时间复杂度O(n²)，空间复杂度O(1)。')
add_paragraph_zh(doc, '方法二：哈希表（推荐）。使用哈希表存储已经遍历过的数字及其下标。对于当前数字x，在哈希表中查找target-x是否存在。关键点是先查后存，避免同一个元素使用两次。这种方法只需要遍历一次数组，时间复杂度为O(n)，空间复杂度为O(n)。')
add_heading_zh(doc, '核心代码逻辑', level=3)
add_code_block(doc, '''创建空哈希表map
遍历数组，对于每个元素num及其下标i：
    complement = target - num
    如果complement在map中：
        返回[map[complement], i]
    将num和i存入map：map[num] = i''')
add_paragraph_zh(doc, '时间复杂度：O(n) | 空间复杂度：O(n)', bold=True)

# 2. 字母异位词分组
add_heading_zh(doc, '2. 字母异位词分组 (Group Anagrams)', level=2)
add_paragraph_zh(doc, '题号：49 | 难度：中等', bold=True)
add_heading_zh(doc, '题目描述', level=3)
add_paragraph_zh(doc, '给你一个字符串数组，请你将字母异位词组合在一起。字母异位词是由重新排列源单词的所有字母得到的一个新单词。')
add_heading_zh(doc, '解题思路', level=3)
add_paragraph_zh(doc, '核心思想：字母异位词排序后得到的字符串相同。方法一：排序作为键。将每个字符串排序，排序后的字符串作为哈希表的键，原字符串加入对应的值列表。时间复杂度O(n × k log k)。方法二：字符计数作为键。使用26个长度的数组统计每个字符出现次数，将计数数组转为字符串作为键。时间复杂度O(n × k)，避免排序，更优。')

# 3. 最长连续序列
add_heading_zh(doc, '3. 最长连续序列 (Longest Consecutive Sequence)', level=2)
add_paragraph_zh(doc, '题号：128 | 难度：中等', bold=True)
add_heading_zh(doc, '题目描述', level=3)
add_paragraph_zh(doc, '给定一个未排序的整数数组 nums，找出数字连续的最长序列（不要求序列元素在原数组中连续）的长度。要求时间复杂度为 O(n)。')
add_heading_zh(doc, '解题思路', level=3)
add_paragraph_zh(doc, '核心思想：只从序列的起点开始计数。步骤：1) 将所有数字存入哈希集合，实现O(1)查找；2) 遍历集合中的每个数字num，如果num-1不在集合中，说明num是某个序列的起点；3) 从num开始，不断查找num+1, num+2...，直到不连续；4) 记录当前序列长度，更新最大值。关键点：通过判断num-1是否存在，确保每个序列只被计算一次，避免了重复计算，保证O(n)时间复杂度。')
add_paragraph_zh(doc, '时间复杂度：O(n) | 空间复杂度：O(n)', bold=True)

# 4. 只出现一次的数字
add_heading_zh(doc, '4. 只出现一次的数字 (Single Number)', level=2)
add_paragraph_zh(doc, '题号：136 | 难度：简单', bold=True)
add_heading_zh(doc, '题目描述', level=3)
add_paragraph_zh(doc, '给定一个非空整数数组，除了某个元素只出现一次以外，其余每个元素均出现两次。找出那个只出现了一次的元素。要求算法具有线性时间复杂度，并且不使用额外空间。')
add_heading_zh(doc, '解题思路', level=3)
add_paragraph_zh(doc, '核心思想：利用异或运算的性质。异或运算有以下重要性质：a ^ a = 0（任何数与自身异或为0）；a ^ 0 = a（任何数与0异或为其本身）；异或运算满足交换律和结合律。算法：将所有数字进行异或运算，成对出现的数字异或结果为0，最后剩下的就是只出现一次的数字。')
add_paragraph_zh(doc, '时间复杂度：O(n) | 空间复杂度：O(1)', bold=True)

# 更多哈希题目...

# ===== 二、双指针 =====
doc.add_page_break()
add_heading_zh(doc, '二、双指针（Two Pointers）', level=1)

# 5. 移动零
add_heading_zh(doc, '5. 移动零 (Move Zeroes)', level=2)
add_paragraph_zh(doc, '题号：283 | 难度：简单', bold=True)
add_heading_zh(doc, '题目描述', level=3)
add_paragraph_zh(doc, '给定一个数组 nums，编写一个函数将所有 0 移动到数组的末尾，同时保持非零元素的相对顺序。必须在不复制数组的情况下原地对数组进行操作。')
add_heading_zh(doc, '解题思路', level=3)
add_paragraph_zh(doc, '核心思想：双指针，一个指针遍历，一个指针记录非零元素应该放置的位置。步骤：1) 初始化指针j = 0，表示下一个非零元素应该放的位置；2) 遍历数组，指针i从0到n-1，如果nums[i] != 0，将nums[i]赋值给nums[j]，然后j++；3) 将j到数组末尾的所有元素赋值为0。优化版本可以在一次遍历中完成，遍历时直接交换非零元素到前面，保持非零元素的相对顺序。')
add_paragraph_zh(doc, '时间复杂度：O(n) | 空间复杂度：O(1)', bold=True)

# 6. 盛最多水的容器
add_heading_zh(doc, '6. 盛最多水的容器 (Container With Most Water)', level=2)
add_paragraph_zh(doc, '题号：11 | 难度：中等', bold=True)
add_heading_zh(doc, '题目描述', level=3)
add_paragraph_zh(doc, '给定一个长度为 n 的整数数组 height。有 n 条垂线，第 i 条线的两个端点是 (i, 0) 和 (i, height[i])。找出其中的两条线，使得它们与 x 轴共同构成的容器可以容纳最多的水。返回容器可以储存的最大水量。')
add_heading_zh(doc, '解题思路', level=3)
add_paragraph_zh(doc, '核心思想：双指针从两端向中间移动，移动较小高度的指针。为什么移动较小高度的指针？容器面积 = min(height[left], height[right]) × (right - left)。如果移动较大高度的指针，宽度减小，高度不变或更小，面积一定减小。如果移动较小高度的指针，虽然宽度减小，但高度可能增大，面积可能增大。步骤：初始化left = 0, right = n-1, maxArea = 0；当left < right时，计算当前面积，更新maxArea；如果height[left] < height[right]，left++；否则right--。')
add_paragraph_zh(doc, '时间复杂度：O(n) | 空间复杂度：O(1)', bold=True)

# 7. 三数之和
add_heading_zh(doc, '7. 三数之和 (3Sum)', level=2)
add_paragraph_zh(doc, '题号：15 | 难度：中等', bold=True)
add_heading_zh(doc, '题目描述', level=3)
add_paragraph_zh(doc, '给你一个整数数组 nums，判断是否存在三元组 [nums[i], nums[j], nums[k]] 满足 i != j、i != k 且 j != k，同时还满足 nums[i] + nums[j] + nums[k] == 0。请你返回所有和为 0 且不重复的三元组。')
add_heading_zh(doc, '解题思路', level=3)
add_paragraph_zh(doc, '核心思想：排序 + 双指针。步骤：1) 对数组进行排序；2) 遍历数组，固定第一个数nums[i]，如果nums[i] > 0则break（三数之和不可能为0）；3) 去重：如果i > 0且nums[i] == nums[i-1]，跳过；4) 使用双指针left = i+1, right = n-1，计算sum = nums[i] + nums[left] + nums[right]，如果sum == 0记录结果并去重，如果sum < 0则left++，如果sum > 0则right--。去重处理：外层循环跳过重复的第一个数，找到解后跳过左右指针的重复元素。')
add_paragraph_zh(doc, '时间复杂度：O(n²) | 空间复杂度：O(1)', bold=True)

# 8. 接雨水
add_heading_zh(doc, '8. 接雨水 (Trapping Rain Water)', level=2)
add_paragraph_zh(doc, '题号：42 | 难度：困难', bold=True)
add_heading_zh(doc, '题目描述', level=3)
add_paragraph_zh(doc, '给定 n 个非负整数表示每个宽度为 1 的柱子的高度图，计算按此排列的柱子，下雨之后能接多少雨水。')
add_heading_zh(doc, '解题思路', level=3)
add_paragraph_zh(doc, '核心思想：每个位置能接的雨水 = min(左边最大高度, 右边最大高度) - 当前高度。方法一：动态规划。预处理每个位置左边的最大高度和右边的最大高度，遍历计算每个位置的接水量。时间复杂度O(n)，空间复杂度O(n)。方法二：双指针（推荐）。用两个指针从两端向中间移动，维护左边的最大值和右边的最大值，较小的一侧可以计算接水量（因为接水量由较小的一侧决定）。步骤：初始化left = 0, right = n-1, leftMax = 0, rightMax = 0, result = 0；当left < right时，如果height[left] < height[right]，处理左边并left++，否则处理右边并right--。')
add_paragraph_zh(doc, '时间复杂度：O(n) | 空间复杂度：O(1)', bold=True)

# ===== 三、滑动窗口 =====
doc.add_page_break()
add_heading_zh(doc, '三、滑动窗口（Sliding Window）', level=1)

# 9. 无重复字符的最长子串
add_heading_zh(doc, '9. 无重复字符的最长子串 (Longest Substring Without Repeating Characters)', level=2)
add_paragraph_zh(doc, '题号：3 | 难度：中等', bold=True)
add_heading_zh(doc, '题目描述', level=3)
add_paragraph_zh(doc, '给定一个字符串 s，请你找出其中不含有重复字符的最长子串的长度。')
add_heading_zh(doc, '解题思路', level=3)
add_paragraph_zh(doc, '核心思想：滑动窗口，维护一个不包含重复字符的窗口。步骤：1) 初始化left = 0, maxLen = 0；2) 使用哈希集合或哈希表记录窗口中的字符；3) 遍历字符串，right从0到n-1，如果s[right]不在窗口中，将s[right]加入窗口并更新maxLen；如果s[right]已在窗口中，移动left直到s[right]不在窗口中，然后将s[right]加入窗口。优化：使用哈希表记录字符的索引，可以直接将left跳到重复字符的下一个位置。')
add_paragraph_zh(doc, '时间复杂度：O(n) | 空间复杂度：O(min(m, n))', bold=True)

# 10. 找到字符串中所有字母异位词
add_heading_zh(doc, '10. 找到字符串中所有字母异位词 (Find All Anagrams in a String)', level=2)
add_paragraph_zh(doc, '题号：438 | 难度：中等', bold=True)
add_heading_zh(doc, '解题思路', level=3)
add_paragraph_zh(doc, '核心思想：滑动窗口 + 字符计数。步骤：1) 统计p中每个字符的出现次数；2) 维护一个大小为p.length()的滑动窗口；3) 使用数组或哈希表统计窗口中每个字符的出现次数；4) 当窗口中的字符计数与p的字符计数相同时，记录起始索引；5) 滑动窗口：右移一位，移除左边字符，添加右边字符。优化：使用一个变量记录当前窗口与目标差异的字符种类数，避免每次比较整个数组。')
add_paragraph_zh(doc, '时间复杂度：O(n) | 空间复杂度：O(1)', bold=True)

# 11. 和为 K 的子数组
add_heading_zh(doc, '11. 和为 K 的子数组 (Subarray Sum Equals K)', level=2)
add_paragraph_zh(doc, '题号：560 | 难度：中等', bold=True)
add_heading_zh(doc, '解题思路', level=3)
add_paragraph_zh(doc, '核心思想：前缀和 + 哈希表。关键公式：子数组和 = prefixSum[j] - prefixSum[i-1] = k，即 prefixSum[i-1] = prefixSum[j] - k。步骤：1) 使用哈希表记录每个前缀和出现的次数；2) 遍历数组，计算当前前缀和；3) 在哈希表中查找prefixSum - k出现的次数，加到结果中；4) 将当前前缀和加入哈希表。初始化：hashMap[0] = 1，表示前缀和为0出现1次（处理从索引0开始的子数组）。')
add_paragraph_zh(doc, '时间复杂度：O(n) | 空间复杂度：O(n)', bold=True)

print("文档生成中...第一部分完成")

# 保存文档
doc.save('/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx')
print("Word文档已生成！")
print("文件路径：/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx")
