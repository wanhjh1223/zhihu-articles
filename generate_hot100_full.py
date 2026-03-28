#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LeetCode HOT100 完整题解 Word 文档生成器
包含100道题的详细题解，按类别分组
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import re

# 题目数据结构：每道题包含id, title, difficulty, link, category, description, solution
PROBLEMS = []

def add_problem(id, title, difficulty, link, category, description, solution):
    """添加一道题目的数据"""
    PROBLEMS.append({
        "id": id,
        "title": title,
        "difficulty": difficulty,
        "link": link,
        "category": category,
        "description": description,
        "solution": solution
    })

# ========== 1. 哈希 (4题) ==========
add_problem(1, "两数之和 (Two Sum)", "简单", "https://leetcode.cn/problems/two-sum/", "哈希",
"给定一个整数数组 nums 和一个整数目标值 target，请你在该数组中找出和为目标值 target 的那两个整数，并返回它们的数组下标。\n你可以假设每种输入只会对应一个答案。但是，数组中同一个元素在答案里不能重复出现。\n你可以按任意顺序返回答案。\n\n示例：\n输入：nums = [2,7,11,15], target = 9\n输出：[0,1]\n解释：因为 nums[0] + nums[1] == 9，返回 [0, 1]。",
"这道题是经典的哈希表应用题。最直观的想法是使用双重循环遍历所有可能的数对，但时间复杂度为O(n²)，效率较低。为了优化查找效率，我们可以利用哈希表的O(1)查找特性。\n\n核心思想：在遍历数组的过程中，对于当前元素num，我们需要查找target-num是否已经在之前遍历过的元素中出现过。如果出现过，说明找到了答案；如果没有，则将当前元素存入哈希表，继续遍历。\n\n为什么选择'先查后存'的顺序？这是为了避免同一个元素被使用两次。如果先存入再查找，当数组中存在target/2这个元素时，可能会错误地认为找到了答案（实际上是同一个元素）。\n\n详细步骤：\n1. 创建一个空的哈希表，用于存储已经遍历过的数字及其下标\n2. 从左到右遍历数组中的每个元素num及其下标i\n3. 计算complement = target - num，这是在数组中需要找到的另一个数\n4. 在哈希表中查找complement是否存在\n   - 如果存在，说明找到了答案，返回[complement的下标, i]\n   - 如果不存在，将num和i存入哈希表，继续遍历\n\n时间复杂度分析：O(n)，其中n是数组长度。我们只需要遍历数组一次，哈希表的插入和查找操作都是O(1)。\n空间复杂度分析：O(n)，最坏情况下哈希表需要存储n-1个元素。")

add_problem(49, "字母异位词分组 (Group Anagrams)", "中等", "https://leetcode.cn/problems/group-anagrams/", "哈希",
"给你一个字符串数组，请你将字母异位词组合在一起。可以按任意顺序返回结果列表。\n字母异位词是由重新排列源单词的所有字母得到的一个新单词。\n\n示例：\n输入: strs = [\"eat\", \"tea\", \"tan\", \"ate\", \"nat\", \"bat\"]\n输出: [[\"bat\"],[\"nat\",\"tan\"],[\"ate\",\"eat\",\"tea\"]]",
"这道题的关键在于如何识别字母异位词。两个字符串是字母异位词，当且仅当它们包含相同的字符且每个字符出现的次数相同。基于这一性质，我们可以设计两种方法来作为哈希表的键。\n\n方法一：排序作为键\n将每个字符串的字符进行排序，排序后的字符串作为哈希表的键。所有字母异位词排序后得到的结果相同，因此会被分到同一组。例如，\"eat\"、\"tea\"、\"ate\"排序后都是\"aet\"，因此它们会被分到同一组。\n\n时间复杂度：O(n × k log k)，其中n是字符串个数，k是字符串的最大长度。\n空间复杂度：O(n × k)，需要存储所有字符串的排序结果。\n\n方法二：字符计数作为键（更优）\n使用一个长度为26的数组统计每个字符的出现次数，然后将这个计数数组转换为字符串作为键。这种方法避免了排序操作。\n\n时间复杂度：O(n × k)，每个字符串只需遍历一次来计算字符计数。\n空间复杂度：O(n × k)\n\n方法二是更优的解法，因为它的时间复杂度更低，且实际运行效率更高。")

add_problem(128, "最长连续序列 (Longest Consecutive Sequence)", "中等", "https://leetcode.cn/problems/longest-consecutive-sequence/", "哈希",
"给定一个未排序的整数数组 nums ，找出数字连续的最长序列（不要求序列元素在原数组中连续）的长度。\n请你设计并实现时间复杂度为 O(n) 的算法解决此问题。\n\n示例：\n输入：nums = [100,4,200,1,3,2]\n输出：4\n解释：最长数字连续序列是 [1, 2, 3, 4]。它的长度为 4。",
"这道题要求O(n)时间复杂度，说明不能对数组进行排序（排序的时间复杂度是O(n log n)）。我们需要利用哈希集合来实现O(1)的查找操作。\n\n核心思想：只从序列的起点开始计数。\n\n关键观察：如果一个数num是某个连续序列的起点，那么num-1一定不在数组中。利用这一性质，我们可以避免重复计算序列长度。\n\n详细步骤：\n1. 将所有数字存入哈希集合，实现O(1)查找\n2. 遍历集合中的每个数字num：\n   - 如果num-1不在集合中，说明num是某个连续序列的起点\n   - 从num开始，不断查找num+1、num+2...是否在集合中，直到不连续\n   - 记录当前序列的长度，更新最大值\n   - 如果num-1在集合中，跳过该数字（它会在作为序列起点时被计算）\n\n为什么这种方法是O(n)？虽然看起来有两层循环，但每个数字最多被访问两次：一次作为起点判断，一次在序列中。因此总时间复杂度是O(n)。\n\n时间复杂度：O(n)，每个数字最多被访问两次\n空间复杂度：O(n)，哈希集合存储所有数字")

add_problem(136, "只出现一次的数字 (Single Number)", "简单", "https://leetcode.cn/problems/single-number/", "哈希",
"给定一个非空整数数组，除了某个元素只出现一次以外，其余每个元素均出现两次。找出那个只出现了一次的元素。\n说明：你的算法应该具有线性时间复杂度。 你可以不使用额外空间来实现吗？\n\n示例：\n输入: [2,2,1]\n输出: 1",
"这道题是一道经典的位运算题目。题目要求线性时间复杂度和常数空间复杂度，这意味着我们不能使用哈希表等额外数据结构。\n\n核心思想：利用异或运算的性质。\n\n异或运算的性质：\n1. a ^ a = 0（任何数与自身异或结果为0）\n2. a ^ 0 = a（任何数与0异或结果为自身）\n3. 异或运算满足交换律和结合律：a ^ b ^ a = a ^ a ^ b = b\n\n算法：将数组中所有数字进行异或运算。根据上述性质，成对出现的数字异或结果为0，最后剩下的就是只出现一次的数字。\n\n例如：[2,2,1]\n2 ^ 2 ^ 1 = 0 ^ 1 = 1\n\n例如：[4,1,2,1,2]\n4 ^ 1 ^ 2 ^ 1 ^ 2 = 4 ^ (1 ^ 1) ^ (2 ^ 2) = 4 ^ 0 ^ 0 = 4\n\n时间复杂度：O(n)，只需要遍历数组一次\n空间复杂度：O(1)，只使用了一个变量来存储异或结果\n\n这是一个非常巧妙的解法，充分利用了异或运算的特性，完美满足了题目的要求。")

# ========== 2. 双指针 (4题) ==========
add_problem(283, "移动零 (Move Zeroes)", "简单", "https://leetcode.cn/problems/move-zeroes/", "双指针",
"给定一个数组 nums，编写一个函数将所有 0 移动到数组的末尾，同时保持非零元素的相对顺序。\n请注意 ，必须在不复制数组的情况下原地对数组进行操作。\n\n示例：\n输入: nums = [0,1,0,3,12]\n输出: [1,3,12,0,0]",
"这道题考察双指针技巧。我们需要在原数组上操作，将所有非零元素移动到数组前面，同时保持它们的相对顺序，最后将剩余位置填充为0。\n\n核心思想：使用双指针，一个指针遍历数组，另一个指针记录非零元素应该放置的位置。\n\n详细步骤：\n1. 初始化指针j = 0，表示下一个非零元素应该放置的位置\n2. 使用指针i遍历数组：\n   - 如果nums[i] != 0：\n     - 将nums[i]赋值给nums[j]\n     - j++（指向下一个应该放置的位置）\n3. 遍历完成后，将数组中从j开始到末尾的所有元素赋值为0\n\n为什么这样能保持相对顺序？因为我们是从左到右遍历，非零元素按照它们在原数组中出现的顺序被依次放置到前面。\n\n优化版本（一次遍历）：\n可以在遍历过程中直接交换非零元素到前面，这样可以省去最后填充0的步骤。\n\n时间复杂度：O(n)，每个元素最多被访问两次\n空间复杂度：O(1)，只使用了常数额外空间")

add_problem(11, "盛最多水的容器 (Container With Most Water)", "中等", "https://leetcode.cn/problems/container-with-most-water/", "双指针",
"给定一个长度为 n 的整数数组 height 。有 n 条垂线，第 i 条线的两个端点是 (i, 0) 和 (i, height[i]) 。\n找出其中的两条线，使得它们与 x 轴共同构成的容器可以容纳最多的水。\n返回容器可以储存的最大水量。\n说明：你不能倾斜容器。\n\n示例：\n输入：[1,8,6,2,5,4,8,3,7]\n输出：49",
"这道题是双指针的经典应用。我们需要找到两条线，使得它们之间的容器能容纳最多的水。\n\n容器的面积计算：面积 = min(height[left], height[right]) × (right - left)\n\n核心思想：双指针从两端向中间移动，每次移动较小高度的指针。\n\n为什么移动较小高度的指针？\n假设height[left] < height[right]，此时容器的面积由height[left]决定。如果我们移动right指针，宽度减小，高度不变或更小，因此面积一定减小。但如果我们移动left指针，虽然宽度减小，但高度可能增大，面积可能增大。因此，应该移动较小高度的指针，才有可能找到更大的面积。\n\n详细步骤：\n1. 初始化left = 0, right = n-1, maxArea = 0\n2. 当left < right时：\n   - 计算当前面积：area = min(height[left], height[right]) × (right - left)\n   - 更新maxArea = max(maxArea, area)\n   - 如果height[left] < height[right]，left++；否则right--\n\n时间复杂度：O(n)，两个指针最多各遍历一次\n空间复杂度：O(1)，只使用了常数额外空间")

add_problem(15, "三数之和 (3Sum)", "中等", "https://leetcode.cn/problems/3sum/", "双指针",
"给你一个整数数组 nums ，判断是否存在三元组 [nums[i], nums[j], nums[k]] 满足 i != j、i != k 且 j != k ，同时还满足 nums[i] + nums[j] + nums[k] == 0 。\n请你返回所有和为 0 且不重复的三元组。\n注意：答案中不可以包含重复的三元组。\n\n示例：\n输入：nums = [-1,0,1,2,-1,-4]\n输出：[[-1,-1,2],[-1,0,1]]",
"这道题是'两数之和'的升级版，要求找出三个数使得它们的和为0。最暴力的解法是三重循环，时间复杂度为O(n³)。我们需要更优的解法。\n\n核心思想：排序 + 双指针\n\n为什么先排序？排序后可以使用双指针技巧，将O(n³)降到O(n²)，同时也方便去重。\n\n详细步骤：\n1. 对数组进行排序\n2. 遍历数组，固定第一个数nums[i]：\n   - 如果nums[i] > 0，直接break（因为数组已排序，三个正数之和不可能为0）\n   - 去重：如果i > 0且nums[i] == nums[i-1]，跳过（避免重复的三元组）\n   - 使用双指针left = i+1, right = n-1：\n     - 计算sum = nums[i] + nums[left] + nums[right]\n     - 如果sum == 0：记录结果，left++，right--，同时去重\n     - 如果sum < 0：left++（需要更大的数）\n     - 如果sum > 0：right--（需要更小的数）\n\n去重处理的三层：\n1. 外层循环跳过重复的nums[i]\n2. 找到解后，跳过重复的nums[left]\n3. 找到解后，跳过重复的nums[right]\n\n时间复杂度：O(n²)，排序O(n log n)，双指针遍历O(n²)\n空间复杂度：O(1)或O(log n)（排序的栈空间）")

add_problem(42, "接雨水 (Trapping Rain Water)", "困难", "https://leetcode.cn/problems/trapping-rain-water/", "双指针",
"给定 n 个非负整数表示每个宽度为 1 的柱子的高度图，计算按此排列的柱子，下雨之后能接多少雨水。\n\n示例：\n输入：height = [0,1,0,2,1,0,1,3,2,1,2,1]\n输出：6",
"这道题是双指针的经典难题。每个位置能接的雨水量取决于其左边和右边最高柱子的较小值。\n\n核心公式：每个位置能接的雨水 = min(左边最大高度, 右边最大高度) - 当前高度\n\n方法一：动态规划\n预处理每个位置左边的最大高度和右边的最大高度。\n1. leftMax[i]表示从0到i的最大高度\n2. rightMax[i]表示从i到n-1的最大高度\n3. 遍历计算每个位置的接水量\n时间复杂度：O(n)，空间复杂度：O(n)\n\n方法二：双指针（推荐）\n用两个指针从两端向中间移动，维护左边的最大值和右边的最大值。\n\n关键观察：\n- 如果leftMax < rightMax，左边可以计算接水量（因为右边一定有更高的柱子）\n- 如果leftMax >= rightMax，右边可以计算接水量\n\n详细步骤：\n1. 初始化left = 0, right = n-1, leftMax = 0, rightMax = 0, result = 0\n2. 当left < right时：\n   - 如果height[left] < height[right]：\n     - 如果height[left] >= leftMax，更新leftMax = height[left]\n     - 否则，result += leftMax - height[left]\n     - left++\n   - 否则：\n     - 如果height[right] >= rightMax，更新rightMax = height[right]\n     - 否则，result += rightMax - height[right]\n     - right--\n\n时间复杂度：O(n)\n空间复杂度：O(1)，只使用了常数额外空间")

# ========== 3. 滑动窗口 (3题) ==========
add_problem(3, "无重复字符的最长子串 (Longest Substring Without Repeating Characters)", "中等", "https://leetcode.cn/problems/longest-substring-without-repeating-characters/", "滑动窗口",
"给定一个字符串 s ，请你找出其中不含有重复字符的 最长子串 的长度。\n\n示例：\n输入: s = \"abcabcbb\"\n输出: 3 \n解释: 因为无重复字符的最长子串是 \"abc\"，所以其长度为 3。",
"这道题是滑动窗口的经典题目。我们需要维护一个不包含重复字符的窗口，不断扩展右边界，当发现重复字符时收缩左边界。\n\n核心思想：滑动窗口，使用双指针维护一个不包含重复字符的窗口[left, right]。\n\n详细步骤：\n1. 初始化left = 0, maxLen = 0\n2. 使用哈希集合或哈希表记录窗口中的字符\n3. 遍历字符串，right从0到n-1：\n   - 如果s[right]不在窗口中：\n     - 将s[right]加入窗口\n     - 更新maxLen = max(maxLen, right - left + 1)\n   - 如果s[right]已在窗口中：\n     - 移动left，直到s[right]不在窗口中\n     - 将s[right]加入窗口\n\n优化版本：使用哈希表记录字符的索引\n如果使用哈希表记录每个字符最后出现的位置，当发现重复字符时，可以直接将left跳到重复字符的下一个位置，而不需要逐个移动。\n\n注意使用max(left, ...)是为了防止left回退。\n\n时间复杂度：O(n)，每个字符最多被访问两次\n空间复杂度：O(min(m, n))，其中m是字符集大小")

add_problem(438, "找到字符串中所有字母异位词 (Find All Anagrams in a String)", "中等", "https://leetcode.cn/problems/find-all-anagrams-in-a-string/", "滑动窗口",
"给定两个字符串 s 和 p，找到 s 中所有 p 的 异位词 的子串，返回这些子串的起始索引。不考虑答案输出的顺序。\n\n示例：\n输入: s = \"cbaebabacd\", p = \"abc\"\n输出: [0,6]\n解释:\n起始索引等于 0 的子串是 \"cba\"，它是 \"abc\" 的异位词。\n起始索引等于 6 的子串是 \"bac\"，它是 \"abc\" 的异位词。",
"这道题是滑动窗口加字符计数的应用。我们需要在字符串s中找到所有与p是字母异位词的子串。\n\n核心思想：滑动窗口 + 字符计数\n\n什么是字母异位词？两个字符串包含相同的字符且每个字符出现的次数相同。因此，我们可以使用字符计数来判断两个字符串是否是字母异位词。\n\n详细步骤：\n1. 统计p中每个字符的出现次数，存入数组pCount（大小为26，对应26个字母）\n2. 初始化一个大小为p.length()的滑动窗口，统计窗口中每个字符的出现次数\n3. 使用数组windowCount记录窗口中的字符计数\n4. 滑动窗口遍历字符串s：\n   - 右移窗口：添加右边的新字符到windowCount\n   - 移除左边滑出的字符\n   - 比较windowCount和pCount，如果相同，说明找到了一个字母异位词，记录起始索引\n\n优化：使用一个变量diff记录当前窗口与目标差异的字符种类数，避免每次比较整个数组。\n\n时间复杂度：O(n)，其中n是s的长度\n空间复杂度：O(1)，固定大小的数组（26个字母）")

add_problem(560, "和为 K 的子数组 (Subarray Sum Equals K)", "中等", "https://leetcode.cn/problems/subarray-sum-equals-k/", "滑动窗口",
"给你一个整数数组 nums 和一个整数 k ，请你统计并返回 该数组中和为 k 的子数组的个数 。\n子数组是数组中元素的连续非空序列。\n\n示例：\n输入：nums = [1,1,1], k = 2\n输出：2",
"这道题不能简单地使用滑动窗口，因为数组中可能包含负数，这意味着扩大窗口不一定增加和，缩小窗口不一定减少和。\n\n核心思想：前缀和 + 哈希表\n\n关键公式：\n- 子数组和 = prefixSum[j] - prefixSum[i-1] = k\n- 即 prefixSum[i-1] = prefixSum[j] - k\n\n详细步骤：\n1. 使用哈希表记录每个前缀和出现的次数，key是前缀和，value是出现次数\n2. 初始化hashMap[0] = 1，表示前缀和为0出现1次（处理从索引0开始的子数组）\n3. 遍历数组，维护当前前缀和prefixSum：\n   - 在哈希表中查找prefixSum - k出现的次数，加到结果中\n   - 将当前prefixSum加入哈希表（次数+1）\n\n为什么hashMap[0] = 1？\n考虑从数组开头开始的子数组，例如nums = [1, 2, 3], k = 6。当遍历到索引2时，prefixSum = 6，我们需要找到prefixSum - k = 0。如果不初始化hashMap[0] = 1，就会漏掉这个解。\n\n示例分析：\nnums = [1, 1, 1], k = 2\ni=0: prefixSum=1, 查找1-2=-1(不存在), hashMap={0:1, 1:1}\ni=1: prefixSum=2, 查找2-2=0(存在1次), result=1, hashMap={0:1, 1:1, 2:1}\ni=2: prefixSum=3, 查找3-2=1(存在1次), result=2, hashMap={0:1, 1:1, 2:1, 3:1}\n最终result=2\n\n时间复杂度：O(n)\n空间复杂度：O(n)，哈希表存储前缀和")

# ========== 4. 普通数组 (5题) ==========
add_problem(53, "最大子数组和 (Maximum Subarray)", "中等", "https://leetcode.cn/problems/maximum-subarray/", "普通数组",
"给你一个整数数组 nums ，请你找出一个具有最大和的连续子数组（子数组最少包含一个元素），返回其最大和。\n子数组 是数组中的一个连续部分。\n\n示例：\n输入：nums = [-2,1,-3,4,-1,2,1,-5,4]\n输出：6\n解释：连续子数组 [4,-1,2,1] 的和最大，为 6 。",
"这道题是著名的Kadane算法（卡丹算法）的应用，是动态规划的经典入门题目。\n\n核心思想：动态规划。对于每个位置，我们要么从当前元素重新开始一个新的子数组，要么将当前元素加入到之前的子数组中。\n\n状态定义：dp[i]表示以nums[i]结尾的最大子数组和\n\n状态转移方程：\ndp[i] = max(nums[i], dp[i-1] + nums[i])\n解释：要么从当前元素重新开始（nums[i]），要么延续前面的子数组（dp[i-1] + nums[i]）\n\n空间优化：\n我们不需要存储整个dp数组，只需要维护当前值和最大值即可。\n\n详细步骤：\n1. 初始化currentSum = nums[0], maxSum = nums[0]\n2. 从i=1遍历到n-1：\n   - currentSum = max(nums[i], currentSum + nums[i])\n   - maxSum = max(maxSum, currentSum)\n\n为什么这样是正确的？\n对于每个位置i，以它结尾的最大子数组和只取决于两个选择：\n1. 只包含nums[i]本身\n2. 包含以i-1结尾的最大子数组 + nums[i]\n我们取这两个选择中的较大值，同时维护全局最大值。\n\n时间复杂度：O(n)，只需要遍历数组一次\n空间复杂度：O(1)，只使用了常数额外空间")

add_problem(56, "合并区间 (Merge Intervals)", "中等", "https://leetcode.cn/problems/merge-intervals/", "普通数组",
"以数组 intervals 表示若干个区间的集合，其中单个区间为 intervals[i] = [starti, endi] 。请你合并所有重叠的区间，并返回一个不重叠的区间数组，该数组需恰好覆盖输入中的所有区间。\n\n示例：\n输入：intervals = [[1,3],[2,6],[8,10],[15,18]]\n输出：[[1,6],[8,10],[15,18]]\n解释：区间 [1,3] 和 [2,6] 重叠, 将它们合并为 [1,6]。",
"这道题是区间问题的经典题目。合并重叠区间的关键是先对区间进行排序，然后逐个检查是否重叠。\n\n核心思想：排序 + 贪心\n\n为什么先排序？按区间的左端点排序后，重叠的区间会相邻出现，方便我们逐个处理。\n\n详细步骤：\n1. 按区间的左端点升序排序\n2. 初始化结果列表，将第一个区间加入结果\n3. 遍历剩余的区间：\n   - 如果当前区间的左端点 <= 结果列表最后一个区间的右端点：\n     - 说明有重叠，合并区间\n     - 更新最后一个区间的右端点为max(当前右端点, 最后一个右端点)\n   - 否则：\n     - 无重叠，将当前区间加入结果列表\n\n合并逻辑示例：\n已排序区间：[[1,3], [2,6], [8,10], [15,18]]\n- 结果初始：[[1,3]]\n- 处理[2,6]：2 <= 3，有重叠，合并为[1, max(3,6)] = [1,6]\n- 结果：[[1,6]]\n- 处理[8,10]：8 > 6，无重叠，加入\n- 结果：[[1,6], [8,10]]\n- 处理[15,18]：15 > 10，无重叠，加入\n- 最终结果：[[1,6], [8,10], [15,18]]\n\n时间复杂度：O(n log n)，主要是排序的时间\n空间复杂度：O(log n)或O(n)，取决于排序的实现")

add_problem(189, "轮转数组 (Rotate Array)", "中等", "https://leetcode.cn/problems/rotate-array/", "普通数组",
"给定一个整数数组 nums，将数组中的元素向右轮转 k 个位置，其中 k 是非负数。\n\n示例：\n输入: nums = [1,2,3,4,5,6,7], k = 3\n输出: [5,6,7,1,2,3,4]\n解释:\n向右轮转 1 步: [7,1,2,3,4,5,6]\n向右轮转 2 步: [6,7,1,2,3,4,5]\n向右轮转 3 步: [5,6,7,1,2,3,4]",
"这道题有多种解法，我们需要找到最优的空间复杂度解法。\n\n方法一：使用额外数组（不推荐）\n创建一个新数组，将原数组的元素放到正确的位置。\n时间复杂度：O(n)，空间复杂度：O(n)\n\n方法二：数组反转（推荐）\n利用数组反转的性质来实现轮转。这是一个非常巧妙的技巧。\n\n核心思想：先整体反转，再分别反转前后两部分。\n\n以[1,2,3,4,5,6,7], k=3为例：\n1. 整体反转：[7,6,5,4,3,2,1]\n2. 反转前k个：[5,6,7,4,3,2,1]\n3. 反转剩余部分：[5,6,7,1,2,3,4]\n\n数学原理：\n向右轮转k位，相当于将后k个元素移到前面，前n-k个元素移到后面。通过三次反转，我们实现了这一效果。\n\n详细步骤：\n1. k = k % n（处理k大于n的情况）\n2. 反转整个数组\n3. 反转[0, k-1]\n4. 反转[k, n-1]\n\n时间复杂度：O(n)，三次反转各需要O(n)\n空间复杂度：O(1)，原地操作，只使用了常数额外空间")

add_problem(238, "除自身以外数组的乘积 (Product of Array Except Self)", "中等", "https://leetcode.cn/problems/product-of-array-except-self/", "普通数组",
"给你一个整数数组 nums，返回 数组 answer ，其中 answer[i] 等于 nums 中除 nums[i] 之外其余各元素的乘积 。\n题目数据 保证 数组 nums之中任意元素的全部前缀元素和后缀的乘积都在 32 位 整数范围内。\n请**不要使用除法**，且在 O(n) 时间复杂度内完成此题。\n\n示例：\n输入: nums = [1,2,3,4]\n输出: [24,12,8,6]",
"这道题乍一看很简单，直接用总乘积除以每个元素即可。但题目要求不能使用除法，而且如果数组中有0，除法方案也会失效。\n\n核心思想：前缀积 × 后缀积\n\n对于answer[i]，它等于nums[0...i-1]的乘积 × nums[i+1...n-1]的乘积。\n- nums[0...i-1]的乘积称为前缀积\n- nums[i+1...n-1]的乘积称为后缀积\n\n详细步骤：\n1. 创建结果数组answer\n2. 第一次遍历（从左到右）：计算前缀积\n   - answer[i] = nums[0] × nums[1] × ... × nums[i-1]\n   - 即answer[i] = answer[i-1] × nums[i-1]\n   - answer[0]初始化为1（空积）\n3. 第二次遍历（从右到左）：计算后缀积并乘到结果中\n   - 维护一个变量suffix，表示从右往左的后缀积\n   - answer[i] *= suffix\n   - suffix *= nums[i]\n\n示例分析：\nnums = [1, 2, 3, 4]\n\n第一次遍历后（前缀积）：\nanswer = [1, 1, 2, 6]\n\n第二次遍历：\ni=3: answer[3]=6×1=6, suffix=1×4=4\ni=2: answer[2]=2×4=8, suffix=4×3=12\ni=1: answer[1]=1×12=12, suffix=12×2=24\ni=0: answer[0]=1×24=24, suffix=24×1=24\n\n最终结果：[24, 12, 8, 6]\n\n时间复杂度：O(n)，两次遍历\n空间复杂度：O(1)，输出数组不算额外空间")

add_problem(41, "缺失的第一个正数 (First Missing Positive)", "困难", "https://leetcode.cn/problems/first-missing-positive/", "普通数组",
"给你一个未排序的整数数组 nums ，请你找出其中没有出现的最小的正整数。\n请你实现时间复杂度为 O(n) 并且只使用常数级别额外空间的解决方案。\n\n示例：\n输入：nums = [1,2,0]\n输出：3",
"这道题要求在O(n)时间复杂度和O(1)空间复杂度内完成，是一个非常有挑战性的原地哈希问题。\n\n核心思想：原地哈希。利用数组本身作为哈希表，将每个数放到它应该在的位置上。\n\n关键观察：\n- 我们要找的是缺失的最小正整数，范围是[1, n+1]，其中n是数组长度\n- 对于一个长度为n的数组，如果1到n都出现了，那么答案就是n+1；否则答案是1到n中缺失的最小数\n- 因此，我们可以忽略所有负数、0和大于n的数\n\n详细步骤：\n1. 遍历数组，对于每个数nums[i]：\n   - 如果nums[i]在[1, n]范围内，且nums[i] != nums[nums[i]-1]，则交换\n   - 即将nums[i]放到索引为nums[i]-1的位置上\n   - 交换后，继续检查新的nums[i]，直到不满足条件\n2. 再次遍历数组，找到第一个nums[i] != i+1的位置，返回i+1\n3. 如果1到n都存在，返回n+1\n\n为什么这样有效？\n通过将每个数放到正确的位置，我们实际上是在原地构建了一个哈希表。正确位置上的数满足nums[i] = i+1。\n\n时间复杂度：O(n)，虽然看起来有嵌套循环，但每个元素最多被交换两次\n空间复杂度：O(1)，原地操作")

# ========== 5. 矩阵 (4题) ==========
add_problem(73, "矩阵置零 (Set Matrix Zeroes)", "中等", "https://leetcode.cn/problems/set-matrix-zeroes/", "矩阵",
"给定一个 m x n 的矩阵，如果一个元素为 0 ，则将其所在行和列的所有元素都设为 0 。请使用 原地 算法。\n\n示例：\n输入：matrix = [[1,1,1],[1,0,1],[1,1,1]]\n输出：[[1,0,1],[0,0,0],[1,0,1]]",
"这道题的关键在于使用原地算法，即不使用额外的矩阵空间。\n\n方法一：使用标记数组（空间复杂度O(m+n)）\n使用两个数组分别记录哪些行和列需要置零。这不是最优解。\n\n方法二：使用矩阵的第一行和第一列作为标记（空间复杂度O(1)）\n\n核心思想：用矩阵的第一行和第一列来记录哪些行和列需要置零。\n\n详细步骤：\n1. 使用两个变量firstRowZero和firstColZero记录第一行和第一列是否需要置零\n2. 遍历矩阵（从第二行第二列开始），如果matrix[i][j] == 0：\n   - 将matrix[i][0]设为0（标记第i行需要置零）\n   - 将matrix[0][j]设为0（标记第j列需要置零）\n3. 再次遍历矩阵（从第二行第二列开始），根据第一行和第一列的标记置零\n4. 根据firstRowZero和firstColZero决定是否置零第一行和第一列\n\n为什么这样有效？\n第一行和第一列充当了'标记数组'的角色，用O(1)的额外空间实现了标记功能。需要注意的是，第一行和第一列本身是否置零需要单独处理。\n\n时间复杂度：O(m × n)，需要遍历矩阵多次\n空间复杂度：O(1)，只使用了常数额外空间")

add_problem(54, "螺旋矩阵 (Spiral Matrix)", "中等", "https://leetcode.cn/problems/spiral-matrix/", "矩阵",
"给你一个 m 行 n 列的矩阵 matrix ，请按照 顺时针螺旋顺序 ，返回矩阵中的所有元素。\n\n示例：\n输入：matrix = [[1,2,3],[4,5,6],[7,8,9]]\n输出：[1,2,3,6,9,8,7,4,5]",
"这道题考察对矩阵边界的控制。按照顺时针螺旋顺序遍历矩阵，即：右→下→左→上→右→...的循环。\n\n核心思想：维护四个边界（上、下、左、右），按照顺序遍历，每遍历完一条边就收缩对应的边界。\n\n详细步骤：\n1. 初始化四个边界：top=0, bottom=m-1, left=0, right=n-1\n2. 当top <= bottom且left <= right时循环：\n   - 从左到右遍历上边：从left到right，遍历完成后top++\n   - 从上到下遍历右边：从top到bottom，遍历完成后right--\n   - 如果top <= bottom，从右到左遍历下边：从right到left，遍历完成后bottom--\n   - 如果left <= right，从下到上遍历左边：从bottom到top，遍历完成后left++\n\n为什么需要检查top <= bottom和left <= right？\n当矩阵是长方形（如1×n或m×1）时，在遍历完某些边后，边界可能已经交叉，需要停止循环。\n\n时间复杂度：O(m × n)，每个元素被访问一次\n空间复杂度：O(1)，不包括输出数组")

add_problem(48, "旋转图像 (Rotate Image)", "中等", "https://leetcode.cn/problems/rotate-image/", "矩阵",
"给定一个 n × n 的二维矩阵 matrix 表示一个图像。请你将图像顺时针旋转 90 度。\n你必须在 原地 旋转图像，这意味着你需要直接修改输入的二维矩阵。请不要 使用另一个矩阵来旋转图像。\n\n示例：\n输入：matrix = [[1,2,3],[4,5,6],[7,8,9]]\n输出：[[7,4,1],[8,5,2],[9,6,3]]",
"这道题要求原地旋转矩阵，不能创建新的矩阵。\n\n核心思想：先转置矩阵，再反转每一行。\n\n顺时针旋转90度的数学关系：\n- 对于元素matrix[i][j]，旋转后的位置是matrix[j][n-1-i]\n- 这等价于：先转置（matrix[i][j] -> matrix[j][i]），再反转每一行\n\n详细步骤：\n1. 转置矩阵：遍历对角线上方的元素，与对角线对称的元素交换\n   - for i in range(n):\n   -   for j in range(i+1, n):\n   -     swap(matrix[i][j], matrix[j][i])\n2. 反转每一行：对每一行进行反转\n   - for i in range(n):\n   -   reverse(matrix[i])\n\n时间复杂度：O(n²)，需要遍历整个矩阵\n空间复杂度：O(1)，原地操作")

add_problem(240, "搜索二维矩阵 II (Search a 2D Matrix II)", "中等", "https://leetcode.cn/problems/search-a-2d-matrix-ii/", "矩阵",
"编写一个高效的算法来搜索 m x n 矩阵 matrix 中的一个目标值 target 。该矩阵具有以下特性：\n- 每行的元素从左到右升序排列。\n- 每列的元素从上到下升序排列。\n\n示例：\n输入：matrix = [[1,4,7,11,15],[2,5,8,12,19],[3,6,9,16,22],[10,13,14,17,24],[18,21,23,26,30]], target = 5\n输出：true",
"这道题的关键在于利用矩阵的有序性来设计高效的搜索算法。暴力搜索是O(m×n)，但我们可以做得更好。\n\n核心思想：从右上角或左下角开始搜索，利用有序性排除行或列。\n\n详细步骤（从右上角开始）：\n1. 初始化row = 0, col = n-1（右上角）\n2. 当row < m且col >= 0时循环：\n   - 如果matrix[row][col] == target，返回true\n   - 如果matrix[row][col] > target，col--（当前列的所有元素都大于target，排除这一列）\n   - 如果matrix[row][col] < target，row++（当前行的所有元素都小于target，排除这一行）\n3. 如果循环结束还没找到，返回false\n\n为什么从右上角开始？\n从右上角开始，我们有一个性质：左边的元素都小于当前元素，下边的元素都大于当前元素。这样每次比较都可以排除一行或一列。\n\n时间复杂度：O(m + n)，每次比较排除一行或一列\n空间复杂度：O(1)，只使用了常数额外空间")

# ========== 6. 链表 (12题) ==========
add_problem(160, "相交链表 (Intersection of Two Linked Lists)", "简单", "https://leetcode.cn/problems/intersection-of-two-linked-lists/", "链表",
"给你两个单链表的头节点 headA 和 headB ，请你找出并返回两个单链表相交的起始节点。如果两个链表不存在相交节点，返回 null 。\n\n示例：\n输入：intersectVal = 8, listA = [4,1,8,4,5], listB = [5,6,1,8,4,5], skipA = 2, skipB = 3\n输出：Intersected at '8'",
"这道题有多种解法，包括使用哈希表、计算长度差等。最优的解法是双指针法。\n\n核心思想：双指针。让两个指针分别遍历两个链表，当到达末尾时从另一个链表的头继续遍历。如果两个链表相交，两个指针最终会在相交节点相遇。\n\n为什么这样有效？\n设链表A的长度为a+c，链表B的长度为b+c，其中c是公共部分的长度。\n- 指针A走过的路程：a + c + b\n- 指针B走过的路程：b + c + a\n两者相等！所以两个指针一定会同时到达相交节点（或同时到达null）。\n\n时间复杂度：O(a + b)，其中a和b分别是两个链表的长度\n空间复杂度：O(1)，只使用了两个指针")

add_problem(206, "反转链表 (Reverse Linked List)", "简单", "https://leetcode.cn/problems/reverse-linked-list/", "链表",
"给你单链表的头节点 head ，请你反转链表，并返回反转后的链表。\n\n示例：\n输入：head = [1,2,3,4,5]\n输出：[5,4,3,2,1]",
"这道题是链表的基础操作题，是理解链表指针操作的经典题目。\n\n核心思想：迭代法，逐个反转指针的方向。\n\n详细步骤：\n1. 初始化prev = null, curr = head\n2. 当curr不为null时循环：\n   - 保存next = curr.next（暂存下一个节点）\n   - 反转curr.next = prev（当前节点指向prev）\n   - 移动prev = curr（prev前移）\n   - 移动curr = next（curr前移）\n3. 返回prev（新的头节点）\n\n为什么需要保存next？\n在反转curr.next之前，我们需要先保存curr.next，否则反转后就找不到下一个节点了。\n\n时间复杂度：O(n)，需要遍历链表一次\n空间复杂度：O(1)，只使用了三个指针")

add_problem(234, "回文链表 (Palindrome Linked List)", "简单", "https://leetcode.cn/problems/palindrome-linked-list/", "链表",
"给你一个单链表的头节点 head ，请你判断该链表是否为回文链表。如果是，返回 true ；否则，返回 false 。\n\n示例：\n输入：head = [1,2,2,1]\n输出：true",
"这道题要求判断链表是否为回文，最优解需要O(n)时间复杂度和O(1)空间复杂度。\n\n核心思想：找到中点 + 反转后半部分 + 比较\n\n详细步骤：\n1. 使用快慢指针找到链表的中点\n   - slow每次走一步，fast每次走两步\n   - 当fast到达末尾时，slow到达中点\n2. 反转链表的后半部分\n3. 比较前半部分和反转后的后半部分\n4. （可选）恢复链表\n\n快慢指针原理：\n- 当fast到达末尾时（fast == null或fast.next == null）\n- slow正好到达中间位置（奇数长度）或中间偏左（偶数长度）\n\n时间复杂度：O(n)\n空间复杂度：O(1)，只使用了几个指针")

add_problem(141, "环形链表 (Linked List Cycle)", "简单", "https://leetcode.cn/problems/linked-list-cycle/", "链表",
"给你一个链表的头节点 head ，判断链表中是否有环。\n\n如果链表中有某个节点，可以通过连续跟踪 next 指针再次到达，则链表中存在环。 为了表示给定链表中的环，评测系统内部使用整数 pos 来表示链表尾连接到链表中的位置（索引从 0 开始）。注意：pos 不作为参数进行传递 。仅仅是为了标识链表的实际情况。\n\n如果链表中存在环 ，则返回 true 。 否则，返回 false 。\n\n示例：\n输入：head = [3,2,0,-4], pos = 1\n输出：true\n解释：链表中有一个环，其尾部连接到第二个节点。",
"这道题是经典的Floyd判圈算法（龟兔赛跑算法）的应用。\n\n核心思想：快慢指针。如果有环，快指针最终会追上慢指针。\n\n详细步骤：\n1. 初始化slow = head, fast = head\n2. 当fast不为null且fast.next不为null时循环：\n   - slow = slow.next（慢指针走一步）\n   - fast = fast.next.next（快指针走两步）\n   - 如果slow == fast，返回true（有环）\n3. 如果循环结束，返回false（无环）\n\n为什么这样有效？\n- 如果无环，fast会先到达null\n- 如果有环，fast会进入环中，最终追上slow\n\n时间复杂度：O(n)，无环时fast到达末尾，有环时fast追上slow\n空间复杂度：O(1)，只使用了两个指针")

add_problem(142, "环形链表 II (Linked List Cycle II)", "中等", "https://leetcode.cn/problems/linked-list-cycle-ii/", "链表",
"给定一个链表的头节点  head ，返回链表开始入环的第一个节点。 如果链表无环，则返回 null。\n\n如果链表中有某个节点，可以通过连续跟踪 next 指针再次到达，则链表中存在环。 为了表示给定链表中的环，评测系统内部使用整数 pos 来表示链表尾连接到链表中的位置（索引从 0 开始）。如果 pos 是 -1，则在该链表中没有环。注意：pos 不作为参数进行传递，仅仅是为了标识链表的实际情况。\n\n不允许修改 链表。\n\n示例：\n输入：head = [3,2,0,-4], pos = 1\n输出：返回索引为 1 的链表节点\n解释：链表中有一个环，其尾部连接到第二个节点。",
"这道题是环形链表的进阶版，不仅要判断是否有环，还要找到环的入口。\n\n核心思想：Floyd判圈算法 + 数学推导\n\n详细步骤：\n1. 使用快慢指针判断是否有环（同141题）\n2. 如果有环，找到相遇点后，让一个指针回到头节点，两个指针都以速度1前进\n3. 再次相遇的位置就是环的入口\n\n数学推导：\n设链表头到环入口的距离为a，环入口到相遇点的距离为b，相遇点回到环入口的距离为c。\n\n当快慢指针相遇时：\n- 慢指针走的距离：a + b\n- 快指针走的距离：a + b + k(b+c) = 2(a+b)\n\n由a + b + k(b+c) = 2(a+b)可得：a = k(b+c) - b = (k-1)(b+c) + c\n\n这意味着：从头节点走a步 = 从相遇点走(k-1)圈 + c步\n\n所以让一个指针从头开始，另一个从相遇点开始，都会在环入口相遇。\n\n时间复杂度：O(n)\n空间复杂度：O(1)")

add_problem(21, "合并两个有序链表 (Merge Two Sorted Lists)", "简单", "https://leetcode.cn/problems/merge-two-sorted-lists/", "链表",
"将两个升序链表合并为一个新的 升序 链表并返回。新链表是通过拼接给定的两个链表的所有节点组成的。\n\n示例：\n输入：l1 = [1,2,4], l2 = [1,3,4]\n输出：[1,1,2,3,4,4]",
"这道题是链表合并的基础题，类似于归并排序中的合并步骤。\n\n核心思想：使用一个虚拟头节点，逐个比较两个链表的节点，将较小的节点接到结果链表后面。\n\n详细步骤：\n1. 创建一个虚拟头节点dummy，以及一个当前指针curr指向dummy\n2. 当l1和l2都不为空时：\n   - 如果l1.val <= l2.val：\n     - curr.next = l1\n     - l1 = l1.next\n   - 否则：\n     - curr.next = l2\n     - l2 = l2.next\n   - curr = curr.next\n3. 如果l1不为空，curr.next = l1（接上剩余部分）\n4. 如果l2不为空，curr.next = l2（接上剩余部分）\n5. 返回dummy.next\n\n为什么使用虚拟头节点？\n虚拟头节点可以简化边界情况的处理，避免对头节点的特殊处理。\n\n时间复杂度：O(n + m)，其中n和m分别是两个链表的长度\n空间复杂度：O(1)，只使用了几个指针（新链表使用原有节点）")

add_problem(2, "两数相加 (Add Two Numbers)", "中等", "https://leetcode.cn/problems/add-two-numbers/", "链表",
"给你两个 非空 的链表，表示两个非负的整数。它们每位数字都是按照 逆序 的方式存储的，并且每个节点只能存储 一位 数字。\n\n请你将两个数相加，并以相同形式返回一个表示和的链表。\n\n你可以假设除了数字 0 之外，这两个数都不会以 0 开头。\n\n示例：\n输入：l1 = [2,4,3], l2 = [5,6,4]\n输出：[7,0,8]\n解释：342 + 465 = 807。",
"这道题模拟了加法运算的过程，只不过数字是以逆序链表的形式存储的。\n\n核心思想：模拟加法运算，注意进位处理。\n\n详细步骤：\n1. 初始化虚拟头节点dummy，当前节点curr = dummy，进位carry = 0\n2. 当l1不为空或l2不为空或carry不为0时循环：\n   - val1 = l1.val if l1 else 0\n   - val2 = l2.val if l2 else 0\n   - sum = val1 + val2 + carry\n   - carry = sum // 10\n   - 创建新节点sum % 10，接到curr后面\n   - curr = curr.next\n   - 如果l1不为空，l1 = l1.next\n   - 如果l2不为空，l2 = l2.next\n3. 返回dummy.next\n\n为什么需要检查carry不为0？\n处理最后的进位，例如99 + 1 = 100，最后会多出一个节点。\n\n时间复杂度：O(max(m, n))，其中m和n分别是两个链表的长度\n空间复杂度：O(max(m, n))，结果链表的长度")

add_problem(19, "删除链表的倒数第 N 个结点 (Remove Nth Node From End of List)", "中等", "https://leetcode.cn/problems/remove-nth-node-from-end-of-list/", "链表",
"给你一个链表，删除链表的倒数第 n 个结点，并且返回链表的头结点。\n\n示例：\n输入：head = [1,2,3,4,5], n = 2\n输出：[1,2,3,5]",
"这道题考察链表的删除操作，关键是找到倒数第n个节点。\n\n核心思想：快慢指针。让快指针先走n步，然后快慢指针一起前进，当快指针到达末尾时，慢指针正好在倒数第n+1个位置。\n\n详细步骤：\n1. 创建虚拟头节点dummy，dummy.next = head\n2. 初始化slow = dummy, fast = dummy\n3. 让fast先走n+1步（这样slow会停在要删除节点的前一个位置）\n4. 当fast不为null时：\n   - slow = slow.next\n   - fast = fast.next\n5. 现在slow.next就是要删除的节点\n6. slow.next = slow.next.next（删除节点）\n7. 返回dummy.next\n\n为什么使用虚拟头节点？\n如果要删除的是头节点，使用虚拟头节点可以统一处理逻辑。\n\n时间复杂度：O(L)，其中L是链表长度\n空间复杂度：O(1)，只使用了几个指针")

add_problem(24, "两两交换链表中的节点 (Swap Nodes in Pairs)", "中等", "https://leetcode.cn/problems/swap-nodes-in-pairs/", "链表",
"给你一个链表，两两交换其中相邻的节点，并返回交换后链表的头节点。你必须在不修改节点内部的值的情况下完成本题（即，只能进行节点交换）。\n\n示例：\n输入：head = [1,2,3,4]\n输出：[2,1,4,3]",
"这道题是链表节点交换的基础题，考察对链表指针的掌握。\n\n核心思想：使用虚拟头节点，遍历链表，每次交换一对节点。\n\n详细步骤：\n1. 创建虚拟头节点dummy，dummy.next = head\n2. 初始化prev = dummy\n3. 当prev.next和prev.next.next都不为空时循环：\n   - first = prev.next\n   - second = prev.next.next\n   - 交换：\n     - prev.next = second\n     - first.next = second.next\n     - second.next = first\n   - prev = first（移动到下一对的前面）\n4. 返回dummy.next\n\n为什么需要prev指针？\nprev指向要交换的一对节点的前一个节点，这样才能正确地修改指针。\n\n时间复杂度：O(n)，需要遍历链表一次\n空间复杂度：O(1)，只使用了几个指针")

add_problem(25, "K 个一组翻转链表 (Reverse Nodes in k-Group)", "困难", "https://leetcode.cn/problems/reverse-nodes-in-k-group/", "链表",
"给你链表的头节点 head ，每 k 个节点一组进行翻转，请你返回修改后的链表。\n\nk 是一个正整数，它的值小于或等于链表的长度。如果节点总数不是 k 的整数倍，那么请将最后剩余的节点保持原有顺序。\n\n你不能只是单纯的改变节点内部的值，而是需要实际进行节点交换。\n\n示例：\n输入：head = [1,2,3,4,5], k = 2\n输出：[2,1,4,3,5]",
"这道题是链表反转的进阶版，要求每k个节点反转一次。\n\n核心思想：分组反转链表。对于每k个节点，先判断是否够k个，然后进行反转。\n\n详细步骤：\n1. 创建虚拟头节点dummy\n2. 初始化prev = dummy\n3. 当还有足够的节点时循环：\n   - 检查是否有k个节点，如果没有，break\n   - 反转这k个节点\n   - 更新prev指向反转后的尾节点\n\n反转k个节点的步骤：\n1. 使用三个指针prev, curr, next进行反转\n2. 反转k次\n3. 记录反转前的头节点（反转后变成尾节点）\n4. 连接反转后的部分\n\n时间复杂度：O(n)，每个节点被访问常数次\n空间复杂度：O(1)，只使用了几个指针")

add_problem(138, "随机链表的复制 (Copy List with Random Pointer)", "中等", "https://leetcode.cn/problems/copy-list-with-random-pointer/", "链表",
"给你一个长度为 n 的链表，每个节点包含一个额外增加的随机指针 random ，该指针可以指向链表中的任何节点或空节点。\n\n构造这个链表的 深拷贝。 深拷贝应该正好由 n 个 全新 节点组成，其中每个新节点的值都设为其对应的原节点的值。新节点的 next 指针和 random 指针也都应指向复制链表中的新节点，并使原链表和复制链表中的这些指针能够表示相同的链表状态。复制链表中的指针都不应指向原链表中的节点 。\n\n例如，如果原链表中有 X 和 Y 两个节点，其中 X.random --> Y 。那么在复制链表中对应的两个节点 x 和 y ，同样有 x.random --> y 。\n\n返回复制链表的头节点。\n\n示例：\n输入：head = [[7,null],[13,0],[11,4],[10,2],[1,0]]\n输出：[[7,null],[13,0],[11,4],[10,2],[1,0]]",
"这道题是复杂链表的深拷贝，难点在于如何处理random指针，因为random可能指向后面的节点。\n\n方法一：哈希表（O(n)空间）\n1. 第一次遍历，创建所有新节点，使用哈希表存储原节点到新节点的映射\n2. 第二次遍历，根据映射关系设置next和random指针\n\n方法二：原地插入（O(1)空间，推荐）\n1. 第一次遍历：在每个原节点后面插入复制节点\n   - 原：A → B → C\n   - 新：A → A' → B → B' → C → C'\n2. 第二次遍历：设置复制节点的random指针\n   - curr.next.random = curr.random.next（如果curr.random存在）\n3. 第三次遍历：分离两个链表\n\n时间复杂度：O(n)，需要遍历链表三次\n空间复杂度：O(1)，只使用了几个指针")

add_problem(148, "排序链表 (Sort List)", "中等", "https://leetcode.cn/problems/sort-list/", "链表",
"给你链表的头结点 head ，请将其按 升序 排列并返回 排序后的链表 。\n\n示例：\n输入：head = [4,2,1,3]\n输出：[1,2,3,4]",
"这道题要求对链表进行排序，最优解法是归并排序，可以达到O(n log n)时间复杂度和O(1)空间复杂度。\n\n核心思想：归并排序（自底向上或自顶向下）\n\n自顶向下归并排序：\n1. 找到链表中点（使用快慢指针）\n2. 递归排序左半部分\n3. 递归排序右半部分\n4. 合并两个有序链表\n\n自底向上归并排序（O(1)空间）：\n1. 计算链表长度length\n2. 从步长step=1开始，每次翻倍\n3. 对每个步长，将链表分成多个组进行合并\n4. 直到step >= length\n\n时间复杂度：O(n log n)\n空间复杂度：\n- 自顶向下：O(log n)递归栈\n- 自底向上：O(1)")

add_problem(23, "合并 K 个升序链表 (Merge k Sorted Lists)", "困难", "https://leetcode.cn/problems/merge-k-sorted-lists/", "链表",
"给你一个链表数组，每个链表都已经按升序排列。\n\n请你将所有链表合并到一个升序链表中，返回合并后的链表。\n\n示例：\n输入：lists = [[1,4,5],[1,3,4],[2,6]]\n输出：[1,1,2,3,4,4,5,6]",
"这道题是合并两个有序链表的扩展，需要合并k个链表。\n\n方法一：顺序合并（O(kn)）\n逐个合并链表，时间复杂度较高。\n\n方法二：分治合并（O(n log k)，推荐）\n使用归并排序的思想，两两合并链表。\n\n方法三：优先队列/最小堆（O(n log k)）\n1. 将所有链表的头节点加入最小堆\n2. 每次取出堆顶元素（最小值），接到结果链表\n3. 如果该节点有下一个节点，将下一个节点加入堆\n4. 重复直到堆为空\n\n时间复杂度：O(n log k)，其中n是所有链表的总节点数\n空间复杂度：O(log k)，递归栈空间")

add_problem(146, "LRU 缓存 (LRU Cache)", "中等", "https://leetcode.cn/problems/lru-cache/", "链表",
"请你设计并实现一个满足  LRU (最近最少使用) 缓存 约束的数据结构。\n\n实现 LRUCache 类：\nLRUCache(int capacity) 以 正整数 作为容量 capacity 初始化 LRU 缓存\nint get(int key) 如果关键字 key 存在于缓存中，则返回关键字的值，否则返回 -1 。\nvoid put(int key, int value) 如果关键字 key 已经存在，则变更其数据值 value ；如果不存在，则向缓存中插入该组 key-value 。如果插入操作导致关键字数量超过 capacity ，则应该 逐出 最久未使用的关键字。\n函数 get 和 put 必须以 O(1) 的平均时间复杂度运行。\n\n示例：\n输入\n[\"LRUCache\", \"put\", \"put\", \"get\", \"put\", \"get\", \"put\", \"get\", \"get\", \"get\"]\n[[2], [1, 1], [2, 2], [1], [3, 3], [2], [4, 4], [1], [3], [4]]\n输出\n[null, null, null, 1, null, -1, null, -1, 3, 4]",
"这道题要求设计一个LRU缓存，关键是get和put操作都要达到O(1)时间复杂度。\n\n核心思想：哈希表 + 双向链表\n\n为什么需要双向链表？\n- 需要在O(1)时间内移动节点到头部（最近使用）\n- 需要在O(1)时间内删除尾部节点（最久未使用）\n- 需要在O(1)时间内删除任意节点（当节点被更新时）\n\n为什么需要哈希表？\n- 需要在O(1)时间内根据key找到对应的节点\n\n详细设计：\n1. 使用双向链表维护节点的使用顺序，头部是最新使用的，尾部是最久未使用的\n2. 使用哈希表存储key到节点的映射\n\nget操作：\n1. 如果key不存在，返回-1\n2. 如果key存在，将对应节点移动到链表头部，返回value\n\nput操作：\n1. 如果key存在，更新value，将节点移动到头部\n2. 如果key不存在：\n   - 创建新节点，加入哈希表，插入链表头部\n   - 如果超出容量，删除链表尾部节点及其在哈希表中的映射\n\n时间复杂度：get和put都是O(1)\n空间复杂度：O(capacity)")

# ========== 7. 二叉树 (15题) ==========
add_problem(104, "二叉树的最大深度 (Maximum Depth of Binary Tree)", "简单", "https://leetcode.cn/problems/maximum-depth-of-binary-tree/", "二叉树",
"给定一个二叉树 root ，返回其最大深度。\n二叉树的 最大深度 是指从根节点到最远叶子节点的最长路径上的节点数。\n\n示例：\n输入：root = [3,9,20,null,null,15,7]\n输出：3",
"这道题是二叉树的基础题，求二叉树的最大深度。\n\n方法一：递归（DFS）\n递归地求左右子树的最大深度，取较大值加1。\ndepth(root) = max(depth(root.left), depth(root.right)) + 1\n\n时间复杂度：O(n)，需要遍历所有节点\n空间复杂度：O(h)，递归栈空间，h是树的高度\n\n方法二：迭代（BFS）\n使用层序遍历，每遍历一层深度加1。\n\n时间复杂度：O(n)\n空间复杂度：O(w)，w是树的最大宽度")

add_problem(100, "相同的树 (Same Tree)", "简单", "https://leetcode.cn/problems/same-tree/", "二叉树",
"给你两棵二叉树的根节点 p 和 q ，编写一个函数来检验这两棵树是否相同。\n如果两个树在结构上相同，并且节点具有相同的值，则认为它们是相同的。\n\n示例：\n输入：p = [1,2,3], q = [1,2,3]\n输出：true",
"这道题考察二叉树的递归遍历。\n\n核心思想：递归比较两棵树的每个节点。\n\n递归条件：\n1. 如果两个节点都为空，返回true\n2. 如果一个为空一个不为空，返回false\n3. 如果两个节点的值不同，返回false\n4. 递归比较左子树和右子树\n\nisSame(p, q) = p.val == q.val AND isSame(p.left, q.left) AND isSame(p.right, q.right)\n\n时间复杂度：O(min(m, n))，其中m和n是两棵树的节点数\n空间复杂度：O(min(h1, h2))，递归栈空间")

add_problem(226, "翻转二叉树 (Invert Binary Tree)", "简单", "https://leetcode.cn/problems/invert-binary-tree/", "二叉树",
"给你一棵二叉树的根节点 root ，翻转这棵二叉树，并返回其根节点。\n\n示例：\n输入：root = [4,2,7,1,3,6,9]\n输出：[4,7,2,9,6,3,1]",
"这道题是二叉树的基础操作，翻转二叉树。\n\n核心思想：递归地交换每个节点的左右子树。\n\n方法一：递归\n1. 如果root为空，返回null\n2. 交换root的左右子树\n3. 递归翻转左子树\n4. 递归翻转右子树\n5. 返回root\n\n方法二：迭代（BFS）\n使用队列进行层序遍历，每次出队时交换左右子节点。\n\n时间复杂度：O(n)，需要遍历所有节点\n空间复杂度：\n- 递归：O(h)\n- 迭代：O(w)")

add_problem(101, "对称二叉树 (Symmetric Tree)", "简单", "https://leetcode.cn/problems/symmetric-tree/", "二叉树",
"给你一个二叉树的根节点 root ， 检查它是否轴对称。\n\n示例：\n输入：root = [1,2,2,3,4,4,3]\n输出：true",
"这道题考察判断二叉树是否对称。\n\n核心思想：递归或迭代比较左右子树是否镜像对称。\n\n递归方法：\n定义一个函数check(left, right)比较两个子树是否镜像：\n1. 如果left和right都为空，返回true\n2. 如果一个为空一个不为空，返回false\n3. 如果left.val != right.val，返回false\n4. 递归比较left.left和right.right，以及left.right和right.left\n\n迭代方法：\n使用队列，每次取出两个节点进行比较，然后将它们的子节点按对称顺序入队。\n\n时间复杂度：O(n)\n空间复杂度：O(n)")

add_problem(543, "二叉树的直径 (Diameter of Binary Tree)", "简单", "https://leetcode.cn/problems/diameter-of-binary-tree/", "二叉树",
"给你一棵二叉树的根节点 root ，返回该树的 直径 。\n二叉树的 直径 是指树中任意两个节点之间最长路径的 长度 。这条路径可能经过也可能不经过根节点 root 。\n两节点之间路径的 长度 由它们之间边数表示。\n\n示例：\n输入：root = [1,2,3,4,5]\n输出：3\n解释：3 ，取路径 [4,2,1,3] 或 [5,2,1,3] 的长度。",
"这道题是二叉树直径的计算，直径可能不经过根节点。\n\n核心思想：后序遍历，计算每个节点的左右子树深度，更新直径。\n\n关键观察：\n对于每个节点，经过它的最长路径 = 左子树深度 + 右子树深度\n\n详细步骤：\n1. 定义全局变量diameter记录最大直径\n2. 定义递归函数depth(node)计算以node为根的子树深度\n3. 在depth函数中：\n   - 递归计算左子树深度left\n   - 递归计算右子树深度right\n   - 更新diameter = max(diameter, left + right)\n   - 返回max(left, right) + 1\n4. 调用depth(root)\n5. 返回diameter\n\n时间复杂度：O(n)\n空间复杂度：O(h)")

add_problem(102, "二叉树的层序遍历 (Binary Tree Level Order Traversal)", "中等", "https://leetcode.cn/problems/binary-tree-level-order-traversal/", "二叉树",
"给你二叉树的根节点 root ，返回其节点值的 层序遍历 。 （即逐层地，从左到右访问所有节点）。\n\n示例：\n输入：root = [3,9,20,null,null,15,7]\n输出：[[3],[9,20],[15,7]]",
"这道题是二叉树的层序遍历（BFS）的经典应用。\n\n核心思想：使用队列进行广度优先搜索，按层遍历节点。\n\n详细步骤：\n1. 如果root为空，返回空列表\n2. 初始化队列，加入root\n3. 当队列不为空时循环：\n   - 记录当前层的节点数levelSize\n   - 创建一个列表存储当前层的节点值\n   - 循环levelSize次：\n     - 出队一个节点\n     - 将节点值加入当前层列表\n     - 将左右子节点入队（如果不为空）\n   - 将当前层列表加入结果\n4. 返回结果\n\n时间复杂度：O(n)\n空间复杂度：O(w)，w是树的最大宽度")

add_problem(108, "将有序数组转换为二叉搜索树 (Convert Sorted Array to Binary Search Tree)", "简单", "https://leetcode.cn/problems/convert-sorted-array-to-binary-search-tree/", "二叉树",
"给你一个整数数组 nums ，其中元素已经按 升序 排列，请你将其转换为一棵 高度平衡 二叉搜索树。\n高度平衡 二叉树是一棵满足「每个节点的左右两个子树的高度差的绝对值不超过 1 」的二叉树。\n\n示例：\n输入：nums = [-10,-3,0,5,9]\n输出：[0,-3,9,-10,null,5]\n解释：[0,-10,5,null,-3,null,9] 也将被视为正确答案。",
"这道题要求将有序数组转换为高度平衡的二叉搜索树。\n\n核心思想：递归地选择中间元素作为根节点。\n\n为什么选择中间元素？\n- 二叉搜索树的性质：左子树 < 根 < 右子树\n- 选择中间元素可以保证左右子树的节点数大致相等，从而保证平衡\n\n详细步骤：\n1. 定义递归函数build(left, right)构建nums[left:right+1]的子树\n2. 如果left > right，返回null\n3. mid = (left + right) // 2\n4. 以nums[mid]为根节点\n5. 递归构建左子树：root.left = build(left, mid-1)\n6. 递归构建右子树：root.right = build(mid+1, right)\n7. 返回root\n\n时间复杂度：O(n)，每个元素被访问一次\n空间复杂度：O(log n)，递归栈空间")

add_problem(98, "验证二叉搜索树 (Validate Binary Search Tree)", "中等", "https://leetcode.cn/problems/validate-binary-search-tree/", "二叉树",
"给你一个二叉树的根节点 root ，判断其是否是一个有效的二叉搜索树。\n有效 二叉搜索树定义如下：\n节点的左子树只包含 小于 当前节点的数。\n节点的右子树只包含 大于 当前节点的数。\n所有左子树和右子树自身必须也是二叉搜索树。\n\n示例：\n输入：root = [2,1,3]\n输出：true",
"这道题要求验证二叉搜索树，容易犯的错误是只比较当前节点和左右子节点。\n\n核心思想：递归时维护当前节点允许的取值范围。\n\n为什么需要范围？\n只比较当前节点和左右子节点是不够的。例如：[5,1,4,null,null,3,6]，4大于5，违反了BST性质，但只看4和它的子节点是合法的。\n\n详细步骤：\n1. 定义递归函数validate(node, min, max)\n2. 如果node为空，返回true\n3. 如果node.val <= min或node.val >= max，返回false\n4. 递归验证左子树：validate(node.left, min, node.val)\n5. 递归验证右子树：validate(node.right, node.val, max)\n6. 返回两个递归结果的与\n\n初始调用：validate(root, -inf, +inf)\n\n时间复杂度：O(n)\n空间复杂度：O(h)")

add_problem(230, "二叉搜索树中第K小的元素 (Kth Smallest Element in a BST)", "中等", "https://leetcode.cn/problems/kth-smallest-element-in-a-bst/", "二叉树",
"给定一个二叉搜索树的根节点 root ，和一个整数 k ，请你设计一个算法查找其中第 k 小的元素（从 1 开始计数）。\n\n示例：\n输入：root = [3,1,4,null,2], k = 1\n输出：1",
"这道题利用二叉搜索树的中序遍历性质。\n\n核心思想：BST的中序遍历是有序的（升序），第k个元素就是第k小的元素。\n\n方法一：中序遍历\n1. 进行中序遍历\n2. 记录遍历计数\n3. 当计数等于k时，找到答案\n\n方法二：迭代中序遍历（推荐）\n使用栈进行迭代中序遍历，可以提前终止，不需要遍历完整棵树。\n\n优化：如果频繁查询\n可以在TreeNode中增加count字段，记录以该节点为根的子树的节点数。\n\n时间复杂度：O(h + k)，h是树的高度\n空间复杂度：O(h)")

add_problem(199, "二叉树的右视图 (Binary Tree Right Side View)", "中等", "https://leetcode.cn/problems/binary-tree-right-side-view/", "二叉树",
"给定一个二叉树的 根节点 root，想象自己站在它的右侧，按照从顶部到底部的顺序，返回从右侧所能看到的节点值。\n\n示例：\n输入：root = [1,2,3,null,5,null,4]\n输出：[1,3,4]",
"这道题要求从右侧看二叉树，能看到的是每层最右边的节点。\n\n方法一：BFS层序遍历\n进行层序遍历，每层取最后一个节点。\n\n方法二：DFS先右后左\n进行深度优先搜索，先访问右子节点，记录第一个访问的每层节点。\n\nDFS详细步骤：\n1. 定义递归函数dfs(node, depth)\n2. 如果node为空，返回\n3. 如果depth等于结果列表的长度，说明这是该层的第一个节点，加入结果\n4. 先递归右子节点：dfs(node.right, depth+1)\n5. 再递归左子节点：dfs(node.left, depth+1)\n\n时间复杂度：O(n)\n空间复杂度：O(h)")

add_problem(114, "二叉树展开为链表 (Flatten Binary Tree to Linked List)", "中等", "https://leetcode.cn/problems/flatten-binary-tree-to-linked-list/", "二叉树",
"给你二叉树的根结点 root ，请你将它展开为一个单链表：\n展开后的单链表应该同样使用 TreeNode ，其中 right 子指针指向链表中下一个结点，而左子指针始终为 null 。\n展开后的单链表应该与二叉树 先序遍历 顺序相同。\n\n示例：\n输入：root = [1,2,5,3,4,null,6]\n输出：[1,null,2,null,3,null,4,null,5,null,6]",
"这道题要求将二叉树展开为链表，顺序是先序遍历。\n\n方法一：递归（需要额外空间存储先序遍历结果）\n1. 先序遍历，存储所有节点\n2. 按顺序连接节点\n\n方法二：原地展开（O(1)空间）\n利用先序遍历的性质，找到左子树的最右节点，将其右指针指向右子树。\n\n原地展开详细步骤：\n1. 初始化curr = root\n2. 当curr不为空时循环：\n   - 如果curr有左子节点：\n     - 找到左子树的最右节点predecessor\n     - 将predecessor的右指针指向curr的右子树\n     - 将curr的右指针指向curr的左子树\n     - 将curr的左指针设为null\n   - curr = curr.right\n\n时间复杂度：O(n)\n空间复杂度：O(1)")

add_problem(105, "从前序与中序遍历序列构造二叉树 (Construct Binary Tree from Preorder and Inorder Traversal)", "中等", "https://leetcode.cn/problems/construct-binary-tree-from-preorder-and-inorder-traversal/", "二叉树",
"给定两个整数数组 preorder 和 inorder ，其中 preorder 是二叉树的先序遍历， inorder 是同一棵树的中序遍历，请构造二叉树并返回其根节点。\n\n示例：\n输入: preorder = [3,9,20,15,7], inorder = [9,3,15,20,7]\n输出: [3,9,20,null,null,15,7]",
"这道题是根据遍历序列构造二叉树的经典题目。\n\n核心思想：利用前序遍历和中序遍历的性质。\n\n性质：\n- 前序遍历：根节点 -> 左子树 -> 右子树\n- 中序遍历：左子树 -> 根节点 -> 右子树\n\n推导：\n1. 前序的第一个元素是根节点\n2. 在中序中找到根节点的位置，左边是左子树，右边是右子树\n3. 根据左子树的节点数，在前序中划分左子树和右子树\n4. 递归构造左右子树\n\n详细步骤：\n1. 使用前序的首元素创建根节点\n2. 在中序中找到该元素的位置index\n3. 左子树的节点数为index - inorder_left\n4. 递归构造左子树和右子树\n\n优化：使用哈希表存储中序的值到索引的映射，避免每次在中序中线性查找。\n\n时间复杂度：O(n)\n空间复杂度：O(n)")

add_problem(437, "路径总和 III (Path Sum III)", "中等", "https://leetcode.cn/problems/path-sum-iii/", "二叉树",
"给定一个二叉树的根节点 root ，和一个整数 targetSum ，求该二叉树里节点值之和等于 targetSum 的 路径 的数目。\n路径 不需要从根节点开始，也不需要在叶子节点结束，但是路径方向必须是向下的（只能从父节点到子节点）。\n\n示例：\n输入：root = [10,5,-3,3,2,null,11,3,-2,null,1], targetSum = 8\n输出：3\n解释：和等于 8 的路径有 3 条。",
"这道题是路径和的进阶版，路径可以从任意节点开始，到任意节点结束。\n\n方法一：双重递归（O(n²)）\n对每个节点，计算以它为起点的路径数。\n\n方法二：前缀和 + 哈希表（O(n)，推荐）\n类似数组中求和为k的子数组个数的方法。\n\n前缀和方法详细步骤：\n1. 定义哈希表prefixSumCount，记录每个前缀和出现的次数\n2. 初始化prefixSumCount[0] = 1\n3. 定义递归函数dfs(node, currSum)：\n   - 如果node为空，返回0\n   - currSum += node.val\n   - count = prefixSumCount.get(currSum - targetSum, 0)\n   - prefixSumCount[currSum] = prefixSumCount.get(currSum, 0) + 1\n   - count += dfs(node.left, currSum) + dfs(node.right, currSum)\n   - prefixSumCount[currSum] -= 1（回溯）\n   - 返回count\n4. 调用dfs(root, 0)\n\n为什么需要回溯？\n因为路径是单向向下的，从根到当前节点的路径是唯一的。离开当前节点时，需要恢复状态，避免影响其他分支。\n\n时间复杂度：O(n)\n空间复杂度：O(h)")

add_problem(236, "二叉树的最近公共祖先 (Lowest Common Ancestor of a Binary Tree)", "中等", "https://leetcode.cn/problems/lowest-common-ancestor-of-a-binary-tree/", "二叉树",
"给定一个二叉树, 找到该树中两个指定节点的最近公共祖先。\n百度百科中最近公共祖先的定义为：\"对于有根树 T 的两个节点 p、q，最近公共祖先表示为一个节点 x，满足 x 是 p、q 的祖先且 x 的深度尽可能大（一个节点也可以是它自己的祖先）。\"\n\n示例：\n输入：root = [3,5,1,6,2,0,8,null,null,7,4], p = 5, q = 1\n输出：3\n解释：节点 5 和节点 1 的最近公共祖先是节点 3 。",
"这道题是二叉树中经典的LCA问题。\n\n核心思想：递归查找，如果当前节点是p或q，或者左右子树分别包含p和q，则当前节点是LCA。\n\n详细步骤：\n1. 如果root为空，返回null\n2. 如果root等于p或q，返回root\n3. 递归在左子树中查找\n4. 递归在右子树中查找\n5. 如果left和right都不为空，说明p和q分别在左右子树，返回root\n6. 如果left为空，返回right\n7. 如果right为空，返回left\n\n为什么这样有效？\n- 如果p和q分别在当前节点的左右子树，当前节点就是LCA\n- 如果都在左子树，LCA在左子树中\n- 如果都在右子树，LCA在右子树中\n- 如果当前节点就是p或q，当前节点就是LCA\n\n时间复杂度：O(n)\n空间复杂度：O(h)")

add_problem(124, "二叉树中的最大路径和 (Binary Tree Maximum Path Sum)", "困难", "https://leetcode.cn/problems/binary-tree-maximum-path-sum/", "二叉树",
"二叉树中的 路径 被定义为一条节点序列，序列中每对相邻节点之间都存在一条边。同一个节点在一条路径序列中 至多 出现一次 。该路径 至少 包含一个 节点，且不一定经过根节点。\n路径和 是路径中各节点值的总和。\n给你一个二叉树的根节点 root ，返回其 最大路径和 。\n\n示例：\n输入：root = [1,2,3]\n输出：6\n解释：最优路径是 2 -> 1 -> 3 ，路径和为 2 + 1 + 3 = 6",
"这道题是二叉树路径问题的难点，路径不需要经过根节点。\n\n核心思想：后序遍历，计算每个节点作为拐点时的最大路径和，同时返回单条路径的最大和。\n\n关键区别：\n- 以当前节点为拐点的路径：可以包含左右子树\n- 返回给父节点的路径：只能包含左右子树中的一个（或都不包含）\n\n详细步骤：\n1. 定义全局变量maxSum记录最大路径和\n2. 定义递归函数maxGain(node)返回从node出发的单条路径的最大和\n3. 在maxGain中：\n   - 如果node为空，返回0\n   - 递归计算左子树的最大增益leftGain = max(maxGain(node.left), 0)\n   - 递归计算右子树的最大增益rightGain = max(maxGain(node.right), 0)\n   - 以当前节点为拐点的路径和：priceNewPath = node.val + leftGain + rightGain\n   - 更新maxSum = max(maxSum, priceNewPath)\n   - 返回node.val + max(leftGain, rightGain)（只能选择一边）\n4. 调用maxGain(root)\n5. 返回maxSum\n\n为什么使用max(..., 0)？\n如果子树的路径和为负，不选它比选它更好（贡献为0）。\n\n时间复杂度：O(n)\n空间复杂度：O(h)")

add_problem(297, "二叉树的序列化与反序列化 (Serialize and Deserialize Binary Tree)", "困难", "https://leetcode.cn/problems/serialize-and-deserialize-binary-tree/", "二叉树",
"序列化是将一个数据结构或者对象转换为连续的比特位的操作，进而可以将转换后的数据存储在一个文件或者内存中，同时也可以通过网络传输到另一个计算机环境，采取相反方式重构得到原数据。\n请设计一个算法来实现二叉树的序列化与反序列化。这里不限定你的序列 / 反序列化算法执行逻辑，你只需要保证一个二叉树可以被序列化为一个字符串并且将这个字符串反序列化为原始的树结构。\n\n示例：\n输入：root = [1,2,3,null,null,4,5]\n输出：[1,2,3,null,null,4,5]",
"这道题要求设计二叉树的序列化和反序列化方法，常见的有前序遍历、后序遍历、层序遍历等方法。\n\n方法一：前序遍历\n序列化：前序遍历二叉树，空节点用特殊字符表示（如\"#\"），节点间用分隔符（如\",\"）\n反序列化：按分隔符分割字符串，递归构建二叉树\n\n方法二：层序遍历（推荐，与LeetCode格式一致）\n序列化：使用队列进行层序遍历，空节点也加入序列\n反序列化：使用队列，按顺序重建节点\n\n层序反序列化详细步骤：\n1. 按分隔符分割字符串得到数组\n2. 如果第一个元素是空，返回null\n3. 创建根节点\n4. 使用队列，将根节点入队\n5. 按顺序读取数组，每次取两个值作为当前节点的左右子节点\n6. 如果子节点不为空，入队\n7. 重复直到数组遍历完毕\n\n时间复杂度：O(n)\n空间复杂度：O(n)")

print(f"已添加 {len(PROBLEMS)} 道题目")

# ========== 8. 图论 (4题) ==========
add_problem(200, "岛屿数量 (Number of Islands)", "中等", "https://leetcode.cn/problems/number-of-islands/", "图论",
"给你一个由 '1'（陆地）和 '0'（水）组成的的二维网格，请你计算网格中岛屿的数量。\n岛屿总是被水包围，并且每座岛屿只能由水平方向和/或竖直方向上相邻的陆地连接形成。\n此外，你可以假设该网格的四条边均被水包围。\n\n示例：\n输入：grid = [\n  [\"1\",\"1\",\"1\",\"1\",\"0\"],\n  [\"1\",\"1\",\"0\",\"1\",\"0\"],\n  [\"1\",\"1\",\"0\",\"0\",\"0\"],\n  [\"0\",\"0\",\"0\",\"0\",\"0\"]\n]\n输出：1",
"这道题是图的连通分量计数问题。\n\n核心思想：遍历网格，当遇到陆地时，使用DFS或BFS标记整个岛屿。\n\n方法一：DFS\n1. 遍历网格的每个位置\n2. 如果当前位置是'1'（未访问的陆地）：\n   - 岛屿计数加1\n   - 使用DFS将当前岛屿的所有陆地标记为'0'（或访问标记）\n3. 返回岛屿计数\n\nDFS函数：\n1. 如果越界或当前位置是'0'，返回\n2. 将当前位置标记为'0'\n3. 递归访问上下左右四个方向\n\n方法二：BFS\n类似DFS，只是使用队列代替递归。\n\n方法三：并查集\n将所有陆地加入并查集，最后统计连通分量的数量。\n\n时间复杂度：O(m × n)\n空间复杂度：\n- DFS：O(m × n)递归栈，最坏情况全是陆地\n- BFS：O(min(m, n))队列\n- 并查集：O(m × n)")

add_problem(207, "课程表 (Course Schedule)", "中等", "https://leetcode.cn/problems/course-schedule/", "图论",
"你这个学期必须选修 numCourses 门课程，记为 0 到 numCourses - 1 。\n在选修某些课程之前需要一些先修课程。 先修课程按数组 prerequisites 给出，其中 prerequisites[i] = [ai, bi] ，表示如果要学习课程 ai 则 必须 先学习课程  bi 。\n例如，先修课程对 [0, 1] 表示：想要学习课程 0 ，你需要先完成课程 1 。\n请你判断是否可能完成所有课程的学习？如果可以，返回 true ；否则，返回 false 。\n\n示例：\n输入：numCourses = 2, prerequisites = [[1,0]]\n输出：true\n解释：总共有 2 门课程。学习课程 1 之前，你需要完成课程 0 。这是可能的。",
"这道题是经典的拓扑排序问题，判断有向图是否有环。\n\n核心思想：拓扑排序，检查是否存在环。\n\n方法一：Kahn算法（BFS）\n1. 计算每个节点的入度\n2. 将所有入度为0的节点加入队列\n3. 当队列不为空时：\n   - 出队一个节点\n   - 将其所有邻居的入度减1\n   - 如果邻居入度变为0，入队\n4. 如果处理的节点数等于总节点数，说明无环\n\n方法二：DFS\n使用三种状态标记节点：\n- 0：未访问\n- 1：正在访问（在当前DFS路径中）\n- 2：已访问（不在当前DFS路径中）\n如果在DFS过程中遇到状态为1的节点，说明有环。\n\n时间复杂度：O(V + E)，V是节点数，E是边数\n空间复杂度：O(V + E)")

add_problem(208, "实现 Trie (前缀树) (Implement Trie Prefix Tree)", "中等", "https://leetcode.cn/problems/implement-trie-prefix-tree/", "图论",
"Trie（发音类似 \"try\"）或者说 前缀树 是一种树形数据结构，用于高效地存储和检索字符串数据集中的键。这一数据结构有相当多的应用情景，如自动补完和拼写检查。\n请你实现 Trie 类：\nTrie() 初始化前缀树对象。\nvoid insert(String word) 向前缀树中插入字符串 word 。\nboolean search(String word) 如果字符串 word 在前缀树中，返回 true（即，在检索之前已经插入）；否则，返回 false 。\nboolean startsWith(String prefix) 如果之前已经插入的字符串 word 的前缀之一为 prefix ，返回 true ；否则，返回 false 。\n\n示例：\n输入\n[\"Trie\", \"insert\", \"search\", \"search\", \"startsWith\", \"insert\", \"search\"]\n[[], [\"apple\"], [\"apple\"], [\"app\"], [\"app\"], [\"app\"], [\"app\"]]\n输出\n[null, null, true, false, true, null, true]",
"Trie（前缀树）是一种多叉树，每个节点代表一个字符，从根到某个节点的路径表示一个前缀。\n\nTrie节点结构：\n- children：子节点数组或哈希表（char -> TrieNode）\n- isEnd：标记是否是一个完整单词的结尾\n\n操作详解：\n\ninsert(word)：\n1. 从根节点开始\n2. 对于word中的每个字符c：\n   - 如果当前节点没有c的子节点，创建新节点\n   - 移动到c的子节点\n3. 将最后一个节点的isEnd设为true\n\nsearch(word)：\n1. 从根节点开始遍历word中的字符\n2. 如果某个字符不存在，返回false\n3. 遍历完所有字符后，返回最后一个节点的isEnd\n\nstartsWith(prefix)：\n类似search，但不要求isEnd为true，只需要能遍历完所有字符。\n\n时间复杂度：\n- insert：O(m)，m是word的长度\n- search：O(m)\n- startsWith：O(m)\n\n空间复杂度：O(m × n)，n是插入的单词数")

add_problem(994, "腐烂的橘子 (Rotting Oranges)", "中等", "https://leetcode.cn/problems/rotting-oranges/", "图论",
"在给定的 m x n 网格 grid 中，每个单元格可以有以下三个值之一：\n值 0 代表空单元格；\n值 1 代表新鲜橘子；\n值 2 代表腐烂的橘子。\n每分钟，腐烂的橘子 周围 4 个方向上相邻 的新鲜橘子都会腐烂。\n返回 直到单元格中没有新鲜橘子为止所必须经过的最小分钟数。如果不可能，返回 -1 。\n\n示例：\n输入：grid = [[2,1,1],[1,1,0],[0,1,1]]\n输出：4",
"这道题是多源BFS的经典应用。\n\n核心思想：BFS层级遍历，每一分钟对应BFS的一层。\n\n详细步骤：\n1. 遍历网格，找到所有腐烂的橘子，加入队列\n2. 统计新鲜橘子的数量freshCount\n3. 初始化minutes = -1（如果一开始就有腐烂橘子，最后要+1）\n4. BFS：\n   - minutes++\n   - 当前层的所有腐烂橘子向四周传播\n   - 如果有新鲜橘子被腐烂，freshCount--，加入队列\n5. 如果freshCount > 0，返回-1\n6. 如果minutes == -1（没有腐烂橘子），返回0\n7. 返回minutes\n\n时间复杂度：O(m × n)\n空间复杂度：O(m × n)")

print(f"已添加 {len(PROBLEMS)} 道题目")

# ========== 9. 回溯 (6题) ==========
add_problem(78, "子集 (Subsets)", "中等", "https://leetcode.cn/problems/subsets/", "回溯",
"给你一个整数数组 nums ，数组中的元素 互不相同 。返回该数组所有可能的子集（幂集）。\n解集 不能 包含重复的子集。你可以按 任意顺序 返回解集。\n\n示例：\n输入：nums = [1,2,3]\n输出：[[],[1],[2],[1,2],[3],[1,3],[2,3],[1,2,3]]",
"这道题是回溯法的经典应用，求数组的所有子集。\n\n核心思想：对于每个元素，选择选或不选，使用回溯遍历所有可能。\n\n方法一：回溯法（选/不选）\n1. 定义递归函数backtrack(start, currentSubset)\n2. 将currentSubset加入结果\n3. 从start开始遍历nums：\n   - 将nums[i]加入currentSubset\n   - 递归调用backtrack(i+1, currentSubset)\n   - 回溯，将nums[i]从currentSubset移除\n\n方法二：迭代法\n初始结果：[[]]\n对于每个num：\n   - 将num加到已有每个子集中，形成新的子集\n   - 加入结果\n\n时间复杂度：O(n × 2^n)，共有2^n个子集，每个子集平均长度n/2\n空间复杂度：O(n)，递归栈空间")

add_problem(17, "电话号码的字母组合 (Letter Combinations of a Phone Number)", "中等", "https://leetcode.cn/problems/letter-combinations-of-a-phone-number/", "回溯",
"给定一个仅包含数字 2-9 的字符串，返回所有它能表示的字母组合。答案可以按 任意顺序 返回。\n给出数字到字母的映射如下（与电话按键相同）。注意 1 不对应任何字母。\n\n示例：\n输入：digits = \"23\"\n输出：[\"ad\",\"ae\",\"af\",\"bd\",\"be\",\"bf\",\"cd\",\"ce\",\"cf\"]",
"这道题是回溯法的经典应用，求所有可能的字母组合。\n\n核心思想：回溯遍历所有可能的组合。\n\n详细步骤：\n1. 建立数字到字母的映射表\n2. 定义递归函数backtrack(index, currentString)\n   - index：当前处理的digits的索引\n   - currentString：当前构建的字符串\n3. 如果index等于digits的长度：\n   - 将currentString加入结果\n   - 返回\n4. 获取当前数字对应的字母列表\n5. 遍历每个字母：\n   - 将字母加入currentString\n   - 递归调用backtrack(index+1, ...)\n   - 回溯，移除最后一个字母\n\n时间复杂度：O(3^m × 4^n)，m是对应3个字母的数字个数，n是对应4个字母的数字个数\n空间复杂度：O(m+n)，递归栈深度")

add_problem(39, "组合总和 (Combination Sum)", "中等", "https://leetcode.cn/problems/combination-sum/", "回溯",
"给你一个 无重复元素 的整数数组 candidates 和一个目标整数 target ，找出 candidates 中可以使数字和为目标数 target 的 所有 不同组合 ，并以列表形式返回。你可以按 任意顺序 返回这些组合。\ncandidates 中的 同一个 数字可以 无限制重复被选取 。如果至少一个数字的被选数量不同，则两种组合是不同的。\n\n示例：\n输入：candidates = [2,3,6,7], target = 7\n输出：[[2,2,3],[7]]",
"这道题是回溯法的经典应用，求和为目标值的所有组合。\n\n核心思想：回溯搜索，注意可以重复使用元素。\n\n详细步骤：\n1. 定义递归函数backtrack(start, target, currentCombination)\n   - start：从candidates的哪个位置开始选择（避免重复组合）\n   - target：剩余的目标值\n   - currentCombination：当前组合\n2. 如果target == 0：\n   - 将currentCombination加入结果\n   - 返回\n3. 如果target < 0：\n   - 返回（剪枝）\n4. 从start开始遍历candidates：\n   - 将candidates[i]加入currentCombination\n   - 递归调用backtrack(i, target-candidates[i], ...)（注意i不是i+1，因为可以重复使用）\n   - 回溯\n\n为什么start参数可以避免重复组合？\n如果每次都从0开始，会产生[2,3]和[3,2]这样的重复组合。通过start参数，我们强制组合中的数字按非递减顺序出现。\n\n时间复杂度：O(S)，S是所有可行解的长度之和\n空间复杂度：O(target)，递归栈深度")

add_problem(22, "括号生成 (Generate Parentheses)", "中等", "https://leetcode.cn/problems/generate-parentheses/", "回溯",
"数字 n 代表生成括号的对数，请你设计一个函数，用于能够生成所有可能的并且 有效的 括号组合。\n\n示例：\n输入：n = 3\n输出：[\"((()))\",\"(()())\",\"(())()\",\"()(())\",\"()()()\"]",
"这道题是回溯法的经典应用，生成有效的括号组合。\n\n核心思想：回溯生成，维护左右括号的数量，确保有效性。\n\n有效括号的条件：\n1. 左括号的数量等于右括号的数量等于n\n2. 在任意前缀中，左括号的数量 >= 右括号的数量\n\n详细步骤：\n1. 定义递归函数backtrack(currentString, leftCount, rightCount)\n2. 如果currentString的长度等于2n：\n   - 加入结果\n   - 返回\n3. 如果leftCount < n：\n   - 添加'('，递归\n   - 回溯\n4. 如果rightCount < leftCount：\n   - 添加')'，递归\n   - 回溯\n\n为什么rightCount < leftCount时才添加')'？\n这保证了任意前缀中左括号数量 >= 右括号数量，即括号的有效性。\n\n时间复杂度：第n个卡特兰数，约为O(4^n / √n)\n空间复杂度：O(n)，递归栈深度")

add_problem(79, "单词搜索 (Word Search)", "中等", "https://leetcode.cn/problems/word-search/", "回溯",
"给定一个 m x n 二维字符网格 board 和一个字符串单词 word 。如果 word 存在于网格中，返回 true ；否则，返回 false 。\n单词必须按照字母顺序，通过相邻的单元格内的字母构成，其中\"相邻\"单元格是那些水平相邻或垂直相邻的单元格。同一个单元格内的字母不允许被重复使用。\n\n示例：\n输入：board = [[\"A\",\"B\",\"C\",\"E\"],[\"S\",\"F\",\"C\",\"S\"],[\"A\",\"D\",\"E\",\"E\"]], word = \"ABCCED\"\n输出：true",
"这道题是回溯法在二维网格上的应用。\n\n核心思想：从每个位置开始DFS，尝试匹配单词的每个字符。\n\n详细步骤：\n1. 遍历网格的每个位置作为起点\n2. 定义递归函数dfs(i, j, k)，表示从(i,j)开始匹配word[k:]\n3. 如果k等于word的长度，返回true\n4. 如果(i,j)越界或board[i][j] != word[k]，返回false\n5. 临时标记board[i][j]为已访问（如设为特殊字符或哈希集合）\n6. 递归尝试四个方向，如果任一方向成功，返回true\n7. 恢复board[i][j]（回溯）\n8. 返回false\n\n为什么需要恢复board[i][j]？\n这是回溯的关键，确保不同路径可以共享同一个网格位置。\n\n时间复杂度：O(m × n × 3^L)，L是word的长度，每个位置有3个方向\n空间复杂度：O(L)，递归栈深度")

add_problem(131, "分割回文串 (Palindrome Partitioning)", "中等", "https://leetcode.cn/problems/palindrome-partitioning/", "回溯",
"给你一个字符串 s，请你将 s 分割成一些子串，使每个子串都是 回文串 。返回 s 所有可能的分割方案。\n\n示例：\n输入：s = \"aab\"\n输出：[[\"a\",\"a\",\"b\"],[\"aa\",\"b\"]]",
"这道题是回溯法的应用，求所有将字符串分割为回文子串的方案。\n\n核心思想：回溯，尝试所有可能的分割位置，检查是否为回文。\n\n详细步骤：\n1. 定义递归函数backtrack(start, currentPartition)\n2. 如果start等于s的长度：\n   - 将currentPartition加入结果\n   - 返回\n3. 从start开始遍历到末尾：\n   - 如果s[start:i+1]是回文：\n     - 将其加入currentPartition\n     - 递归调用backtrack(i+1, ...)\n     - 回溯\n\n回文判断优化：\n可以使用动态规划预处理，dp[i][j]表示s[i:j+1]是否为回文。\n\n时间复杂度：O(n × 2^n)，共有O(2^n)种分割方式，每种需要O(n)判断回文\n空间复杂度：O(n)，递归栈空间")

add_problem(51, "N 皇后 (N-Queens)", "困难", "https://leetcode.cn/problems/n-queens/", "回溯",
"按照国际象棋的规则，皇后可以攻击与之处在同一行或同一列或同一斜线上的棋子。\nn 皇后问题 研究的是如何将 n 个皇后放置在 n×n 的棋盘上，并且使皇后彼此之间不能相互攻击。\n给你一个整数 n ，返回所有不同的 n 皇后问题 的解决方案。\n每一种解法包含一个不同的 n 皇后问题 的棋子放置方案，该方案中 'Q' 和 '.' 分别代表了皇后和空位。\n\n示例：\n输入：n = 4\n输出：[[\".Q..\",\"...Q\",\"Q...\",\"..Q.\"],[\"..Q.\",\"Q...\",\"...Q\",\".Q..\"]]",
"这道题是回溯法的经典难题，N皇后问题。\n\n核心思想：逐行放置皇后，检查是否冲突。\n\n冲突检测：\n- 列冲突：记录每列是否已有皇后\n- 对角线冲突（左上到右下）：行-列的值相同\n- 反对角线冲突（右上到左下）：行+列的值相同\n\n详细步骤：\n1. 定义递归函数backtrack(row)\n2. 如果row等于n：\n   - 找到一个解，加入结果\n   - 返回\n3. 遍历当前行的每一列：\n   - 如果在该位置放置皇后不冲突：\n     - 放置皇后，标记列、对角线、反对角线\n     - 递归调用backtrack(row+1)\n     - 回溯，移除标记\n\n使用三个集合来记录：\n- cols：已占用的列\n- diagonals：已占用的对角线（row - col）\n- antiDiagonals：已占用的反对角线（row + col）\n\n时间复杂度：O(n!)，第一行n种选择，第二行最多n-1种，以此类推\n空间复杂度：O(n)，递归栈深度")

print(f"已添加 {len(PROBLEMS)} 道题目")

# ========== 10. 二分查找 (6题) ==========
add_problem(35, "搜索插入位置 (Search Insert Position)", "简单", "https://leetcode.cn/problems/search-insert-position/", "二分查找",
"给定一个排序数组和一个目标值，在数组中找到目标值，并返回其索引。如果目标值不存在于数组中，返回它将会被按顺序插入的位置。\n请必须使用时间复杂度为 O(log n) 的算法。\n\n示例：\n输入：nums = [1,3,5,6], target = 5\n输出：2",
"这道题是二分查找的基础题。\n\n核心思想：二分查找，如果找不到，返回left（即插入位置）。\n\n详细步骤：\n1. 初始化left = 0, right = len(nums) - 1\n2. 当left <= right时循环：\n   - mid = (left + right) // 2\n   - 如果nums[mid] == target，返回mid\n   - 如果nums[mid] < target，left = mid + 1\n   - 否则，right = mid - 1\n3. 返回left\n\n为什么返回left？\n当循环结束时，left是大于target的第一个位置，也就是target应该插入的位置。\n\n时间复杂度：O(log n)\n空间复杂度：O(1)")

add_problem(74, "搜索二维矩阵 (Search a 2D Matrix)", "中等", "https://leetcode.cn/problems/search-a-2d-matrix/", "二分查找",
"给你一个满足下述两条属性的 m x n 整数矩阵：\n每行中的整数从左到右按非严格递增顺序排列。\n每行的第一个整数大于前一行的最后一个整数。\n现给你一个整数 target ，如果 target 在矩阵中，返回 true ；否则，返回 false 。\n\n示例：\n输入：matrix = [[1,3,5,7],[10,11,16,20],[23,30,34,60]], target = 3\n输出：true",
"这道题是二分查找在二维矩阵上的应用。\n\n核心思想：由于矩阵的行优先有序，可以将二维矩阵视为一维有序数组进行二分查找。\n\n一维索引到二维坐标的转换：\n- row = index // n\n- col = index % n\n\n详细步骤：\n1. 初始化left = 0, right = m × n - 1\n2. 当left <= right时循环：\n   - mid = (left + right) // 2\n   - row = mid // n, col = mid % n\n   - 如果matrix[row][col] == target，返回true\n   - 如果matrix[row][col] < target，left = mid + 1\n   - 否则，right = mid - 1\n3. 返回false\n\n时间复杂度：O(log(m × n))\n空间复杂度：O(1)")

add_problem(34, "在排序数组中查找元素的第一个和最后一个位置 (Find First and Last Position of Element in Sorted Array)", "中等", "https://leetcode.cn/problems/find-first-and-last-position-of-element-in-sorted-array/", "二分查找",
"给你一个按照非递减顺序排列的整数数组 nums，和一个目标值 target。请你找出给定目标值在数组中的开始位置和结束位置。\n如果数组中不存在目标值 target，返回 [-1, -1]。\n你必须设计并实现时间复杂度为 O(log n) 的算法解决此问题。\n\n示例：\n输入：nums = [5,7,7,8,8,10], target = 8\n输出：[3,4]",
"这道题是二分查找的进阶，需要找到目标值的左右边界。\n\n核心思想：两次二分查找，分别找左边界和右边界。\n\n找左边界的二分查找：\n1. 初始化left = 0, right = len(nums) - 1\n2. 当left <= right时循环：\n   - mid = (left + right) // 2\n   - 如果nums[mid] < target，left = mid + 1\n   - 否则，right = mid - 1\n3. 检查left是否越界以及nums[left]是否等于target\n\n找右边界的二分查找：\n1. 初始化left = 0, right = len(nums) - 1\n2. 当left <= right时循环：\n   - mid = (left + right) // 2\n   - 如果nums[mid] <= target，left = mid + 1\n   - 否则，right = mid - 1\n3. 检查right是否越界以及nums[right]是否等于target\n\n注意两个二分查找的区别：\n- 找左边界时，nums[mid] == target时也要收缩右边界\n- 找右边界时，nums[mid] == target时也要收缩左边界\n\n时间复杂度：O(log n)，两次二分查找\n空间复杂度：O(1)")

add_problem(33, "搜索旋转排序数组 (Search in Rotated Sorted Array)", "中等", "https://leetcode.cn/problems/search-in-rotated-sorted-array/", "二分查找",
"整数数组 nums 按升序排列，数组中的值 互不相同 。\n在传递给函数之前，nums 在预先未知的某个下标 k（0 <= k < nums.length）上进行了 旋转，使数组变为 [nums[k], nums[k+1], ..., nums[n-1], nums[0], nums[1], ..., nums[k-1]]（下标 从 0 开始 计数）。例如， [0,1,2,4,5,6,7] 在下标 3 处经旋转后可能变为 [4,5,6,7,0,1,2] 。\n给你 旋转后 的数组 nums 和一个整数 target ，如果 nums 中存在这个目标值 target ，则返回它的下标，否则返回 -1 。\n你必须设计一个时间复杂度为 O(log n) 的算法解决此问题。\n\n示例：\n输入：nums = [4,5,6,7,0,1,2], target = 0\n输出：4",
"这道题是二分查找的变形，数组被旋转过，但部分仍然有序。\n\n核心思想：判断哪一半是有序的，然后决定搜索范围。\n\n详细步骤：\n1. 初始化left = 0, right = len(nums) - 1\n2. 当left <= right时循环：\n   - mid = (left + right) // 2\n   - 如果nums[mid] == target，返回mid\n   - 判断哪一半是有序的：\n     - 如果nums[left] <= nums[mid]，左半部分有序\n       - 如果target在[nums[left], nums[mid])范围内，right = mid - 1\n       - 否则，left = mid + 1\n     - 否则，右半部分有序\n       - 如果target在(nums[mid], nums[right]]范围内，left = mid + 1\n       - 否则，right = mid - 1\n3. 返回-1\n\n关键：虽然整个数组不是有序的，但总有一半是有序的，可以根据target是否在该有序范围内来决定搜索方向。\n\n时间复杂度：O(log n)\n空间复杂度：O(1)")

add_problem(153, "寻找旋转排序数组中的最小值 (Find Minimum in Rotated Sorted Array)", "中等", "https://leetcode.cn/problems/find-minimum-in-rotated-sorted-array/", "二分查找",
"已知一个长度为 n 的数组，预先按照升序排列，经由 1 到 n 次 旋转 后，得到输入数组。例如，原数组 nums = [0,1,2,4,5,6,7] 在变化后可能得到：\n若旋转 4 次，则可以得到 [4,5,6,7,0,1,2]\n若旋转 7 次，则可以得到 [0,1,2,4,5,6,7]\n注意，数组 [a[0], a[1], a[2], ..., a[n-1]] 旋转一次 的结果为数组 [a[n-1], a[0], a[1], a[2], ..., a[n-2]] 。\n给你一个元素值 互不相同 的数组 nums ，它原来是一个升序排列的数组，并按上述情形进行了多次旋转。请你找出并返回数组中的 最小元素 。\n你必须设计一个时间复杂度为 O(log n) 的算法解决此问题。\n\n示例：\n输入：nums = [3,4,5,1,2]\n输出：1\n解释：原数组为 [1,2,3,4,5] ，旋转 3 次得到输入数组。",
"这道题是二分查找在旋转数组上找最小值。\n\n核心思想：最小值位于无序的那一半。\n\n详细步骤：\n1. 初始化left = 0, right = len(nums) - 1\n2. 当left < right时循环：\n   - mid = (left + right) // 2\n   - 如果nums[mid] > nums[right]，最小值在右半部分，left = mid + 1\n   - 如果nums[mid] < nums[right]，最小值在左半部分（包括mid），right = mid\n   - 如果nums[mid] == nums[right]，无法判断，right--\n3. 返回nums[left]\n\n为什么与nums[right]比较？\n因为最小值一定在某个递增区间的边界，通过比较可以确定最小值在哪一半。\n\n时间复杂度：O(log n)，最坏情况（有重复元素）O(n)\n空间复杂度：O(1)")

add_problem(4, "寻找两个正序数组的中位数 (Median of Two Sorted Arrays)", "困难", "https://leetcode.cn/problems/median-of-two-sorted-arrays/", "二分查找",
"给定两个大小分别为 m 和 n 的正序（从小到大）数组 nums1 和 nums2。请你找出并返回这两个正序数组的 中位数 。\n算法的时间复杂度应该为 O(log (m+n)) 。\n\n示例：\n输入：nums1 = [1,3], nums2 = [2]\n输出：2.00000\n解释：合并数组 = [1,2,3] ，中位数 2",
"这道题是二分查找的经典难题，要求在O(log(m+n))时间内找到两个有序数组的中位数。\n\n核心思想：二分查找第k小的元素。\n\n转化为求第k小的问题：\n- 如果总长度为奇数，中位数是第(m+n)//2 + 1小的元素\n- 如果总长度为偶数，中位数是第(m+n)//2和第(m+n)//2 + 1小的元素的平均值\n\n详细步骤（找第k小的元素）：\n1. 每次从两个数组中各取前k//2个元素比较\n2. 如果nums1[k//2 - 1] < nums2[k//2 - 1]，说明nums1的前k//2个元素不可能是第k小的元素，可以排除\n3. 否则，排除nums2的前k//2个元素\n4. 更新k，重复直到k=1\n5. 返回min(nums1[0], nums2[0])\n\n边界处理：\n- 如果一个数组为空，直接从另一个数组找第k小的元素\n- 如果k=1，返回两个数组首元素的较小值\n\n时间复杂度：O(log(m+n))\n空间复杂度：O(1)")

print(f"已添加 {len(PROBLEMS)} 道题目")

# ========== 11. 栈 (5题) ==========
add_problem(20, "有效的括号 (Valid Parentheses)", "简单", "https://leetcode.cn/problems/valid-parentheses/", "栈",
"给定一个只包括 '('，')'，'{'，'}'，'['，']' 的字符串 s ，判断字符串是否有效。\n有效字符串需满足：\n左括号必须用相同类型的右括号闭合。\n左括号必须以正确的顺序闭合。\n每个右括号都有一个对应的相同类型的左括号。\n\n示例：\n输入：s = \"()\"\n输出：true",
"这道题是栈的经典应用，判断括号是否匹配。\n\n核心思想：使用栈，遇到左括号入栈，遇到右括号检查栈顶。\n\n详细步骤：\n1. 创建一个空栈\n2. 遍历字符串中的每个字符：\n   - 如果是左括号（'(', '{', '['），入栈\n   - 如果是右括号（')', '}', ']'）：\n     - 如果栈为空，返回false\n     - 弹出栈顶，检查是否与当前右括号匹配\n     - 如果不匹配，返回false\n3. 遍历结束后，如果栈为空，返回true；否则返回false\n\n时间复杂度：O(n)\n空间复杂度：O(n)，栈的空间")

add_problem(155, "最小栈 (Min Stack)", "中等", "https://leetcode.cn/problems/min-stack/", "栈",
"设计一个支持 push ，pop ，top 操作，并能在常数时间内检索到最小元素的栈。\n实现 MinStack 类:\nMinStack() 初始化堆栈对象。\nvoid push(int val) 将元素val推入堆栈。\nvoid pop() 删除堆栈顶部的元素。\nint top() 获取堆栈顶部的元素。\nint getMin() 获取堆栈中的最小元素。\n\n示例：\n输入：\n[\"MinStack\",\"push\",\"push\",\"push\",\"getMin\",\"pop\",\"top\",\"getMin\"]\n[[],[-2],[0],[-3],[],[],[],[]]\n输出：\n[null,null,null,null,-3,null,0,-2]",
"这道题要求在常数时间内获取栈的最小值。\n\n核心思想：使用辅助栈存储每个状态下的最小值。\n\n方法一：辅助栈\n- 使用两个栈，一个存储所有元素，一个存储当前最小值\n- push时，如果新元素小于等于辅助栈顶，也push到辅助栈\n- pop时，如果弹出的元素等于辅助栈顶，辅助栈也pop\n\n方法二：一个栈存储(值, 当前最小值)的元组\n- 每个元素入栈时，同时记录当前栈中的最小值\n- getMin直接返回栈顶的当前最小值\n\n时间复杂度：所有操作都是O(1)\n空间复杂度：O(n)")

add_problem(394, "字符串解码 (Decode String)", "中等", "https://leetcode.cn/problems/decode-string/", "栈",
"给定一个经过编码的字符串，返回它解码后的字符串。\n编码规则为: k[encoded_string]，表示其中方括号内部的 encoded_string 正好重复 k 次。注意 k 保证为正整数。\n你可以认为输入字符串总是有效的；输入字符串中没有额外的空格，且输入的方括号总是符合格式要求的。\n此外，你可以认为原始数据不包含数字，所有的数字只表示重复的次数 k ，例如不会出现像 3a 或 2[4] 的输入。\n\n示例：\n输入：s = \"3[a]2[bc]\"\n输出：\"aaabcbc\"",
"这道题是栈的应用，处理嵌套的编码字符串。\n\n核心思想：使用栈处理嵌套结构，遇到']'时解码当前层。\n\n详细步骤：\n1. 使用两个栈，一个存储重复次数，一个存储之前的字符串\n2. 遍历字符串：\n   - 如果是数字，计算完整的数字（可能是多位数）\n   - 如果是'['，将当前数字和当前字符串分别入栈，重置\n   - 如果是']'，弹出重复次数和之前的字符串，解码当前层并拼接\n   - 如果是字母，加入当前字符串\n3. 返回最终字符串\n\n时间复杂度：O(n × 最大重复次数)\n空间复杂度：O(n)")

add_problem(739, "每日温度 (Daily Temperatures)", "中等", "https://leetcode.cn/problems/daily-temperatures/", "栈",
"给定一个整数数组 temperatures ，表示每天的温度，返回一个数组 answer ，其中 answer[i] 是指对于第 i 天，下一个更高温度出现在几天后。如果气温在这之后都不会升高，请在该位置用 0 来代替。\n\n示例：\n输入: temperatures = [73,74,75,71,69,72,76,73]\n输出: [1,1,4,2,1,1,0,0]",
"这道题是单调栈的经典应用，找下一个更大元素。\n\n核心思想：使用单调递减栈，栈中存储索引。对于每个元素，弹出所有比它小的元素，这些元素的\"下一个更大元素\"就是当前元素。\n\n详细步骤：\n1. 初始化结果数组，全为0\n2. 初始化空栈\n3. 遍历温度数组：\n   - 当栈不为空且当前温度大于栈顶索引对应的温度：\n     - 弹出栈顶索引idx\n     - answer[idx] = i - idx（当前天数减去栈顶天数）\n   - 将当前索引入栈\n4. 返回结果数组\n\n为什么栈保持单调递减？\n这样栈中的元素都是\"等待找到下一个更大元素\"的，一旦发现更大的元素，就可以立即计算答案。\n\n时间复杂度：O(n)，每个元素最多入栈出栈一次\n空间复杂度：O(n)")

add_problem(84, "柱状图中最大的矩形 (Largest Rectangle in Histogram)", "困难", "https://leetcode.cn/problems/largest-rectangle-in-histogram/", "栈",
"给定 n 个非负整数，用来表示柱状图中各个柱子的高度。每个柱子彼此相邻，且宽度为 1 。\n求在该柱状图中，能够勾勒出来的矩形的最大面积。\n\n示例：\n输入：heights = [2,1,5,6,2,3]\n输出：10\n解释：最大的矩形为图中红色区域，面积为 10",
"这道题是单调栈的经典难题，求柱状图中最大的矩形面积。\n\n核心思想：对于每个柱子，找到左右两边第一个比它小的柱子，确定矩形的宽度。\n\n方法一：暴力（O(n²)）\n对于每个柱子，向两边扩展直到遇到更小的柱子，计算面积。\n\n方法二：单调栈（O(n)）\n使用单调递增栈，找到每个柱子的左边界和右边界。\n\n详细步骤：\n1. 在原数组两边添加0（方便处理边界）\n2. 初始化栈，加入0（左边界）\n3. 遍历每个柱子（从1到n+1）：\n   - 当当前柱子高度小于栈顶柱子高度：\n     - 弹出栈顶mid作为矩形的高\n     - 新的栈顶就是左边界left\n     - 右边界就是当前索引i\n     - 面积 = heights[mid] × (i - left - 1)\n     - 更新最大面积\n   - 将当前索引入栈\n4. 返回最大面积\n\n时间复杂度：O(n)，每个元素最多入栈出栈一次\n空间复杂度：O(n)")

print(f"已添加 {len(PROBLEMS)} 道题目")

# ========== 12. 堆 (3题) ==========
add_problem(215, "数组中的第K个最大元素 (Kth Largest Element in an Array)", "中等", "https://leetcode.cn/problems/kth-largest-element-in-an-array/", "堆",
"给定整数数组 nums 和整数 k，请返回数组中第 k 个最大的元素。\n请注意，你需要找的是数组排序后的第 k 个最大的元素，而不是第 k 个不同的元素。\n\n示例：\n输入: [3,2,1,5,6,4] 和 k = 2\n输出: 5",
"这道题要求在数组中找到第k大的元素。\n\n方法一：排序（O(n log n)）\n直接排序，取倒数第k个元素。\n\n方法二：最小堆（O(n log k)）\n维护一个大小为k的最小堆，堆顶就是第k大的元素。\n\n方法三：快速选择（O(n)平均，O(n²)最坏）\n利用快速排序的partition思想，在平均情况下可以达到O(n)。\n\n时间复杂度：\n- 排序：O(n log n)\n- 最小堆：O(n log k)\n- 快速选择：O(n)平均\n\n空间复杂度：\n- 排序：O(1)或O(log n)\n- 最小堆：O(k)\n- 快速选择：O(log n)递归栈")

add_problem(347, "前 K 个高频元素 (Top K Frequent Elements)", "中等", "https://leetcode.cn/problems/top-k-frequent-elements/", "堆",
"给你一个整数数组 nums 和一个整数 k ，请你返回其中出现频率前 k 高的元素。你可以按 任意顺序 返回答案。\n\n示例：\n输入: nums = [1,1,1,2,2,3], k = 2\n输出: [1,2]",
"这道题是Top K问题的经典应用。\n\n核心思想：先用哈希表统计频率，然后用堆找前k个高频元素。\n\n方法一：哈希表 + 排序（O(n log n)）\n1. 统计每个元素的频率\n2. 按频率排序\n3. 取前k个\n\n方法二：哈希表 + 最小堆（O(n log k)）\n1. 统计每个元素的频率\n2. 维护一个大小为k的最小堆，按频率比较\n3. 堆中的元素就是前k个高频元素\n\n方法三：桶排序（O(n)）\n1. 统计每个元素的频率\n2. 使用频率作为桶的索引，将元素放入对应桶中\n3. 从高频到低频收集元素\n\n时间复杂度：\n- 哈希表+排序：O(n log n)\n- 哈希表+最小堆：O(n log k)\n- 桶排序：O(n)\n\n空间复杂度：O(n)")

add_problem(295, "数据流的中位数 (Find Median from Data Stream)", "困难", "https://leetcode.cn/problems/find-median-from-data-stream/", "堆",
"中位数是有序整数列表中的中间值。如果列表的大小是偶数，则没有中间值，中位数是两个中间值的平均值。\n例如 arr = [2,3,4] 的中位数是 3 。\n例如 arr = [2,3] 的中位数是 (2 + 3) / 2 = 2.5 。\n实现 MedianFinder 类:\nMedianFinder() 初始化 MedianFinder 对象。\nvoid addNum(int num) 将数据流中的整数 num 添加到数据结构中。\ndouble findMedian() 返回到目前为止所有元素的中位数。\n\n示例：\n输入\n[\"MedianFinder\", \"addNum\", \"addNum\", \"findMedian\", \"addNum\", \"findMedian\"]\n[[], [1], [2], [], [3], []]\n输出\n[null, null, null, 1.5, null, 2.0]",
"这道题要求在数据流中动态维护中位数。\n\n核心思想：使用两个堆，一个大根堆（存储较小的一半），一个小根堆（存储较大的一半）。\n\n设计：\n- 大根堆maxHeap：存储较小的一半元素，堆顶是这半部分的最大值\n- 小根堆minHeap：存储较大的一半元素，堆顶是这半部分的最小值\n- 保持两个堆的大小差不超过1\n\naddNum(num)：\n1. 如果maxHeap为空或num小于等于maxHeap顶，加入maxHeap\n2. 否则，加入minHeap\n3. 平衡两个堆的大小（确保大小差不超过1）\n\nfindMedian()：\n1. 如果两个堆大小相等，返回(maxHeap顶 + minHeap顶) / 2\n2. 否则，返回元素较多的那个堆的堆顶\n\n时间复杂度：\n- addNum：O(log n)\n- findMedian：O(1)\n\n空间复杂度：O(n)")

print(f"已添加 {len(PROBLEMS)} 道题目")

# ========== 13. 贪心算法 (4题) ==========
add_problem(121, "买卖股票的最佳时机 (Best Time to Buy and Sell Stock)", "简单", "https://leetcode.cn/problems/best-time-to-buy-and-sell-stock/", "贪心算法",
"给定一个数组 prices ，它的第 i 个元素 prices[i] 表示一支给定股票第 i 天的价格。\n你只能选择 某一天 买入这只股票，并选择在 未来的某一个不同的日子 卖出该股票。设计一个算法来计算你所能获取的最大利润。\n返回你可以从这笔交易中获取的最大利润。如果你不能获取任何利润，返回 0 。\n\n示例：\n输入：[7,1,5,3,6,4]\n输出：5\n解释：在第 2 天（股票价格 = 1）的时候买入，在第 5 天（股票价格 = 6）的时候卖出，最大利润 = 6-1 = 5 。",
"这道题是股票买卖的基础题，只能进行一次交易。\n\n核心思想：在最低点买入，在之后的最高点卖出。\n\n详细步骤：\n1. 初始化minPrice为无穷大，maxProfit为0\n2. 遍历每一天的价格：\n   - 如果当前价格小于minPrice，更新minPrice\n   - 否则，计算当前价格与minPrice的差，更新maxProfit\n3. 返回maxProfit\n\n为什么这样有效？\n对于每一天，我们只需要知道之前最低的买入价格，以及当前能卖出的最大利润。\n\n时间复杂度：O(n)\n空间复杂度：O(1)")

add_problem(55, "跳跃游戏 (Jump Game)", "中等", "https://leetcode.cn/problems/jump-game/", "贪心算法",
"给定一个非负整数数组 nums ，你最初位于数组的 第一个下标 。\n数组中的每个元素代表你在该位置可以跳跃的最大长度。\n判断你是否能够到达最后一个下标。\n\n示例：\n输入：nums = [2,3,1,1,4]\n输出：true\n解释：可以先跳 1 步，从下标 0 到达下标 1, 然后再从下标 1 跳 3 步到达最后一个下标。",
"这道题是贪心算法的经典应用，判断是否能到达终点。\n\n核心思想：维护当前能到达的最远位置。\n\n详细步骤：\n1. 初始化maxReach = 0\n2. 遍历数组的每个位置i：\n   - 如果i > maxReach，说明当前位置无法到达，返回false\n   - 更新maxReach = max(maxReach, i + nums[i])\n   - 如果maxReach >= n-1，返回true\n3. 返回true\n\n为什么这样有效？\nmaxReach表示从起点出发，经过所有已遍历的位置，能够到达的最远位置。如果当前位置已经超过maxReach，说明无法到达这里，更无法到达终点。\n\n时间复杂度：O(n)\n空间复杂度：O(1)")

add_problem(45, "跳跃游戏 II (Jump Game II)", "中等", "https://leetcode.cn/problems/jump-game-ii/", "贪心算法",
"给定一个长度为 n 的 0 索引整数数组 nums。初始位置为 nums[0]。\n每个元素 nums[i] 表示从索引 i 向前跳转的最大长度。换句话说，如果你在 nums[i] 处，你可以跳转到任意 nums[i + j] 处:\n0 <= j <= nums[i]\ni + j < n\n返回到达 nums[n - 1] 的最小跳跃次数。生成的测试用例可以到达 nums[n - 1]。\n\n示例：\n输入: nums = [2,3,1,1,4]\n输出: 2\n解释: 跳到最后一个位置的最小跳跃数是 2。\n     从下标为 0 跳到下标为 1 的位置，跳 1 步，然后跳 3 步到达数组的最后一个位置。",
"这道题是跳跃游戏的进阶版，要求最小跳跃次数。\n\n核心思想：贪心，每次在当前可达范围内选择能跳最远的位置作为下一次起跳点。\n\n详细步骤：\n1. 初始化jumps = 0, currentEnd = 0, farthest = 0\n2. 遍历数组（除了最后一个位置）：\n   - 更新farthest = max(farthest, i + nums[i])\n   - 如果i == currentEnd：\n     - jumps++\n     - currentEnd = farthest\n3. 返回jumps\n\n为什么这样有效？\n- currentEnd表示当前这一跳能到达的最远位置\n- 在到达currentEnd之前，我们收集所有位置能到达的最远距离farthest\n- 当到达currentEnd时，必须进行下一次跳跃，跳到farthest\n\n时间复杂度：O(n)\n空间复杂度：O(1)")

add_problem(763, "划分字母区间 (Partition Labels)", "中等", "https://leetcode.cn/problems/partition-labels/", "贪心算法",
"给你一个字符串 s 。我们要把这个字符串划分为尽可能多的片段，同一字母最多出现在一个片段中。\n注意，划分结果需要满足：将所有划分结果按顺序连接，得到的字符串仍然是 s 。\n返回一个表示每个字符串片段的长度的列表。\n\n示例：\n输入：s = \"ababcbacadefegdehijhklij\"\n输出：[9,7,8]\n解释：\n划分结果为 \"ababcbaca\", \"defegde\", \"hijhklij\" 。\n每个字母最多出现在一个片段中。\n像 \"ababcbacadefegde\", \"hijhklij\" 的划分是错误的，因为划分的片段数较少。",
"这道题是贪心算法的应用，划分字符串使得每个字母只出现在一个片段中。\n\n核心思想：对于每个字母，找到它最后出现的位置，这决定了片段的边界。\n\n详细步骤：\n1. 使用哈希表记录每个字母最后出现的索引\n2. 初始化start = 0, end = 0\n3. 遍历字符串的每个位置i：\n   - 更新end = max(end, lastIndex[s[i]])\n   - 如果i == end：\n     - 找到了一个片段，长度为end - start + 1\n     - 更新start = i + 1\n4. 返回所有片段的长度\n\n为什么这样有效？\n每个字母必须完全包含在一个片段中，所以片段的右边界至少要是当前所有已遍历字母的最后位置的最大值。当遍历位置到达这个边界时，就可以切分。\n\n时间复杂度：O(n)\n空间复杂度：O(1)，最多26个字母")

print(f"已添加 {len(PROBLEMS)} 道题目")

# ========== Word文档生成代码 ==========
def create_word_document():
    """创建Word文档"""
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    
    # 添加标题
    title = doc.add_heading('LeetCode 热题 HOT 100 完整题解', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加副标题
    subtitle = doc.add_paragraph('包含全部100道题的详细题目描述与解题思路')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 添加日期
    date_para = doc.add_paragraph('整理时间：2026年3月')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 添加目录标题
    toc_title = doc.add_heading('目录', level=1)
    
    # 获取所有类别
    categories = sorted(set([p["category"] for p in PROBLEMS]))
    
    # 添加目录内容
    for i, cat in enumerate(categories, 1):
        cat_problems = [p for p in PROBLEMS if p["category"] == cat]
        cat_para = doc.add_paragraph(f'{i}. {cat}（{len(cat_problems)}题）')
        cat_para.paragraph_format.left_indent = Inches(0.3)
    
    doc.add_page_break()
    
    # 按类别分组添加题目
    for category in categories:
        # 添加类别标题
        cat_heading = doc.add_heading(category, level=1)
        
        cat_problems = [p for p in PROBLEMS if p["category"] == category]
        cat_problems.sort(key=lambda x: x['id'])
        
        for problem in cat_problems:
            # 添加题目标题
            p_heading = doc.add_heading(f"{problem['id']}. {problem['title']}", level=2)
            
            # 添加题目信息
            info_para = doc.add_paragraph()
            info_para.add_run('难度：').bold = True
            difficulty = problem["difficulty"]
            difficulty_run = info_para.add_run(difficulty)
            if difficulty == "简单":
                difficulty_run.font.color.rgb = RGBColor(0, 176, 80)
            elif difficulty == "中等":
                difficulty_run.font.color.rgb = RGBColor(255, 192, 0)
            else:  # 困难
                difficulty_run.font.color.rgb = RGBColor(255, 0, 0)
            info_para.add_run(f'    链接：').bold = True
            info_para.add_run(problem["link"])
            
            # 添加题目描述标题
            desc_title = doc.add_paragraph()
            desc_title.add_run('题目描述：').bold = True
            
            # 添加题目描述
            desc_para = doc.add_paragraph(problem["description"])
            desc_para.paragraph_format.first_line_indent = Inches(0.3)
            
            # 添加解题思路标题
            sol_title = doc.add_paragraph()
            sol_title.add_run('解题思路：').bold = True
            
            # 添加解题思路
            solution_para = doc.add_paragraph(problem["solution"])
            solution_para.paragraph_format.first_line_indent = Inches(0.3)
            
            # 添加分隔线
            doc.add_paragraph('_' * 60)
    
    return doc

def main():
    """主函数"""
    print("正在生成LeetCode HOT100完整题解Word文档...")
    
    # 创建Word文档
    doc = create_word_document()
    
    # 保存文档
    output_path = '/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx'
    doc.save(output_path)
    
    print(f"文档已生成：{output_path}")
    print(f"共包含 {len(PROBLEMS)} 道题")

# ========== 补充题目以达到100道 ==========
add_problem(338, "比特位计数 (Counting Bits)", "简单", "https://leetcode.cn/problems/counting-bits/", "技巧",
"给你一个整数 n ，对于 0 <= i <= n 中的每个 i ，计算其二进制表示中 1 的个数 ，返回一个长度为 n + 1 的数组 ans 作为答案。\n\n示例：\n输入：n = 2\n输出：[0,1,1]\n解释：\n0 --> 0\n1 --> 1\n2 --> 10",
"这道题是动态规划与位运算的结合。\n\n核心思想：利用之前计算过的结果。\n\n状态转移方程：\ndp[i] = dp[i >> 1] + (i & 1)\n\n解释：\n- i >> 1相当于i/2\n- i & 1表示i的最低位是否为1\n\n或者使用：\ndp[i] = dp[i & (i-1)] + 1\ni & (i-1)可以消除i的最后一位1\n\n时间复杂度：O(n)\n空间复杂度：O(1)，不包括输出数组")

# ========== 补充题目以达到100道 ==========
add_problem(338, "比特位计数 (Counting Bits)", "简单", "https://leetcode.cn/problems/counting-bits/", "技巧",
"给你一个整数 n ，对于 0 <= i <= n 中的每个 i ，计算其二进制表示中 1 的个数 ，返回一个长度为 n + 1 的数组 ans 作为答案。\n\n示例：\n输入：n = 2\n输出：[0,1,1]\n解释：\n0 --> 0\n1 --> 1\n2 --> 10",
"这道题是动态规划与位运算的结合。\n\n核心思想：利用之前计算过的结果。\n\n状态转移方程：\ndp[i] = dp[i >> 1] + (i & 1)\n\n解释：\n- i >> 1相当于i/2\n- i & 1表示i的最低位是否为1\n\n时间复杂度：O(n)\n空间复杂度：O(1)，不包括输出数组")

add_problem(75, "颜色分类 (Sort Colors)", "中等", "https://leetcode.cn/problems/sort-colors/", "普通数组",
"给定一个包含红色、白色和蓝色、共 n 个元素的数组 nums ，原地对它们进行排序，使得相同颜色的元素相邻，并按照红色、白色、蓝色顺序排列。\n我们使用整数 0、 1 和 2 分别表示红色、白色和蓝色。\n必须在不使用库内置的 sort 函数的情况下解决这个问题。\n\n示例：\n输入：nums = [2,0,2,1,1,0]\n输出：[0,0,1,1,2,2]",
"这道题是荷兰国旗问题，要求原地排序三种颜色。\n\n核心思想：三指针法，将数组分为三部分：0在左边，1在中间，2在右边。\n\n详细步骤（单指针两次遍历）：\n1. 第一次遍历，将所有0移到前面\n2. 第二次遍历，将所有1移到0的后面\n\n或者（双指针一次遍历）：\n1. p0指向0应该放的位置，p2指向2应该放的位置\n2. 遍历数组，遇到0交换到p0，遇到2交换到p2\n\n时间复杂度：O(n)\n空间复杂度：O(1)")

add_problem(76, "最小覆盖子串 (Minimum Window Substring)", "困难", "https://leetcode.cn/problems/minimum-window-substring/", "滑动窗口",
"给你一个字符串 s 、一个字符串 t 。返回 s 中涵盖 t 所有字符的最小子串。如果不存在符合条件的子串，则返回空字符串 \"\" 。\n注意：\n对于 t 中重复字符，我们寻找的子字符串中该字符数量必须不少于 t 中该字符数量。\n如果 s 中存在这样的子串，我们保证它是唯一的答案。\n\n示例：\n输入：s = \"ADOBECODEBANC\", t = \"ABC\"\n输出：\"BANC\"\n解释：最小覆盖子串 \"BANC\" 包含来自字符串 t 的 'A'、'B' 和 'C'。",
"这道题是滑动窗口的经典难题。\n\n核心思想：使用滑动窗口，右指针扩展窗口直到包含所有需要的字符，然后左指针收缩窗口找到最小子串。\n\n详细步骤：\n1. 使用哈希表need记录t中每个字符需要的数量\n2. 使用哈希表window记录当前窗口中每个字符的数量\n3. 初始化left = 0, valid = 0（记录窗口中满足条件的字符数）\n4. 右指针遍历字符串s：\n   - 将s[right]加入window\n   - 如果s[right]在need中且数量满足要求，valid++\n   - 当valid等于need.size()时，尝试收缩左边界\n5. 返回最小子串\n\n时间复杂度：O(|s| + |t|)\n空间复杂度：O(|Σ|)，字符集大小")

add_problem(448, "找到所有数组中消失的数字 (Find All Numbers Disappeared in an Array)", "简单", "https://leetcode.cn/problems/find-all-numbers-disappeared-in-an-array/", "普通数组",
"给你一个含 n 个整数的数组 nums ，其中 nums[i] 在区间 [1, n] 内。请你找出所有在 [1, n] 范围内但没有出现在 nums 中的数字，并以数组的形式返回。\n\n示例：\n输入：nums = [4,3,2,7,8,2,3,1]\n输出：[5,6]",
"这道题要求在O(n)时间复杂度和不使用额外空间的情况下找出缺失的数字。\n\n核心思想：原地标记。对于出现的数字x，将索引x-1位置的数标记为负数。\n\n详细步骤：\n1. 遍历数组，对于每个数nums[i]，将索引abs(nums[i])-1位置的数变为负数\n2. 再次遍历数组，如果nums[i]为正数，说明i+1没有出现\n\n为什么这样有效？\n数组的索引范围是[0, n-1]，数字范围是[1, n]，正好是索引+1。\n\n时间复杂度：O(n)\n空间复杂度：O(1)")

add_problem(581, "最短无序连续子数组 (Shortest Unsorted Continuous Subarray)", "中等", "https://leetcode.cn/problems/shortest-unsorted-continuous-subarray/", "普通数组",
"给你一个整数数组 nums ，你需要找出一个 连续子数组 ，如果对这个子数组进行升序排序，那么整个数组都会变为升序排序。\n请你找出符合题意的 最短 子数组，并输出它的长度。\n\n示例：\n输入：nums = [2,6,4,8,10,9,15]\n输出：5\n解释：你只需要对 [6, 4, 8, 10, 9] 进行升序排序，那么整个表都会变为升序排序。",
"这道题要求找出需要排序的最短子数组。\n\n核心思想：找到数组中乱序的部分。\n\n方法一：排序比较\n1. 复制数组并排序\n2. 从左到右找到第一个不同的位置\n3. 从右到左找到第一个不同的位置\n4. 两者之间的距离就是答案\n\n方法二：一次遍历（O(n)）\n1. 从左到右遍历，维护最大值，如果当前值小于最大值，记录位置\n2. 从右到左遍历，维护最小值，如果当前值大于最小值，记录位置\n3. 两个位置之间的区域就是需要排序的部分\n\n时间复杂度：O(n)\n空间复杂度：O(1)")

add_problem(617, "合并二叉树 (Merge Two Binary Trees)", "简单", "https://leetcode.cn/problems/merge-two-binary-trees/", "二叉树",
"给你两棵二叉树： root1 和 root2 。\n想象一下，当你将其中一棵覆盖到另一棵之上时，两棵树上的一些节点将会重叠（而另一些不会）。你需要将这两棵树合并成一棵新二叉树。合并的规则是：如果两个节点重叠，那么将这两个节点的值相加作为合并后节点的新值；否则，不为 null 的节点将直接作为新二叉树的节点。\n返回合并后的二叉树。\n注意: 合并过程必须从两个树的根节点开始。\n\n示例：\n输入：root1 = [1,3,2,5], root2 = [2,1,3,null,4,null,7]\n输出：[3,4,5,5,4,null,7]",
"这道题是二叉树的递归操作。\n\n核心思想：递归合并两个二叉树。\n\n详细步骤：\n1. 如果root1为空，返回root2\n2. 如果root2为空，返回root1\n3. 创建新节点，值为root1.val + root2.val\n4. 递归合并左子树\n5. 递归合并右子树\n6. 返回新节点\n\n时间复杂度：O(min(m, n))\n空间复杂度：O(min(m, n))")

add_problem(234, "回文链表 (Palindrome Linked List)", "简单", "https://leetcode.cn/problems/palindrome-linked-list/", "链表",
"给你一个单链表的头节点 head ，请你判断该链表是否为回文链表。如果是，返回 true ；否则，返回 false 。\n\n示例：\n输入：head = [1,2,2,1]\n输出：true",
"这道题要求判断链表是否为回文。\n\n核心思想：找到中点 + 反转后半部分 + 比较\n\n详细步骤：\n1. 使用快慢指针找到链表的中点\n2. 反转链表的后半部分\n3. 比较前半部分和反转后的后半部分\n4. （可选）恢复链表\n\n时间复杂度：O(n)\n空间复杂度：O(1)")

add_problem(647, "回文子串 (Palindromic Substrings)", "中等", "https://leetcode.cn/problems/palindromic-substrings/", "动态规划",
"给你一个字符串 s ，请你统计并返回这个字符串中 回文子串 的数目。\n回文字符串 是正着读和倒过来读一样的字符串。\n子字符串 是字符串中的由连续字符组成的一个序列。\n\n示例：\n输入：s = \"aaa\"\n输出：6\n解释：6个回文子串: \"a\", \"a\", \"a\", \"aa\", \"aa\", \"aaa\"",
"这道题要求统计回文子串的数量。\n\n方法一：中心扩展\n枚举回文中心（每个字符或每对字符之间），向两边扩展。\n\n方法二：动态规划\ndp[i][j]表示s[i:j+1]是否为回文串\ndp[i][j] = (s[i] == s[j]) and (j - i < 2 or dp[i+1][j-1])\n\n时间复杂度：O(n²)\n空间复杂度：\n- 中心扩展：O(1)\n- 动态规划：O(n²)")

add_problem(406, "根据身高重建队列 (Queue Reconstruction by Height)", "中等", "https://leetcode.cn/problems/queue-reconstruction-by-height/", "贪心算法",
"假设有打乱顺序的一群人站成一个队列，数组 people 表示队列中一些人的属性（不一定按顺序）。每个 people[i] = [hi, ki] 表示第 i 个人的身高为 hi ，前面 正好 有 ki 个身高大于或等于 hi 的人。\n请你重新构造并返回输入数组 people 所表示的队列。返回的队列应该格式化为数组 queue ，其中 queue[j] = [hj, kj] 是队列中第 j 个人的属性（queue[0] 是排在队列前面的人）。\n\n示例：\n输入：people = [[7,0],[4,4],[7,1],[5,0],[6,1],[5,2]]\n输出：[[5,0],[7,0],[5,2],[6,1],[4,4],[7,1]]",
"这道题是贪心算法的经典题目。\n\n核心思想：先按身高降序排列，身高相同的按k升序排列，然后按k值插入到结果中。\n\n详细步骤：\n1. 按身高降序排序，身高相同的按k升序排序\n2. 创建一个空列表\n3. 遍历排序后的people，将每个人插入到索引为k的位置\n\n为什么这样有效？\n- 先处理高的人，他们的k值就是他们在队列中的位置\n- 后处理矮的人，插入到指定位置不会影响已经处理的高个子的人\n\n时间复杂度：O(n²)，插入操作是O(n)\n空间复杂度：O(n)")

add_problem(538, "把二叉搜索树转换为累加树 (Convert BST to Greater Tree)", "中等", "https://leetcode.cn/problems/convert-bst-to-greater-tree/", "二叉树",
"给出二叉 搜索 树的根节点，该树的节点值各不相同，请你将其转换为累加树（Greater Sum Tree），使每个节点 node 的新值等于原树中大于或等于 node.val 的值之和。\n\n示例：\n输入：[4,1,6,0,2,5,7,null,null,null,3,null,null,null,8]\n输出：[30,36,21,36,35,26,15,null,null,null,33,null,null,null,8]",
"这道题要求将BST转换为累加树。\n\n核心思想：反向中序遍历（右-根-左），维护一个累加和。\n\n详细步骤：\n1. 定义sum = 0\n2. 反向中序遍历BST（先右子树，再根，再左子树）\n3. 对于每个节点：\n   - sum += node.val\n   - node.val = sum\n\n为什么反向中序遍历？\nBST的中序遍历是升序，反向中序遍历是降序。这样遍历到的每个节点都是已经处理过的节点（值更大的节点）。\n\n时间复杂度：O(n)\n空间复杂度：O(h)")

add_problem(621, "任务调度器 (Task Scheduler)", "中等", "https://leetcode.cn/problems/task-scheduler/", "贪心算法",
"给你一个用字符数组 tasks 表示的 CPU 需要执行的任务列表。其中每个字母表示一种不同种类的任务。任务可以以任意顺序执行，并且每个任务都可以在 1 个单位时间内执行完。在任何一个单位时间，CPU 可以完成一个任务，或者处于待命状态。\n然而，两个 相同种类 的任务之间必须有长度为整数 n 的冷却时间，因此至少有连续 n 个单位时间内 CPU 在执行不同的任务，或者在待命状态。\n你需要计算完成所有任务所需要的 最短时间 。\n\n示例：\n输入：tasks = [\"A\",\"A\",\"A\",\"B\",\"B\",\"B\"], n = 2\n输出：8\n解释：A -> B -> (待命) -> A -> B -> (待命) -> A -> B",
"这道题是贪心算法的经典题目。\n\n核心思想：找到出现次数最多的任务，计算其需要的冷却时间。\n\n详细步骤：\n1. 统计每个任务出现的次数\n2. 找到最大出现次数maxCount，以及出现次数等于maxCount的任务数maxCountTasks\n3. 计算最少时间：\n   - 框架：(maxCount - 1) * (n + 1) + maxCountTasks\n   - 如果结果小于tasks的长度，说明不需要待命，答案就是tasks的长度\n\n时间复杂度：O(n)\n空间复杂度：O(1)")

add_problem(494, "目标和 (Target Sum)", "中等", "https://leetcode.cn/problems/target-sum/", "动态规划",
"给你一个非负整数数组 nums 和一个整数 target 。\n向数组中的每个整数前添加 '+' 或 '-' ，然后串联起所有整数，可以构造一个 表达式 ：\n例如，nums = [2, 1] ，可以在 2 之前添加 '+' ，在 1 之前添加 '-' ，然后串联起来得到表达式 \"+2-1\" 。\n返回可以通过上述方法构造的、运算结果等于 target 的不同 表达式 的数目。\n\n示例：\n输入：nums = [1,1,1,1,1], target = 3\n输出：5\n解释：一共有 5 种方法让最终目标和为 3。",
"这道题是0-1背包问题的变种。\n\n问题转化：\n设正数和为P，负数和为N，则：\nP + N = sum\nP - N = target\n解得：P = (sum + target) / 2\n\n所以问题转化为：从数组中选择一些数，使其和为(sum + target) / 2的方案数。\n\n状态定义：\ndp[i]表示和为i的方案数\n\n状态转移：\ndp[i] = dp[i] + dp[i - num] for all num in nums\n\n时间复杂度：O(n × target)\n空间复杂度：O(target)")

add_problem(560, "和为 K 的子数组 (Subarray Sum Equals K)", "中等", "https://leetcode.cn/problems/subarray-sum-equals-k/", "动态规划",
"给你一个整数数组 nums 和一个整数 k ，请你统计并返回 该数组中和为 k 的连续子数组的个数 。\n\n示例：\n输入：nums = [1,1,1], k = 2\n输出：2",
"这道题是前缀和 + 哈希表的经典应用。\n\n核心思想：前缀和 + 哈希表\n\n关键公式：\n- 子数组和 = prefixSum[j] - prefixSum[i-1] = k\n- 即 prefixSum[i-1] = prefixSum[j] - k\n\n详细步骤：\n1. 使用哈希表记录每个前缀和出现的次数\n2. 初始化hashMap[0] = 1\n3. 遍历数组，维护当前前缀和prefixSum：\n   - 在哈希表中查找prefixSum - k出现的次数，加到结果中\n   - 将当前prefixSum加入哈希表\n\n时间复杂度：O(n)\n空间复杂度：O(n)")

add_problem(437, "路径总和 III (Path Sum III)", "中等", "https://leetcode.cn/problems/path-sum-iii/", "动态规划",
"给定一个二叉树的根节点 root ，和一个整数 targetSum ，求该二叉树里节点值之和等于 targetSum 的 路径 的数目。\n路径 不需要从根节点开始，也不需要在叶子节点结束，但是路径方向必须是向下的（只能从父节点到子节点）。\n\n示例：\n输入：root = [10,5,-3,3,2,null,11,3,-2,null,1], targetSum = 8\n输出：3\n解释：和等于 8 的路径有 3 条。",
"这道题是二叉树的路径问题，前缀和 + 哈希表的应用。\n\n核心思想：前缀和 + 哈希表\n\n详细步骤：\n1. 定义哈希表prefixSumCount，记录每个前缀和出现的次数\n2. 初始化prefixSumCount[0] = 1\n3. 定义递归函数dfs(node, currSum)：\n   - 如果node为空，返回0\n   - currSum += node.val\n   - count = prefixSumCount.get(currSum - targetSum, 0)\n   - prefixSumCount[currSum] = prefixSumCount.get(currSum, 0) + 1\n   - count += dfs(node.left, currSum) + dfs(node.right, currSum)\n   - prefixSumCount[currSum] -= 1（回溯）\n   - 返回count\n\n时间复杂度：O(n)\n空间复杂度：O(h)")

add_problem(207, "课程表 (Course Schedule)", "中等", "https://leetcode.cn/problems/course-schedule/", "图论",
"你这个学期必须选修 numCourses 门课程，记为 0 到 numCourses - 1 。\n在选修某些课程之前需要一些先修课程。 先修课程按数组 prerequisites 给出，其中 prerequisites[i] = [ai, bi] ，表示如果要学习课程 ai 则 必须 先学习课程  bi 。\n请你判断是否可能完成所有课程的学习？如果可以，返回 true ；否则，返回 false 。\n\n示例：\n输入：numCourses = 2, prerequisites = [[1,0]]\n输出：true",
"这道题是经典的拓扑排序问题，判断有向图是否有环。\n\n核心思想：拓扑排序\n\n方法一：Kahn算法（BFS）\n1. 计算每个节点的入度\n2. 将所有入度为0的节点加入队列\n3. 当队列不为空时：\n   - 出队一个节点\n   - 将其所有邻居的入度减1\n   - 如果邻居入度变为0，入队\n4. 如果处理的节点数等于总节点数，说明无环\n\n方法二：DFS\n使用三种状态标记节点，如果在DFS过程中遇到状态为1的节点，说明有环。\n\n时间复杂度：O(V + E)\n空间复杂度：O(V + E)")

add_problem(210, "课程表 II (Course Schedule II)", "中等", "https://leetcode.cn/problems/course-schedule-ii/", "图论",
"现在你总共有 numCourses 门课需要选，记为 0 到 numCourses - 1。在选修某些课程之前需要一些先修课程。 先修课程按数组 prerequisites 给出，其中 prerequisites[i] = [ai, bi] ，表示如果要学习课程 ai 则 必须 先学习课程  bi 。\n例如，先修课程对 [0, 1] 表示：想要学习课程 0 ，你需要先完成课程 1 。\n请你判断是否可能完成所有课程的学习？如果可以，返回一个完成所有课程的学习顺序。你可以假设答案总是唯一的。\n\n示例：\n输入：numCourses = 2, prerequisites = [[1,0]]\n输出：[0,1]",
"这道题是课程表的进阶版，要求返回拓扑排序的结果。\n\n核心思想：拓扑排序（Kahn算法或DFS）\n\nKahn算法详细步骤：\n1. 计算每个节点的入度\n2. 将所有入度为0的节点加入队列和结果\n3. 当队列不为空时：\n   - 出队一个节点\n   - 将其所有邻居的入度减1\n   - 如果邻居入度变为0，入队并加入结果\n4. 如果结果的长度等于numCourses，返回结果；否则返回空数组\n\n时间复杂度：O(V + E)\n空间复杂度：O(V + E)")

add_problem(416, "分割等和子集 (Partition Equal Subset Sum)", "中等", "https://leetcode.cn/problems/partition-equal-subset-sum/", "动态规划",
"给你一个 只包含正整数 的 非空 数组 nums 。请你判断是否可以将这个数组分割成两个子集，使得两个子集的元素和相等。\n\n示例：\n输入：nums = [1,5,11,5]\n输出：true\n解释：数组可以分割成 [1, 5, 5] 和 [11] 。",
"这道题是经典的0-1背包问题。\n\n核心思想：判断是否存在一个子集，其和等于总和的一半。\n\n问题转化：\n设总和为sum，如果sum是奇数，直接返回false。\n否则，问题转化为：能否从数组中选出一些数，使其和等于sum/2。\n\n状态定义：\ndp[i]表示是否能从数组中选出一些数，使其和等于i\n\n状态转移：\n对于每个数num，从大到小遍历j：\n  dp[j] = dp[j] or dp[j - num]\n\n初始条件：\ndp[0] = true\n\n时间复杂度：O(n × sum)\n空间复杂度：O(sum)")

add_problem(11, "盛最多水的容器 (Container With Most Water)", "中等", "https://leetcode.cn/problems/container-with-most-water/", "双指针",
"给定一个长度为 n 的整数数组 height 。有 n 条垂线，第 i 条线的两个端点是 (i, 0) 和 (i, height[i]) 。\n找出其中的两条线，使得它们与 x 轴共同构成的容器可以容纳最多的水。\n返回容器可以储存的最大水量。\n说明：你不能倾斜容器。\n\n示例：\n输入：[1,8,6,2,5,4,8,3,7]\n输出：49",
"这道题是双指针的经典应用。\n\n核心思想：双指针从两端向中间移动，每次移动较小高度的指针。\n\n详细步骤：\n1. 初始化left = 0, right = n-1, maxArea = 0\n2. 当left < right时：\n   - 计算当前面积：area = min(height[left], height[right]) × (right - left)\n   - 更新maxArea\n   - 如果height[left] < height[right]，left++；否则right--\n\n为什么移动较小高度的指针？\n因为面积由较小高度决定，移动较小高度的指针才有可能找到更大的面积。\n\n时间复杂度：O(n)\n空间复杂度：O(1)")

add_problem(287, "寻找重复数 (Find the Duplicate Number)", "中等", "https://leetcode.cn/problems/find-the-duplicate-number/", "技巧",
"给定一个包含 n + 1 个整数的数组 nums ，其数字都在 [1, n] 范围内（包括 1 和 n），可知至少存在一个重复的整数。\n假设 nums 只有 一个重复的整数 ，返回 这个重复的数 。\n你设计的解决方案必须 不修改 数组 nums 且只用常量级 O(1) 的额外空间。\n\n示例：\n输入：nums = [1,3,4,2,2]\n输出：2",
"这道题有多种解法，但要求不修改数组且只用O(1)空间，就只能使用二分查找或Floyd判圈算法。\n\n方法一：二分查找（O(n log n)）\n对答案进行二分查找，统计数组中小于等于mid的数的个数。\n\n方法二：Floyd判圈算法（O(n)，推荐）\n将数组看作链表，nums[i]表示从i指向的下一个节点。由于存在重复数字，链表一定有环。\n使用Floyd算法找到环的入口，就是重复的数字。\n\n时间复杂度：O(n)\n空间复杂度：O(1)")

add_problem(17, "电话号码的字母组合 (Letter Combinations of a Phone Number)", "中等", "https://leetcode.cn/problems/letter-combinations-of-a-phone-number/", "回溯",
"给定一个仅包含数字 2-9 的字符串，返回所有它能表示的字母组合。答案可以按 任意顺序 返回。\n给出数字到字母的映射如下（与电话按键相同）。注意 1 不对应任何字母。\n\n示例：\n输入：digits = \"23\"\n输出：[\"ad\",\"ae\",\"af\",\"bd\",\"be\",\"bf\",\"cd\",\"ce\",\"cf\"]",
"这道题是回溯法的经典应用。\n\n核心思想：回溯遍历所有可能的组合。\n\n详细步骤：\n1. 建立数字到字母的映射表\n2. 定义递归函数backtrack(index, currentString)\n3. 如果index等于digits的长度，将currentString加入结果\n4. 获取当前数字对应的字母列表\n5. 遍历每个字母，递归调用backtrack\n\n时间复杂度：O(3^m × 4^n)\n空间复杂度：O(m+n)")

print(f"总共已添加 {len(PROBLEMS)} 道题目")

if __name__ == '__main__':
    main()
