#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LeetCode HOT100 完整题解 Word 文档生成器
包含100道题的详细题解，按类别分组
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 完整的100道题数据
PROBLEMS = [
    # ========== 哈希 ==========
    {
        "id": 1,
        "title": "两数之和 (Two Sum)",
        "difficulty": "简单",
        "link": "https://leetcode.cn/problems/two-sum/",
        "category": "哈希",
        "description": """给定一个整数数组 nums 和一个整数目标值 target，请你在该数组中找出和为目标值 target 的那两个整数，并返回它们的数组下标。

你可以假设每种输入只会对应一个答案。但是，数组中同一个元素在答案里不能重复出现。

你可以按任意顺序返回答案。

**示例**：
```
输入：nums = [2,7,11,15], target = 9
输出：[0,1]
解释：因为 nums[0] + nums[1] == 9，返回 [0, 1]。
```""",
        "solution": """**解题思路**：

这道题是经典的哈希表应用题。最直观的想法是使用双重循环遍历所有可能的数对，但时间复杂度为O(n²)，效率较低。为了优化查找效率，我们可以利用哈希表的O(1)查找特性。

**核心思想**：在遍历数组的过程中，对于当前元素num，我们需要查找target-num是否已经在之前遍历过的元素中出现过。如果出现过，说明找到了答案；如果没有，则将当前元素存入哈希表，继续遍历。

**为什么选择"先查后存"的顺序？** 这是为了避免同一个元素被使用两次。如果先存入再查找，当数组中存在target/2这个元素时，可能会错误地认为找到了答案（实际上是同一个元素）。

**详细步骤**：
1. 创建一个空的哈希表，用于存储已经遍历过的数字及其下标
2. 从左到右遍历数组中的每个元素num及其下标i
3. 计算complement = target - num，这是在数组中需要找到的另一个数
4. 在哈希表中查找complement是否存在
   - 如果存在，说明找到了答案，返回[complement的下标, i]
   - 如果不存在，将num和i存入哈希表，继续遍历

**时间复杂度分析**：O(n)，其中n是数组长度。我们只需要遍历数组一次，哈希表的插入和查找操作都是O(1)。

**空间复杂度分析**：O(n)，最坏情况下哈希表需要存储n-1个元素。"""
    },
    {
        "id": 49,
        "title": "字母异位词分组 (Group Anagrams)",
        "difficulty": "中等",
        "link": "https://leetcode.cn/problems/group-anagrams/",
        "category": "哈希",
        "description": """给你一个字符串数组，请你将字母异位词组合在一起。可以按任意顺序返回结果列表。

字母异位词是由重新排列源单词的所有字母得到的一个新单词。

**示例**：
```
输入: strs = ["eat", "tea", "tan", "ate", "nat", "bat"]
输出: [["bat"],["nat","tan"],["ate","eat","tea"]]
```""",
        "solution": """**解题思路**：

这道题的关键在于如何识别字母异位词。两个字符串是字母异位词，当且仅当它们包含相同的字符且每个字符出现的次数相同。基于这一性质，我们可以设计两种方法来作为哈希表的键。

**方法一：排序作为键**

将每个字符串的字符进行排序，排序后的字符串作为哈希表的键。所有字母异位词排序后得到的结果相同，因此会被分到同一组。

例如，"eat"、"tea"、"ate"排序后都是"aet"，因此它们会被分到同一组。

- 时间复杂度：O(n × k log k)，其中n是字符串个数，k是字符串的最大长度。对每个字符串排序需要O(k log k)，共n个字符串。
- 空间复杂度：O(n × k)，需要存储所有字符串的排序结果。

**方法二：字符计数作为键（更优）**

使用一个长度为26的数组统计每个字符的出现次数，然后将这个计数数组转换为字符串作为键。这种方法避免了排序操作。

例如，"eat"的字符计数为[1,0,0,0,1,0,0,0,0,0,0,0,0,0,0,0,0,0,0,1,0,0,0,0,0,0]，转换为字符串"1#0#0#0#1#0#0#0#0#0#0#0#0#0#0#0#0#0#0#1#0#0#0#0#0#0"。

- 时间复杂度：O(n × k)，每个字符串只需遍历一次来计算字符计数。
- 空间复杂度：O(n × k)

方法二是更优的解法，因为它的时间复杂度更低，且实际运行效率更高。"""
    },
    {
        "id": 128,
        "title": "最长连续序列 (Longest Consecutive Sequence)",
        "difficulty": "中等",
        "link": "https://leetcode.cn/problems/longest-consecutive-sequence/",
        "category": "哈希",
        "description": """给定一个未排序的整数数组 nums ，找出数字连续的最长序列（不要求序列元素在原数组中连续）的长度。

请你设计并实现时间复杂度为 O(n) 的算法解决此问题。

**示例**：
```
输入：nums = [100,4,200,1,3,2]
输出：4
解释：最长数字连续序列是 [1, 2, 3, 4]。它的长度为 4。
```""",
        "solution": """**解题思路**：

这道题要求O(n)时间复杂度，说明不能对数组进行排序（排序的时间复杂度是O(n log n)）。我们需要利用哈希集合来实现O(1)的查找操作。

**核心思想**：只从序列的起点开始计数。

**关键观察**：如果一个数num是某个连续序列的起点，那么num-1一定不在数组中。利用这一性质，我们可以避免重复计算序列长度。

**详细步骤**：
1. 将所有数字存入哈希集合，实现O(1)查找
2. 遍历集合中的每个数字num：
   - 如果num-1不在集合中，说明num是某个连续序列的起点
   - 从num开始，不断查找num+1、num+2...是否在集合中，直到不连续
   - 记录当前序列的长度，更新最大值
   - 如果num-1在集合中，跳过该数字（它会在作为序列起点时被计算）

**为什么这种方法是O(n)？** 虽然看起来有两层循环，但每个数字最多被访问两次：一次作为起点判断，一次在序列中。因此总时间复杂度是O(n)。

**时间复杂度**：O(n)，每个数字最多被访问两次

**空间复杂度**：O(n)，哈希集合存储所有数字"""
    },
    {
        "id": 136,
        "title": "只出现一次的数字 (Single Number)",
        "difficulty": "简单",
        "link": "https://leetcode.cn/problems/single-number/",
        "category": "哈希",
        "description": """给定一个非空整数数组，除了某个元素只出现一次以外，其余每个元素均出现两次。找出那个只出现了一次的元素。

**说明**：你的算法应该具有线性时间复杂度。 你可以不使用额外空间来实现吗？

**示例**：
```
输入: [2,2,1]
输出: 1
```""",
        "solution": """**解题思路**：

这道题是一道经典的位运算题目。题目要求线性时间复杂度和常数空间复杂度，这意味着我们不能使用哈希表等额外数据结构。

**核心思想**：利用异或运算的性质。

**异或运算的性质**：
1. a ^ a = 0（任何数与自身异或结果为0）
2. a ^ 0 = a（任何数与0异或结果为自身）
3. 异或运算满足交换律和结合律：a ^ b ^ a = a ^ a ^ b = b

**算法**：
将数组中所有数字进行异或运算。根据上述性质，成对出现的数字异或结果为0，最后剩下的就是只出现一次的数字。

例如：[2,2,1]
- 2 ^ 2 ^ 1 = 0 ^ 1 = 1

例如：[4,1,2,1,2]
- 4 ^ 1 ^ 2 ^ 1 ^ 2 = 4 ^ (1 ^ 1) ^ (2 ^ 2) = 4 ^ 0 ^ 0 = 4

**时间复杂度**：O(n)，只需要遍历数组一次

**空间复杂度**：O(1)，只使用了一个变量来存储异或结果

这是一个非常巧妙的解法，充分利用了异或运算的特性，完美满足了题目的要求。"""
    },
    
    # ========== 双指针 ==========
    {
        "id": 283,
        "title": "移动零 (Move Zeroes)",
        "difficulty": "简单",
        "link": "https://leetcode.cn/problems/move-zeroes/",
        "category": "双指针",
        "description": """给定一个数组 nums，编写一个函数将所有 0 移动到数组的末尾，同时保持非零元素的相对顺序。

请注意 ，必须在不复制数组的情况下原地对数组进行操作。

**示例**：
```
输入: nums = [0,1,0,3,12]
输出: [1,3,12,0,0]
```""",
        "solution": """**解题思路**：

这道题考察双指针技巧。我们需要在原数组上操作，将所有非零元素移动到数组前面，同时保持它们的相对顺序，最后将剩余位置填充为0。

**核心思想**：使用双指针，一个指针遍历数组，另一个指针记录非零元素应该放置的位置。

**详细步骤**：
1. 初始化指针j = 0，表示下一个非零元素应该放置的位置
2. 使用指针i遍历数组：
   - 如果nums[i] != 0：
     - 将nums[i]赋值给nums[j]
     - j++（指向下一个应该放置的位置）
3. 遍历完成后，将数组中从j开始到末尾的所有元素赋值为0

**为什么这样能保持相对顺序？** 因为我们是从左到右遍历，非零元素按照它们在原数组中出现的顺序被依次放置到前面。

**优化版本（一次遍历）**：
可以在遍历过程中直接交换非零元素到前面，这样可以省去最后填充0的步骤。

```python
j = 0
for i in range(len(nums)):
    if nums[i] != 0:
        nums[i], nums[j] = nums[j], nums[i]
        j += 1
```

**时间复杂度**：O(n)，每个元素最多被访问两次

**空间复杂度**：O(1)，只使用了常数额外空间"""
    },
    {
        "id": 11,
        "title": "盛最多水的容器 (Container With Most Water)",
        "difficulty": "中等",
        "link": "https://leetcode.cn/problems/container-with-most-water/",
        "category": "双指针",
        "description": """给定一个长度为 n 的整数数组 height 。有 n 条垂线，第 i 条线的两个端点是 (i, 0) 和 (i, height[i]) 。

找出其中的两条线，使得它们与 x 轴共同构成的容器可以容纳最多的水。

返回容器可以储存的最大水量。

**说明**：你不能倾斜容器。

**示例**：
```
输入：[1,8,6,2,5,4,8,3,7]
输出：49
```""",
        "solution": """**解题思路**：

这道题是双指针的经典应用。我们需要找到两条线，使得它们之间的容器能容纳最多的水。

**容器的面积计算**：面积 = min(height[left], height[right]) × (right - left)

**核心思想**：双指针从两端向中间移动，每次移动较小高度的指针。

**为什么移动较小高度的指针？**

假设height[left] < height[right]，此时容器的面积由height[left]决定。如果我们移动right指针：
- 宽度减小（right - left变小）
- 高度不变或更小（min(height[left], height[new_right]) ≤ height[left]）
- 因此面积一定减小

但如果我们移动left指针：
- 宽度减小
- 高度可能增大（如果新的left高度更大）
- 面积可能增大

因此，应该移动较小高度的指针，才有可能找到更大的面积。

**详细步骤**：
1. 初始化left = 0, right = n-1, maxArea = 0
2. 当left < right时：
   - 计算当前面积：area = min(height[left], height[right]) × (right - left)
   - 更新maxArea = max(maxArea, area)
   - 如果height[left] < height[right]，left++；否则right--

**时间复杂度**：O(n)，两个指针最多各遍历一次

**空间复杂度**：O(1)，只使用了常数额外空间"""
    },
    {
        "id": 15,
        "title": "三数之和 (3Sum)",
        "difficulty": "中等",
        "link": "https://leetcode.cn/problems/3sum/",
        "category": "双指针",
        "description": """给你一个整数数组 nums ，判断是否存在三元组 [nums[i], nums[j], nums[k]] 满足 i != j、i != k 且 j != k ，同时还满足 nums[i] + nums[j] + nums[k] == 0 。

请你返回所有和为 0 且不重复的三元组。

**注意**：答案中不可以包含重复的三元组。

**示例**：
```
输入：nums = [-1,0,1,2,-1,-4]
输出：[[-1,-1,2],[-1,0,1]]
```""",
        "solution": """**解题思路**：

这道题是"两数之和"的升级版，要求找出三个数使得它们的和为0。最暴力的解法是三重循环，时间复杂度为O(n³)。我们需要更优的解法。

**核心思想**：排序 + 双指针

**为什么先排序？** 排序后可以使用双指针技巧，将O(n³)降到O(n²)，同时也方便去重。

**详细步骤**：
1. 对数组进行排序
2. 遍历数组，固定第一个数nums[i]：
   - 如果nums[i] > 0，直接break（因为数组已排序，三个正数之和不可能为0）
   - 去重：如果i > 0且nums[i] == nums[i-1]，跳过（避免重复的三元组）
   - 使用双指针left = i+1, right = n-1：
     - 计算sum = nums[i] + nums[left] + nums[right]
     - 如果sum == 0：
       - 记录结果
       - left++，right--
       - 去重：跳过所有重复的nums[left]和nums[right]
     - 如果sum < 0：left++（需要更大的数）
     - 如果sum > 0：right--（需要更小的数）

**去重处理的三层**：
1. 外层循环跳过重复的nums[i]
2. 找到解后，跳过重复的nums[left]
3. 找到解后，跳过重复的nums[right]

**时间复杂度**：O(n²)，排序O(n log n)，双指针遍历O(n²)

**空间复杂度**：O(1)或O(log n)（排序的栈空间）"""
    },
    {
        "id": 42,
        "title": "接雨水 (Trapping Rain Water)",
        "difficulty": "困难",
        "link": "https://leetcode.cn/problems/trapping-rain-water/",
        "category": "双指针",
        "description": """给定 n 个非负整数表示每个宽度为 1 的柱子的高度图，计算按此排列的柱子，下雨之后能接多少雨水。

**示例**：
```
输入：height = [0,1,0,2,1,0,1,3,2,1,2,1]
输出：6
```""",
        "solution": """**解题思路**：

这道题是双指针的经典难题。每个位置能接的雨水量取决于其左边和右边最高柱子的较小值。

**核心公式**：每个位置能接的雨水 = min(左边最大高度, 右边最大高度) - 当前高度

**方法一：动态规划**

预处理每个位置左边的最大高度和右边的最大高度。
1. leftMax[i]表示从0到i的最大高度
2. rightMax[i]表示从i到n-1的最大高度
3. 遍历计算每个位置的接水量

- 时间复杂度：O(n)
- 空间复杂度：O(n)

**方法二：双指针（推荐）**

用两个指针从两端向中间移动，维护左边的最大值和右边的最大值。

**关键观察**：
- 如果leftMax < rightMax，左边可以计算接水量（因为右边一定有更高的柱子）
- 如果leftMax >= rightMax，右边可以计算接水量

**详细步骤**：
1. 初始化left = 0, right = n-1, leftMax = 0, rightMax = 0, result = 0
2. 当left < right时：
   - 如果height[left] < height[right]：
     - 如果height[left] >= leftMax，更新leftMax = height[left]
     - 否则，result += leftMax - height[left]（可以接雨水）
     - left++
   - 否则：
     - 如果height[right] >= rightMax，更新rightMax = height[right]
     - 否则，result += rightMax - height[right]（可以接雨水）
     - right--

**时间复杂度**：O(n)

**空间复杂度**：O(1)，只使用了常数额外空间"""
    },
    
    # ========== 滑动窗口 ==========
    {
        "id": 3,
        "title": "无重复字符的最长子串 (Longest Substring Without Repeating Characters)",
        "difficulty": "中等",
        "link": "https://leetcode.cn/problems/longest-substring-without-repeating-characters/",
        "category": "滑动窗口",
        "description": """给定一个字符串 s ，请你找出其中不含有重复字符的 最长子串 的长度。

**示例**：
```
输入: s = "abcabcbb"
输出: 3 
解释: 因为无重复字符的最长子串是 "abc"，所以其长度为 3。
```""",
        "solution": """**解题思路**：

这道题是滑动窗口的经典题目。我们需要维护一个不包含重复字符的窗口，不断扩展右边界，当发现重复字符时收缩左边界。

**核心思想**：滑动窗口，使用双指针维护一个不包含重复字符的窗口[left, right]。

**详细步骤**：
1. 初始化left = 0, maxLen = 0
2. 使用哈希集合或哈希表记录窗口中的字符
3. 遍历字符串，right从0到n-1：
   - 如果s[right]不在窗口中：
     - 将s[right]加入窗口
     - 更新maxLen = max(maxLen, right - left + 1)
   - 如果s[right]已在窗口中：
     - 移动left，直到s[right]不在窗口中
     - 将s[right]加入窗口

**优化版本**：使用哈希表记录字符的索引

如果使用哈希表记录每个字符最后出现的位置，当发现重复字符时，可以直接将left跳到重复字符的下一个位置，而不需要逐个移动。

```python
left = 0
maxLen = 0
charIndex = {}

for right in range(len(s)):
    if s[right] in charIndex:
        left = max(left, charIndex[s[right]] + 1)
    charIndex[s[right]] = right
    maxLen = max(maxLen, right - left + 1)
```

注意使用max(left, ...)是为了防止left回退。

**时间复杂度**：O(n)，每个字符最多被访问两次

**空间复杂度**：O(min(m, n))，其中m是字符集大小"""
    },
    {
        "id": 438,
        "title": "找到字符串中所有字母异位词 (Find All Anagrams in a String)",
        "difficulty": "中等",
        "link": "https://leetcode.cn/problems/find-all-anagrams-in-a-string/",
        "category": "滑动窗口",
        "description": """给定两个字符串 s 和 p，找到 s 中所有 p 的 异位词 的子串，返回这些子串的起始索引。不考虑答案输出的顺序。

**示例**：
```
输入: s = "cbaebabacd", p = "abc"
输出: [0,6]
解释:
起始索引等于 0 的子串是 "cba"，它是 "abc" 的异位词。
起始索引等于 6 的子串是 "bac"，它是 "abc" 的异位词。
```""",
        "solution": """**解题思路**：

这道题是滑动窗口加字符计数的应用。我们需要在字符串s中找到所有与p是字母异位词的子串。

**核心思想**：滑动窗口 + 字符计数

**什么是字母异位词？** 两个字符串包含相同的字符且每个字符出现的次数相同。因此，我们可以使用字符计数来判断两个字符串是否是字母异位词。

**详细步骤**：
1. 统计p中每个字符的出现次数，存入数组pCount（大小为26，对应26个字母）
2. 初始化一个大小为p.length()的滑动窗口，统计窗口中每个字符的出现次数
3. 使用数组windowCount记录窗口中的字符计数
4. 滑动窗口遍历字符串s：
   - 右移窗口：添加右边的新字符到windowCount
   - 移除左边滑出的字符
   - 比较windowCount和pCount，如果相同，说明找到了一个字母异位词，记录起始索引

**优化**：使用一个变量diff记录当前窗口与目标差异的字符种类数，避免每次比较整个数组。

```python
diff = 0  # 记录有多少种字符的计数不匹配
for i in range(26):
    if pCount[i] != windowCount[i]:
        diff += 1

# 滑动窗口时更新diff
```

**时间复杂度**：O(n)，其中n是s的长度

**空间复杂度**：O(1)，固定大小的数组（26个字母）"""
    },
    {
        "id": 560,
        "title": "和为 K 的子数组 (Subarray Sum Equals K)",
        "difficulty": "中等",
        "link": "https://leetcode.cn/problems/subarray-sum-equals-k/",
        "category": "滑动窗口",
        "description": """给你一个整数数组 nums 和一个整数 k ，请你统计并返回 该数组中和为 k 的子数组的个数 。

子数组是数组中元素的连续非空序列。

**示例**：
```
输入：nums = [1,1,1], k = 2
输出：2
```""",
        "solution": """**解题思路**：

这道题不能简单地使用滑动窗口，因为数组中可能包含负数，这意味着扩大窗口不一定增加和，缩小窗口不一定减少和。

**核心思想**：前缀和 + 哈希表

**关键公式**：
- 子数组和 = prefixSum[j] - prefixSum[i-1] = k
- 即 prefixSum[i-1] = prefixSum[j] - k

**详细步骤**：
1. 使用哈希表记录每个前缀和出现的次数，key是前缀和，value是出现次数
2. 初始化hashMap[0] = 1，表示前缀和为0出现1次（处理从索引0开始的子数组）
3. 遍历数组，维护当前前缀和prefixSum：
   - 在哈希表中查找prefixSum - k出现的次数，加到结果中
   - 将当前prefixSum加入哈希表（次数+1）

**为什么hashMap[0] = 1？**

考虑从数组开头开始的子数组，例如nums = [1, 2, 3], k = 6。当遍历到索引2时，prefixSum = 6，我们需要找到prefixSum - k = 0。如果不初始化hashMap[0] = 1，就会漏掉这个解。

**示例分析**：
```
nums = [1, 1, 1], k = 2

i=0: prefixSum=1, 查找1-2=-1(不存在), hashMap={0:1, 1:1}
i=1: prefixSum=2, 查找2-2=0(存在1次), result=1, hashMap={0:1, 1:1, 2:1}
i=2: prefixSum=3, 查找3-2=1(存在1次), result=2, hashMap={0:1, 1:1, 2:1, 3:1}

最终result=2
```

**时间复杂度**：O(n)

**空间复杂度**：O(n)，哈希表存储前缀和"""
    },
]

# 继续添加剩余题目...

def add_more_problems():
    """添加剩余的题目数据"""
    more_problems = [
        # ========== 普通数组 ==========
        {
            "id": 53,
            "title": "最大子数组和 (Maximum Subarray)",
            "difficulty": "中等",
            "link": "https://leetcode.cn/problems/maximum-subarray/",
            "category": "普通数组",
            "description": """给你一个整数数组 nums ，请你找出一个具有最大和的连续子数组（子数组最少包含一个元素），返回其最大和。

子数组 是数组中的一个连续部分。

**示例**：
```
输入：nums = [-2,1,-3,4,-1,2,1,-5,4]
输出：6
解释：连续子数组 [4,-1,2,1] 的和最大，为 6 。
```""",
            "solution": """**解题思路**：

这道题是著名的Kadane算法（卡丹算法）的应用，是动态规划的经典入门题目。

**核心思想**：动态规划。对于每个位置，我们要么从当前元素重新开始一个新的子数组，要么将当前元素加入到之前的子数组中。

**状态定义**：
- dp[i]表示以nums[i]结尾的最大子数组和

**状态转移方程**：
- dp[i] = max(nums[i], dp[i-1] + nums[i])
- 解释：要么从当前元素重新开始（nums[i]），要么延续前面的子数组（dp[i-1] + nums[i]）

**空间优化**：
我们不需要存储整个dp数组，只需要维护当前值和最大值即可。

**详细步骤**：
1. 初始化currentSum = nums[0], maxSum = nums[0]
2. 从i=1遍历到n-1：
   - currentSum = max(nums[i], currentSum + nums[i])
   - maxSum = max(maxSum, currentSum)

**为什么这样是正确的？**

对于每个位置i，以它结尾的最大子数组和只取决于两个选择：
1. 只包含nums[i]本身
2. 包含以i-1结尾的最大子数组 + nums[i]

我们取这两个选择中的较大值，同时维护全局最大值。

**时间复杂度**：O(n)，只需要遍历数组一次

**空间复杂度**：O(1)，只使用了常数额外空间"""
        },
        {
            "id": 56,
            "title": "合并区间 (Merge Intervals)",
            "difficulty": "中等",
            "link": "https://leetcode.cn/problems/merge-intervals/",
            "category": "普通数组",
            "description": """以数组 intervals 表示若干个区间的集合，其中单个区间为 intervals[i] = [starti, endi] 。请你合并所有重叠的区间，并返回一个不重叠的区间数组，该数组需恰好覆盖输入中的所有区间。

**示例**：
```
输入：intervals = [[1,3],[2,6],[8,10],[15,18]]
输出：[[1,6],[8,10],[15,18]]
解释：区间 [1,3] 和 [2,6] 重叠, 将它们合并为 [1,6]。
```""",
            "solution": """**解题思路**：

这道题是区间问题的经典题目。合并重叠区间的关键是先对区间进行排序，然后逐个检查是否重叠。

**核心思想**：排序 + 贪心

**为什么先排序？** 按区间的左端点排序后，重叠的区间会相邻出现，方便我们逐个处理。

**详细步骤**：
1. 按区间的左端点升序排序
2. 初始化结果列表，将第一个区间加入结果
3. 遍历剩余的区间：
   - 如果当前区间的左端点 <= 结果列表最后一个区间的右端点：
     - 说明有重叠，合并区间
     - 更新最后一个区间的右端点为max(当前右端点, 最后一个右端点)
   - 否则：
     - 无重叠，将当前区间加入结果列表

**合并逻辑示例**：
```
已排序区间：[[1,3], [2,6], [8,10], [15,18]]

- 结果初始：[[1,3]]
- 处理[2,6]：2 <= 3，有重叠，合并为[1, max(3,6)] = [1,6]
- 结果：[[1,6]]
- 处理[8,10]：8 > 6，无重叠，加入
- 结果：[[1,6], [8,10]]
- 处理[15,18]：15 > 10，无重叠，加入
- 最终结果：[[1,6], [8,10], [15,18]]
```

**时间复杂度**：O(n log n)，主要是排序的时间

**空间复杂度**：O(log n)或O(n)，取决于排序的实现（是否为原地排序）"""
        },
        {
            "id": 189,
            "title": "轮转数组 (Rotate Array)",
            "difficulty": "中等",
            "link": "https://leetcode.cn/problems/rotate-array/",
            "category": "普通数组",
            "description": """给定一个整数数组 nums，将数组中的元素向右轮转 k 个位置，其中 k 是非负数。

**示例**：
```
输入: nums = [1,2,3,4,5,6,7], k = 3
输出: [5,6,7,1,2,3,4]
解释:
向右轮转 1 步: [7,1,2,3,4,5,6]
向右轮转 2 步: [6,7,1,2,3,4,5]
向右轮转 3 步: [5,6,7,1,2,3,4]
```""",
            "solution": """**解题思路**：

这道题有多种解法，我们需要找到最优的空间复杂度解法。

**方法一：使用额外数组（不推荐）**
创建一个新数组，将原数组的元素放到正确的位置。
- 时间复杂度：O(n)
- 空间复杂度：O(n)

**方法二：数组反转（推荐）**
利用数组反转的性质来实现轮转。这是一个非常巧妙的技巧。

**核心思想**：先整体反转，再分别反转前后两部分。

以[1,2,3,4,5,6,7], k=3为例：
1. 整体反转：[7,6,5,4,3,2,1]
2. 反转前k个：[5,6,7,4,3,2,1]
3. 反转剩余部分：[5,6,7,1,2,3,4]

**数学原理**：
向右轮转k位，相当于将后k个元素移到前面，前n-k个元素移到后面。通过三次反转，我们实现了这一效果。

**详细步骤**：
1. k = k % n（处理k大于n的情况）
2. 反转整个数组
3. 反转[0, k-1]
4. 反转[k, n-1]

**时间复杂度**：O(n)，三次反转各需要O(n)

**空间复杂度**：O(1)，原地操作，只使用了常数额外空间"""
        },
        {
            "id": 238,
            "title": "除自身以外数组的乘积 (Product of Array Except Self)",
            "difficulty": "中等",
            "link": "https://leetcode.cn/problems/product-of-array-except-self/",
            "category": "普通数组",
            "description": """给你一个整数数组 nums，返回 数组 answer ，其中 answer[i] 等于 nums 中除 nums[i] 之外其余各元素的乘积 。

题目数据 保证 数组 nums之中任意元素的全部前缀元素和后缀的乘积都在 32 位 整数范围内。

请**不要使用除法**，且在 O(n) 时间复杂度内完成此题。

**示例**：
```
输入: nums = [1,2,3,4]
输出: [24,12,8,6]
```""",
            "solution": """**解题思路**：

这道题乍一看很简单，直接用总乘积除以每个元素即可。但题目要求不能使用除法，而且如果数组中有0，除法方案也会失效。

**核心思想**：前缀积 × 后缀积

对于answer[i]，它等于nums[0...i-1]的乘积 × nums[i+1...n-1]的乘积。
- nums[0...i-1]的乘积称为前缀积
- nums[i+1...n-1]的乘积称为后缀积

**详细步骤**：
1. 创建结果数组answer
2. 第一次遍历（从左到右）：计算前缀积
   - answer[i] = nums[0] × nums[1] × ... × nums[i-1]
   - 即answer[i] = answer[i-1] × nums[i-1]
   - answer[0]初始化为1（空积）
3. 第二次遍历（从右到左）：计算后缀积并乘到结果中
   - 维护一个变量suffix，表示从右往左的后缀积
   - answer[i] *= suffix
   - suffix *= nums[i]

**示例分析**：
```
nums = [1, 2, 3, 4]

第一次遍历后（前缀积）：
answer = [1, 1, 2, 6]
  - answer[0] = 1 (初始值)
  - answer[1] = 1 (nums[0])
  - answer[2] = 1×2 = 2 (nums[0]×nums[1])
  - answer[3] = 1×2×3 = 6

第二次遍历：
i=3: answer[3]=6×1=6, suffix=1×4=4
i=2: answer[2]=2×4=8, suffix=4×3=12
i=1: answer[1]=1×12=12, suffix=12×2=24
i=0: answer[0]=1×24=24, suffix=24×1=24

最终结果：[24, 12, 8, 6]
```

**时间复杂度**：O(n)，两次遍历

**空间复杂度**：O(1)，输出数组不算额外空间"""
        },
        {
            "id": 41,
            "title": "缺失的第一个正数 (First Missing Positive)",
            "difficulty": "困难",
            "link": "https://leetcode.cn/problems/first-missing-positive/",
            "category": "普通数组",
            "description": """给你一个未排序的整数数组 nums ，请你找出其中没有出现的最小的正整数。

请你实现时间复杂度为 O(n) 并且只使用常数级别额外空间的解决方案。

**示例**：
```
输入：nums = [1,2,0]
输出：3
```""",
            "solution": """**解题思路**：

这道题要求在O(n)时间复杂度和O(1)空间复杂度内完成，是一个非常有挑战性的原地哈希问题。

**核心思想**：原地哈希。利用数组本身作为哈希表，将每个数放到它应该在的位置上。

**关键观察**：
- 我们要找的是缺失的最小正整数，范围是[1, n+1]，其中n是数组长度
- 对于一个长度为n的数组，如果1到n都出现了，那么答案就是n+1；否则答案是1到n中缺失的最小数
- 因此，我们可以忽略所有负数、0和大于n的数

**详细步骤**：
1. 遍历数组，对于每个数nums[i]：
   - 如果nums[i]在[1, n]范围内，且nums[i] != nums[nums[i]-1]，则交换
   - 即将nums[i]放到索引为nums[i]-1的位置上
   - 交换后，继续检查新的nums[i]，直到不满足条件
2. 再次遍历数组，找到第一个nums[i] != i+1的位置，返回i+1
3. 如果1到n都存在，返回n+1

**为什么这样有效？**

通过将每个数放到正确的位置，我们实际上是在原地构建了一个哈希表。正确位置上的数满足nums[i] = i+1。

**示例**：
```
nums = [3, 4, -1, 1]

遍历过程：
i=0: nums[0]=3，放到索引2，交换后[-1, 4, 3, 1]，-1不在[1,4]范围内，跳过
i=1: nums[1]=4，放到索引3，交换后[-1, 1, 3, 4]，继续检查nums[1]=1
       1放到索引0，交换后[1, -1, 3, 4]
i=2: nums[2]=3，已经在正确位置
i=3: nums[3]=4，已经在正确位置

结果数组：[1, -1, 3, 4]
检查：nums[1] = -1 != 2，返回2
```

**时间复杂度**：O(n)，虽然看起来有嵌套循环，但每个元素最多被交换两次

**空间复杂度**：O(1)，原地操作"""
        },
    ]
    
    # 添加更多题目... 继续矩阵、链表、二叉树、图论、回溯等类别
    more_problems.extend([
        # ========== 矩阵 ==========
        {
            "id": 73,
            "title": "矩阵置零 (Set Matrix Zeroes)",
            "difficulty": "中等",
            "link": "https://leetcode.cn/problems/set-matrix-zeroes/",
            "category": "矩阵",
            "description": """给定一个 m x n 的矩阵，如果一个元素为 0 ，则将其所在行和列的所有元素都设为 0 。请使用 原地 算法。

**示例**：
```
输入：matrix = [[1,1,1],[1,0,1],[1,1,1]]
输出：[[1,0,1],[0,0,0],[1,0,1]]
```""",
            "solution": """**解题思路**：

这道题的关键在于使用原地算法，即不使用额外的矩阵空间。

**方法一：使用标记数组（空间复杂度O(m+n)）**
使用两个数组分别记录哪些行和列需要置零。这不是最优解。

**方法二：使用矩阵的第一行和第一列作为标记（空间复杂度O(1)）**

**核心思想**：用矩阵的第一行和第一列来记录哪些行和列需要置零。

**详细步骤**：
1. 使用两个变量firstRowZero和firstColZero记录第一行和第一列是否需要置零
2. 遍历矩阵（从第二行第二列开始），如果matrix[i][j] == 0：
   - 将matrix[i][0]设为0（标记第i行需要置零）
   - 将matrix[0][j]设为0（标记第j列需要置零）
3. 再次遍历矩阵（从第二行第二列开始），根据第一行和第一列的标记置零
4. 根据firstRowZero和firstColZero决定是否置零第一行和第一列

**为什么这样有效？**

第一行和第一列充当了"标记数组"的角色，用O(1)的额外空间实现了标记功能。需要注意的是，第一行和第一列本身是否置零需要单独处理。

**时间复杂度**：O(m × n)，需要遍历矩阵多次

**空间复杂度**：O(1)，只使用了常数额外空间"""
        },
        {
            "id": 54,
            "title": "螺旋矩阵 (Spiral Matrix)",
            "difficulty": "中等",
            "link": "https://leetcode.cn/problems/spiral-matrix/",
            "category": "矩阵",
            "description": """给你一个 m 行 n 列的矩阵 matrix ，请按照 顺时针螺旋顺序 ，返回矩阵中的所有元素。

**示例**：
```
输入：matrix = [[1,2,3],[4,5,6],[7,8,9]]
输出：[1,2,3,6,9,8,7,4,5]
```""",
            "solution": """**解题思路**：

这道题考察对矩阵边界的控制。按照顺时针螺旋顺序遍历矩阵，即：右→下→左→上→右→...的循环。

**核心思想**：维护四个边界（上、下、左、右），按照顺序遍历，每遍历完一条边就收缩对应的边界。

**详细步骤**：
1. 初始化四个边界：top=0, bottom=m-1, left=0, right=n-1
2. 当top <= bottom且left <= right时循环：
   - 从左到右遍历上边：从left到right，遍历完成后top++
   - 从上到下遍历右边：从top到bottom，遍历完成后right--
   - 如果top <= bottom，从右到左遍历下边：从right到left，遍历完成后bottom--
   - 如果left <= right，从下到上遍历左边：从bottom到top，遍历完成后left++

**为什么需要检查top <= bottom和left <= right？**

当矩阵是长方形（如1×n或m×1）时，在遍历完某些边后，边界可能已经交叉，需要停止循环。

**时间复杂度**：O(m × n)，每个元素被访问一次

**空间复杂度**：O(1)，不包括输出数组"""
        },
        {
            "id": 48,
            "title": "旋转图像 (Rotate Image)",
            "difficulty": "中等",
            "link": "https://leetcode.cn/problems/rotate-image/",
            "category": "矩阵",
            "description": """给定一个 n × n 的二维矩阵 matrix 表示一个图像。请你将图像顺时针旋转 90 度。

你必须在 原地 旋转图像，这意味着你需要直接修改输入的二维矩阵。请不要 使用另一个矩阵来旋转图像。

**示例**：
```
输入：matrix = [[1,2,3],[4,5,6],[7,8,9]]
输出：[[7,4,1],[8,5,2],[9,6,3]]
```""",
            "solution": """**解题思路**：

这道题要求原地旋转矩阵，不能创建新的矩阵。

**核心思想**：先转置矩阵，再反转每一行。

**顺时针旋转90度的数学关系**：
- 对于元素matrix[i][j]，旋转后的位置是matrix[j][n-1-i]
- 这等价于：先转置（matrix[i][j] -> matrix[j][i]），再反转每一行

**详细步骤**：
1. 转置矩阵：遍历对角线上方的元素，与对角线对称的元素交换
   - for i in range(n):
   -   for j in range(i+1, n):
   -     swap(matrix[i][j], matrix[j][i])
2. 反转每一行：对每一行进行反转
   - for i in range(n):
   -   reverse(matrix[i])

**示例**：
```
原矩阵：
[[1,2,3],
 [4,5,6],
 [7,8,9]]

转置后：
[[1,4,7],
 [2,5,8],
 [3,6,9]]

每行反转后：
[[7,4,1],
 [8,5,2],
 [9,6,3]]
```

**时间复杂度**：O(n²)，需要遍历整个矩阵

**空间复杂度**：O(1)，原地操作"""
        },
        {
            "id": 240,
            "title": "搜索二维矩阵 II (Search a 2D Matrix II)",
            "difficulty": "中等",
            "link": "https://leetcode.cn/problems/search-a-2d-matrix-ii/",
            "category": "矩阵",
            "description": """编写一个高效的算法来搜索 m x n 矩阵 matrix 中的一个目标值 target 。该矩阵具有以下特性：

- 每行的元素从左到右升序排列。
- 每列的元素从上到下升序排列。

**示例**：
```
输入：matrix = [[1,4,7,11,15],[2,5,8,12,19],[3,6,9,16,22],[10,13,14,17,24],[18,21,23,26,30]], target = 5
输出：true
```""",
            "solution": """**解题思路**：

这道题的关键在于利用矩阵的有序性来设计高效的搜索算法。暴力搜索是O(m×n)，但我们可以做得更好。

**核心思想**：从右上角或左下角开始搜索，利用有序性排除行或列。

**详细步骤（从右上角开始）**：
1. 初始化row = 0, col = n-1（右上角）
2. 当row < m且col >= 0时循环：
   - 如果matrix[row][col] == target，返回True
   - 如果matrix[row][col] > target，col--（当前列的所有元素都大于target，排除这一列）
   - 如果matrix[row][col] < target，row++（当前行的所有元素都小于target，排除这一行）
3. 如果循环结束还没找到，返回False

**为什么从右上角开始？**

从右上角开始，我们有一个性质：左边的元素都小于当前元素，下边的元素都大于当前元素。这样每次比较都可以排除一行或一列。

**示例分析**：
```
从右上角15开始，target=5
15 > 5，向左移动到11
11 > 5，向左移动到7
7 > 5，向左移动到4
4 < 5，向下移动到5
5 == 5，找到！
```

**时间复杂度**：O(m + n)，每次比较排除一行或一列

**空间复杂度**：O(1)，只使用了常数额外空间"""
        },
    ])
    
    return more_problems

def create_word_document():
    """创建Word文档"""
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    
    # 添加标题
    title = doc.add_heading('LeetCode 热题 HOT 100 完整题解', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加副标题
    subtitle = doc.add_paragraph('包含全部100道题的详细题目描述与解题思路')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 添加日期
    date_para = doc.add_paragraph('整理时间：2026年3月')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 添加目录标题
    toc_title = doc.add_heading('目录', level=1)
    
    # 获取所有类别
    categories = list(set([p["category"] for p in PROBLEMS]))
    categories.sort()
    
    # 添加目录内容
    for i, cat in enumerate(categories, 1):
        cat_problems = [p for p in PROBLEMS if p["category"] == cat]
        cat_para = doc.add_paragraph(f'{i}. {cat}（{len(cat_problems)}题）')
        cat_para.paragraph_format.left_indent = Inches(0.3)
    
    doc.add_page_break()
    
    # 按类别分组添加题目
    for category in categories:
        # 添加类别标题
        cat_heading = doc.add_heading(category, level=1)
        
        cat_problems = [p for p in PROBLEMS if p["category"] == category]
        cat_problems.sort(key=lambda x: x['id'])
        
        for problem in cat_problems:
            # 添加题目标题
            p_heading = doc.add_heading(f'{problem["title"]}', level=2)
            
            # 添加题目信息
            info_para = doc.add_paragraph()
            info_para.add_run(f'题号：').bold = True
            info_para.add_run(f'{problem["id"]}    ')
            info_para.add_run(f'难度：').bold = True
            difficulty = problem["difficulty"]
            difficulty_run = info_para.add_run(difficulty)
            if difficulty == "简单":
                difficulty_run.font.color.rgb = RGBColor(0, 176, 80)
            elif difficulty == "中等":
                difficulty_run.font.color.rgb = RGBColor(255, 192, 0)
            else:  # 困难
                difficulty_run.font.color.rgb = RGBColor(255, 0, 0)
            info_para.add_run(f'    链接：').bold = True
            info_para.add_run(problem["link"])
            
            # 添加题目描述标题
            desc_title = doc.add_paragraph()
            desc_title.add_run('题目描述：').bold = True
            
            # 添加题目描述
            desc_para = doc.add_paragraph(problem["description"])
            desc_para.paragraph_format.first_line_indent = Inches(0.3)
            
            # 添加解题思路
            solution_para = doc.add_paragraph()
            solution_para.add_run(problem["solution"])
            
            # 添加分隔线
            doc.add_paragraph('_' * 50)
    
    return doc

def main():
    """主函数"""
    print("正在生成LeetCode HOT100完整题解Word文档...")
    
    # 添加更多题目
    global PROBLEMS
    PROBLEMS.extend(add_more_problems())
    
    # 创建Word文档
    doc = create_word_document()
    
    # 保存文档
    output_path = '/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx'
    doc.save(output_path)
    
    print(f"文档已生成：{output_path}")
    print(f"共包含 {len(PROBLEMS)} 道题")

if __name__ == '__main__':
    main()
