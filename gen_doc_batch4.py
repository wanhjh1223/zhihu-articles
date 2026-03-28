#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def set_font(run, name='宋体', size=10.5):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_h(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    set_font(r, '黑体', sizes.get(level, 10))
    return h

def add_p(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    set_font(r, '宋体', size)
    return p

doc = Document('/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx')

# 继续添加二叉树题目
tree_problems_2 = [
    ('45', '二叉搜索树中第K小的元素', 'Kth Smallest Element in a BST', '中等', '230',
     '给定二叉搜索树根节点root和整数k，返回其中第k小的元素。',
     'BST中序遍历是有序的。递归中序遍历，计数访问的节点，第k个访问的节点即为答案。迭代：用栈模拟中序遍历，提前终止。优化：如果经常查询，可维护子树节点数，快速定位。',
     '时间：O(h+k) | 空间：O(h)'),
    ('46', '二叉树的右视图', 'Binary Tree Right Side View', '中等', '199',
     '给定二叉树根节点root，按从上到下顺序返回从右侧能看到的节点值。',
     'BFS层序遍历，每层取最后一个节点。DFS：先访问右子树再访问左子树，按深度记录第一个访问的节点。',
     '时间：O(n) | 空间：O(w)'),
    ('47', '二叉树展开为链表', 'Flatten Binary Tree to Linked List', '中等', '114',
     '给定二叉树根节点root，将其展开为单链表。展开后的单链表应该使用TreeNode.right指针指向链表下一个节点，左子树指针为null。',
     '递归。将左子树插入到根节点和右子树之间：找到左子树最右节点，将右子树接在其后，将左子树移到右边，左子树置null，递归处理右子树。',
     '时间：O(n) | 空间：O(h)'),
    ('48', '从前序与中序遍历序列构造二叉树', 'Construct Binary Tree from Preorder and Inorder Traversal', '中等', '105',
     '给定两个整数数组preorder和inorder，其中preorder是二叉树先序遍历，inorder是中序遍历，构造二叉树并返回根节点。',
     '递归。前序第一个元素是根，在中序中找到根位置，左边是左子树，右边是右子树。递归构建左右子树。用哈希表存储中序值到索引的映射，快速找到根位置。',
     '时间：O(n) | 空间：O(n)'),
    ('49', '路径总和III', 'Path Sum III', '中等', '437',
     '给定二叉树根节点root和整数targetSum，求路径和等于targetSum的路径数量。路径不需要从根开始或到叶子结束，但必须向下。',
     '前缀和+哈希表。递归遍历，用哈希表记录根到当前节点路径上的前缀和个数。当前前缀和-targetSum即为以当前节点结尾的目标路径数。回溯时恢复哈希表状态。',
     '时间：O(n) | 空间：O(h)'),
    ('50', '二叉树的最近公共祖先', 'Lowest Common Ancestor of a Binary Tree', '中等', '236',
     '给定二叉树根节点root和两个节点p、q，找到它们的最近公共祖先。',
     '递归。递归函数返回以root为根的子树是否包含p或q。如果root为p或q返回root。递归左右子树，如果左右都返回非null则root是LCA，否则返回非null的一侧。',
     '时间：O(n) | 空间：O(h)'),
    ('51', '二叉树中的最大路径和', 'Binary Tree Maximum Path Sum', '困难', '124',
     '给定二叉树根节点root，返回最大路径和。路径是树中任意节点序列，路径和是路径上节点值之和。',
     '递归（树形DP）。递归函数返回以当前节点为端点的最大路径和（只能选一条分支）。递归过程中用全局变量记录最大路径和（可包含两条分支）。当前节点为端点的最大值=max(0,左分支,右分支)+当前值，更新全局最大值。',
     '时间：O(n) | 空间：O(h)'),
]

for p in tree_problems_2:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

# 八、图论
add_h(doc, '八、图论（Graph）', 1)

graph_problems = [
    ('52', '岛屿数量', 'Number of Islands', '中等', '200',
     '给定m×n二维网格grid表示地图，1为陆地，0为水，计算岛屿数量。岛屿由水平或垂直相邻的陆地组成。',
     'DFS或BFS。遍历网格，遇到1则启动DFS/BFS，将所有相邻1标记为0（淹没岛屿），岛屿计数+1。并查集：将所有相邻1合并，统计连通分量数。',
     '时间：O(mn) | 空间：O(mn)'),
    ('53', '腐烂的橘子', 'Rotting Oranges', '简单', '994',
     '给定m×n网格grid，0为空，1为新鲜橘子，2为腐烂橘子。每分钟，腐烂橘子会使相邻新鲜橘子腐烂。返回直到没有新鲜橘子为止的最小分钟数，如果不可能返回-1。',
     'BFS多源。将所有腐烂橘子入队，同时统计新鲜橘子数。BFS每层代表一分钟，腐烂相邻新鲜橘子，新鲜橘子数减一。最后如果还有新鲜橘子返回-1，否则返回分钟数。',
     '时间：O(mn) | 空间：O(mn)'),
    ('54', '课程表', 'Course Schedule', '中等', '207',
     '给定numCourses门课和prerequisites先修课程关系，判断是否可能完成所有课程学习。',
     '拓扑排序。建图，计算每个节点入度。BFS：将入度为0的节点入队，依次出队，减少邻接节点入度，将入度变为0的节点入队。如果处理的节点数等于课程数则无环，可以完成。DFS：检测环。',
     '时间：O(V+E) | 空间：O(V+E)'),
    ('55', '实现Trie(前缀树)', 'Implement Trie (Prefix Tree)', '中等', '208',
     '实现Trie类，支持插入字符串、搜索字符串、搜索前缀。',
     'Trie树节点结构：children字典存储子节点，isEnd标记是否为单词结尾。插入：逐字符创建节点，最后标记isEnd。搜索：逐字符查找，最后返回isEnd。前缀搜索：逐字符查找，存在即返回true。',
     '时间：O(m) | 空间：O(mn)'),
]

for p in graph_problems:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

# 九、回溯
add_h(doc, '九、回溯（Backtracking）', 1)

backtrack_problems = [
    ('56', '全排列', 'Permutations', '中等', '46',
     '给定不含重复数字的数组nums，返回其所有可能的全排列。',
     '回溯。维护一个路径列表记录当前排列，用used数组标记已使用元素。遍历所有元素，如果未使用则加入路径并标记，递归，回溯时取消标记。递归终止条件：路径长度等于nums长度。',
     '时间：O(n×n!) | 空间：O(n)'),
    ('57', '子集', 'Subsets', '中等', '78',
     '给定整数数组nums，返回该数组所有可能的子集（幂集）。解集不能包含重复子集。',
     '回溯。每个元素选或不选。递归函数参数为当前索引，递归时选择当前元素加入路径，然后递归下一个索引；不选则直接递归下一个索引。或者用迭代法：遍历每个元素，将元素添加到已有所有子集后形成新子集。',
     '时间：O(n×2ⁿ) | 空间：O(n)'),
    ('58', '电话号码的字母组合', 'Letter Combinations of a Phone Number', '中等', '17',
     '给定仅包含数字2-9的字符串，返回所有它能表示的字母组合。数字到字母映射与电话按键相同。',
     '回溯。建立数字到字母的映射。递归函数参数为当前处理的数字索引，遍历当前数字对应的所有字母，加入路径，递归处理下一个数字。递归终止条件：路径长度等于数字字符串长度。',
     '时间：O(3ⁿ×4ᵐ) | 空间：O(n)'),
]

for p in backtrack_problems:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

print("已添加45-58题（共58题）")
doc.save('/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx')
print("已保存！")
