#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成LeetCode HOT100完整题解 - 分批生成
"""
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, name='宋体', size=10.5):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_h(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    set_font(r, '黑体', sizes.get(level, 10))
    return h

def add_p(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    set_font(r, '宋体', size)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, 'Courier New', 9)
    r.font.color.rgb = RGBColor(0, 80, 0)
    return p

# 完整的HOT100题目数据
problems = {
    '一、哈希': [
        ('1', '两数之和', 'Two Sum', '简单', '1', 
         '给定整数数组nums和整数目标值target，找出和为目标值的那两个整数的下标。假设每种输入只对应一个答案，同一元素不能重复使用。',
         '使用哈希表存储已遍历过的数字及其下标。对于当前数字x，查找target-x是否存在。关键点是先查后存，避免同一元素使用两次。时间复杂度O(n)，空间复杂度O(n)。',
         '时间：O(n) | 空间：O(n)'),
        ('2', '字母异位词分组', 'Group Anagrams', '中等', '49',
         '给定字符串数组，将字母异位词组合在一起。字母异位词是重新排列源单词所有字母得到的新单词。',
         '核心思想是字母异位词排序后相同。方法一：排序作为键，时间O(n×klogk)。方法二：字符计数作为键，时间O(n×k)更优。',
         '时间：O(n×k) | 空间：O(n×k)'),
        ('3', '最长连续序列', 'Longest Consecutive Sequence', '中等', '128',
         '给定未排序整数数组，找出数字连续的最长序列长度。要求时间复杂度O(n)。',
         '只从序列起点开始计数。将所有数字存入哈希集合，遍历集合，如果num-1不在集合中则说明num是起点，从num开始不断查找num+1直到不连续，记录长度更新最大值。',
         '时间：O(n) | 空间：O(n)'),
        ('4', '只出现一次的数字', 'Single Number', '简单', '136',
         '给定非空整数数组，除某个元素只出现一次外，其余每个元素均出现两次。找出那个只出现一次的元素。',
         '利用异或运算性质：a^a=0，a^0=a，异或满足交换律和结合律。将所有数字异或，成对出现的异或为0，最后剩下的就是只出现一次的数字。',
         '时间：O(n) | 空间：O(1)'),
    ],
    '二、双指针': [
        ('5', '移动零', 'Move Zeroes', '简单', '283',
         '给定数组nums，将所有0移动到数组末尾，同时保持非零元素相对顺序。必须原地操作。',
         '双指针，一个遍历，一个记录非零元素应放的位置。遍历数组，如果nums[i]!=0则赋值给nums[j]然后j++，最后将j到末尾的元素赋值为0。',
         '时间：O(n) | 空间：O(1)'),
        ('6', '盛最多水的容器', 'Container With Most Water', '中等', '11',
         '给定长度为n的整数数组height，n条垂线端点为(i,0)和(i,height[i])。找出两条线，使与x轴构成的容器容纳最多水。',
         '双指针从两端向中间移动，移动较小高度的指针。面积=min(左高,右高)×宽度。移动较大高度的指针面积一定减小，移动较小高度的指针面积可能增大。',
         '时间：O(n) | 空间：O(1)'),
        ('7', '三数之和', '3Sum', '中等', '15',
         '给定整数数组nums，判断是否存在三元组满足nums[i]+nums[j]+nums[k]==0，返回所有和为0且不重复的三元组。',
         '排序+双指针。先排序，遍历固定第一个数，如果nums[i]>0则break。去重：如果i>0且nums[i]==nums[i-1]则跳过。双指针left=i+1,right=n-1，计算sum，根据sum调整指针，找到解后去重。',
         '时间：O(n²) | 空间：O(1)'),
        ('8', '接雨水', 'Trapping Rain Water', '困难', '42',
         '给定n个非负整数表示每个宽度为1的柱子的高度图，计算下雨后能接多少雨水。',
         '每个位置能接的雨水=min(左边最大,右边最大)-当前高度。方法一DP预处理左右最大，时间O(n)空间O(n)。方法二双指针，维护左右最大值，处理较小的一侧，时间O(n)空间O(1)。',
         '时间：O(n) | 空间：O(1)'),
    ],
    '三、滑动窗口': [
        ('9', '无重复字符的最长子串', 'Longest Substring Without Repeating Characters', '中等', '3',
         '给定字符串s，找出不含有重复字符的最长子串的长度。',
         '滑动窗口维护不含重复字符的窗口。初始化left=0,maxLen=0，用哈希集合记录窗口字符。遍历字符串，如果s[right]不在窗口中加入并更新maxLen，否则移动left直到s[right]不在窗口中。优化：用哈希表记录字符索引，直接将left跳到重复字符的下一个位置。',
         '时间：O(n) | 空间：O(min(m,n))'),
        ('10', '找到字符串中所有字母异位词', 'Find All Anagrams in a String', '中等', '438',
         '给定字符串s和p，找到s中所有p的异位词的子串，返回起始索引。',
         '滑动窗口+字符计数。统计p中每个字符出现次数，维护大小为p.length()的窗口，用数组统计窗口字符数，当窗口字符计数与p相同时记录起始索引。优化：用变量记录差异字符种类数，避免每次比较整个数组。',
         '时间：O(n) | 空间：O(1)'),
        ('11', '和为K的子数组', 'Subarray Sum Equals K', '中等', '560',
         '给定整数数组nums和整数k，统计并返回数组中和为k的子数组个数。',
         '前缀和+哈希表。关键公式：prefixSum[i-1]=prefixSum[j]-k。用哈希表记录每个前缀和出现次数，遍历数组计算当前前缀和，查找prefixSum-k出现的次数加到结果中，将当前前缀和加入哈希表。初始化hashMap[0]=1。',
         '时间：O(n) | 空间：O(n)'),
        ('12', '滑动窗口最大值', 'Sliding Window Maximum', '困难', '239',
         '给定整数数组nums和大小为k的滑动窗口，返回滑动窗口中的最大值。',
         '单调队列（双端队列）。队列存储元素下标，保持队列中对应元素单调递减。遍历数组，移除窗口外的元素，从队尾移除小于当前元素的元素，将当前元素入队，当窗口形成后记录队首元素为最大值。',
         '时间：O(n) | 空间：O(k)'),
        ('13', '最小覆盖子串', 'Minimum Window Substring', '困难', '76',
         '给定字符串s和t，返回s中涵盖t所有字符的最小子串。如果不存在返回空串。',
         '滑动窗口+字符计数。用哈希表记录t中字符需求，右指针扩展窗口直到满足条件，然后左指针收缩窗口找最小，更新结果。用变量记录还需要的字符种类数来快速判断是否满足条件。',
         '时间：O(n) | 空间：O(m)'),
    ],
    '四、普通数组': [
        ('14', '最大子数组和', 'Maximum Subarray', '中等', '53',
         '给定整数数组nums，找出具有最大和的连续子数组，返回其最大和。',
         '动态规划（Kadane算法）。dp[i]表示以nums[i]结尾的最大子数组和。状态转移：dp[i]=max(nums[i],dp[i-1]+nums[i])。优化：只需维护当前值和最大值，无需存储整个dp数组。',
         '时间：O(n) | 空间：O(1)'),
        ('15', '合并区间', 'Merge Intervals', '中等', '56',
         '以数组intervals表示若干个区间，合并所有重叠区间，返回不重叠区间数组。',
         '排序+贪心。按区间左端点排序，初始化结果列表加入第一个区间。遍历剩余区间，如果当前区间左端点<=结果列表最后一个区间的右端点则合并，更新右端点为max(当前右端点,最后一个右端点)，否则加入结果列表。',
         '时间：O(nlogn) | 空间：O(logn)'),
        ('16', '轮转数组', 'Rotate Array', '中等', '189',
         '给定整数数组nums，将数组元素向右轮转k个位置，k是非负数。',
         '数组反转。先整体反转，再分别反转前后两部分。例如[1,2,3,4,5,6,7],k=3：整体反转[7,6,5,4,3,2,1]，反转前3个[5,6,7,4,3,2,1]，反转剩余[5,6,7,1,2,3,4]。',
         '时间：O(n) | 空间：O(1)'),
        ('17', '除自身以外数组的乘积', 'Product of Array Except Self', '中等', '238',
         '给定整数数组nums，返回数组answer，answer[i]等于除nums[i]外其余各元素的乘积。不使用除法，O(n)时间完成。',
         '前缀积×后缀积。第一次遍历从左到右计算前缀积，第二次遍历从右到左计算后缀积并乘到结果中，维护变量suffix表示从右往左的后缀积。',
         '时间：O(n) | 空间：O(1)'),
        ('18', '缺失的第一个正数', 'First Missing Positive', '困难', '41',
         '给定未排序整数数组nums，找出缺失的最小正整数。要求时间O(n)，空间O(1)。',
         '原地哈希。核心思想是将数字放到对应的位置，即数字i应该放在索引i-1处。遍历数组，如果nums[i]在[1,n]范围内且不在正确位置，就交换到正确位置。最后遍历找第一个不在正确位置的元素。',
         '时间：O(n) | 空间：O(1)'),
    ],
}

# 创建文档
doc = Document()

# 标题
title = doc.add_heading('LeetCode 热题 HOT 100 完整题解', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs:
    set_font(r, '黑体', 22)
    r.bold = True

add_p(doc, '本文档包含LeetCode HOT 100全部100道经典题目的完整题解')
add_p(doc, '每道题包含：题目描述、详细解题思路、核心算法步骤、时间/空间复杂度分析')
add_p(doc, '整理时间：2026年3月')
doc.add_page_break()

# 添加所有题目
for category, probs in problems.items():
    add_h(doc, category, 1)
    for num, title, eng, diff, pid, desc, sol, comp in probs:
        add_h(doc, f'{num}. {title} ({eng})', 2)
        add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
        add_h(doc, '题目描述', 3)
        add_p(doc, desc)
        add_h(doc, '解题思路', 3)
        add_p(doc, sol)
        add_h(doc, '复杂度分析', 3)
        add_p(doc, comp, bold=True)
        doc.add_paragraph()

print("第一批题目已生成（18题）")
print("继续添加更多题目...")

# 保存
doc.save('/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx')
print("已保存！")
