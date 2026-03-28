#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
继续添加更多题目到Word文档
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, name='宋体', size=10.5):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_h(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    set_font(r, '黑体', sizes.get(level, 10))
    return h

def add_p(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    set_font(r, '宋体', size)
    return p

# 打开现有文档
doc = Document('/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx')

# 五、矩阵
add_h(doc, '五、矩阵（Matrix）', 1)

matrix_problems = [
    ('19', '矩阵置零', 'Set Matrix Zeroes', '中等', '73',
     '给定m×n矩阵，如果一个元素为0，则将其所在行和列所有元素都设为0。请使用原地算法。',
     '用第一行和第一列作为标记位。首先用两个变量记录第一行和第一列是否需要置零，然后遍历非第一行第一列的元素，如果为0则将对应的第一行和第一列位置标记为0。再次遍历根据标记置零。',
     '时间：O(mn) | 空间：O(1)'),
    ('20', '螺旋矩阵', 'Spiral Matrix', '中等', '54',
     '给定m行n列矩阵matrix，请按照顺时针螺旋顺序返回矩阵中的所有元素。',
     '模拟螺旋遍历。定义四个边界top,bottom,left,right，按上、右、下、左的顺序遍历，每遍历完一条边收缩对应边界，直到边界相交。',
     '时间：O(mn) | 空间：O(1)'),
    ('21', '旋转图像', 'Rotate Image', '中等', '48',
     '给定n×n二维矩阵表示图像，将图像顺时针旋转90度。必须原地旋转。',
     '先转置矩阵（沿对角线翻转），再反转每一行。转置：matrix[i][j]与matrix[j][i]交换。反转：每行首尾交换直到中间。',
     '时间：O(n²) | 空间：O(1)'),
    ('22', '搜索二维矩阵II', 'Search a 2D Matrix II', '中等', '240',
     '编写高效算法从m×n矩阵中搜索target值。矩阵每行从左到右升序，每列从上到下升序。',
     '从右上角或左下角开始搜索。从右上角开始：如果target<当前值则向左，如果target>当前值则向下，直到找到target或越界。利用矩阵有序性质每次可以排除一行或一列。',
     '时间：O(m+n) | 空间：O(1)'),
]

for num, title, eng, diff, pid, desc, sol, comp in matrix_problems:
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

# 六、链表
add_h(doc, '六、链表（Linked List）', 1)

list_problems = [
    ('23', '相交链表', 'Intersection of Two Linked Lists', '简单', '160',
     '给定两个单链表的头节点headA和headB，找出并返回两个单链表相交的起始节点。如果不相交返回null。',
     '双指针交换起点法。指针pA遍历完A后从B开始，指针pB遍历完B后从A开始，如果相交则会在交点相遇，如果不相交则会在null相遇。原理：pA和pB走过的总路程相同，要么在交点相遇，要么同时到达null。',
     '时间：O(m+n) | 空间：O(1)'),
    ('24', '反转链表', 'Reverse Linked List', '简单', '206',
     '给定单链表的头节点head，请反转链表并返回反转后的链表。',
     '迭代法：三指针pre,curr,next。初始化pre=null,curr=head，遍历链表，先保存curr.next到next，然后curr.next指向pre，pre移动到curr，curr移动到next。递归法：递归反转后续节点，然后将当前节点接到反转后链表的末尾。',
     '时间：O(n) | 空间：O(1)'),
    ('25', '回文链表', 'Palindrome Linked List', '简单', '234',
     '给定单链表的头节点head，请判断该链表是否为回文链表。',
     '方法一：复制到数组后双指针判断。方法二：快慢指针找中点+反转后半部分+比较。快慢指针：慢指针走1步，快指针走2步，快指针到末尾时慢指针到中点。反转后半部分，然后比较前后两部分。',
     '时间：O(n) | 空间：O(1)'),
    ('26', '环形链表', 'Linked List Cycle', '简单', '141',
     '给定链表头节点head，判断链表中是否有环。如果链表中有某个节点可以通过连续跟踪next指针再次到达，则链表中存在环。',
     '快慢指针判环。慢指针每次走1步，快指针每次走2步。如果存在环，快指针一定会追上慢指针。如果不存在环，快指针会先到达末尾。',
     '时间：O(n) | 空间：O(1)'),
    ('27', '环形链表II', 'Linked List Cycle II', '中等', '142',
     '给定链表头节点head，返回链表开始入环的第一个节点。如果无环返回null。',
     '快慢指针+数学推导。快慢指针相遇后，一个指针回到头节点，两个指针都以1步前进，再次相遇的点就是环入口。原理：设头到入口距离为a，入口到相遇点距离为b，相遇点到入口距离为c，则2(a+b)=a+b+c+b，得a=c。',
     '时间：O(n) | 空间：O(1)'),
    ('28', '合并两个有序链表', 'Merge Two Sorted Lists', '简单', '21',
     '将两个升序链表合并为一个新的升序链表并返回。新链表通过拼接给定的两个链表的所有节点组成。',
     '递归或迭代。递归：比较两个链表头节点，较小者作为新头，递归合并剩余部分。迭代：创建虚拟头节点，用指针遍历两个链表，每次选择较小节点接到结果链表，直到一个链表为空，将另一个链表直接接上。',
     '时间：O(m+n) | 空间：O(1)'),
    ('29', '两数相加', 'Add Two Numbers', '中等', '2',
     '给定两个非空链表表示两个非负整数，数字以逆序存储，每个节点存一位数字，将两数相加并以相同形式返回和。',
     '模拟加法。遍历两个链表，对应位相加并加上进位，创建新节点存储当前位结果（和对10取模），更新进位（和除以10）。如果链表长度不同则补0，最后如果进位不为0则添加新节点。',
     '时间：O(max(m,n)) | 空间：O(max(m,n))'),
    ('30', '删除链表的倒数第N个结点', 'Remove Nth Node From End of List', '中等', '19',
     '给定链表头节点head，删除链表的倒数第n个节点，并返回头节点。',
     '快慢指针。快指针先走n步，然后快慢指针一起走，当快指针到达末尾时，慢指针在倒数第n+1个节点，修改慢指针的next删除目标节点。注意处理删除头节点的情况。',
     '时间：O(n) | 空间：O(1)'),
]

for p in list_problems:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

print("已添加矩阵和链表题目（共30题）")
doc.save('/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx')
print("已保存！")
