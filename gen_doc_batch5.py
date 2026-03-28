#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def set_font(run, name='宋体', size=10.5):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_h(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    set_font(r, '黑体', sizes.get(level, 10))
    return h

def add_p(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    set_font(r, '宋体', size)
    return p

doc = Document('/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx')

# 继续添加回溯题目
backtrack_problems_2 = [
    ('59', '组合总和', 'Combination Sum', '中等', '39',
     '给定无重复元素整数数组candidates和目标整数target，找出candidates中可以使数字和为目标数的所有组合。同一数字可重复使用。',
     '回溯。递归函数参数为当前索引和剩余目标值。遍历从当前索引开始的所有元素，如果元素<=剩余目标值则加入路径，递归（同一元素可重复使用所以索引不变），剩余目标值减去元素值，回溯。终止条件：剩余目标值为0记录结果，小于0返回。',
     '时间：O(n^(target/min)) | 空间：O(target/min)'),
    ('60', '括号生成', 'Generate Parentheses', '中等', '22',
     '给定数字n，生成所有由n对括号组成的有效组合。',
     '回溯。维护当前左括号数和右括号数。如果左括号数<n，可添加左括号；如果右括号数<左括号数，可添加右括号。终止条件：路径长度等于2n。',
     '时间：O(4ⁿ/√n) | 空间：O(n)'),
    ('61', '单词搜索', 'Word Search', '中等', '79',
     '给定m×n网格board和字符串word，如果word存在于网格中返回true。单词按相邻单元格字母构造，相邻是水平或垂直，同一单元格字母不可重复使用。',
     '回溯。遍历网格每个位置作为起点，DFS搜索。标记已访问位置，四个方向搜索，找到完整单词返回true，回溯时恢复标记。剪枝：如果当前字符不匹配直接返回。',
     '时间：O(mn×3ᴸ) | 空间：O(L)'),
    ('62', '分割回文串', 'Palindrome Partitioning', '中等', '131',
     '给定字符串s，将s分割成子串，每个子串都是回文串，返回所有可能的分割方案。',
     '回溯。递归函数参数为当前起始索引。遍历从起始索引开始的所有结束索引，如果子串是回文则加入路径，递归处理剩余部分，回溯。预先用DP计算所有子串是否为回文加速判断。',
     '时间：O(n×2ⁿ) | 空间：O(n)'),
    ('63', 'N皇后', 'N-Queens', '困难', '51',
     '给定整数n，返回所有不同的n皇后问题的解决方案。n皇后问题：nxn棋盘放n个皇后，使其不能互相攻击。',
     '回溯。每行放一个皇后，递归函数参数为当前行。遍历当前行所有列，检查当前位置是否与已放置皇后冲突（同列、对角线），不冲突则放置，递归下一行，回溯。用集合记录已占用的列和对角线加速检查。',
     '时间：O(n!) | 空间：O(n)'),
]

for p in backtrack_problems_2:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

# 十、二分查找
add_h(doc, '十、二分查找（Binary Search）', 1)

bs_problems = [
    ('64', '搜索插入位置', 'Search Insert Position', '简单', '35',
     '给定排序数组和目标值，找到目标值并返回索引，如果不存在返回按顺序插入的位置。',
     '二分查找。初始化left=0,right=n-1，当left<=right时，mid=left+(right-left)//2，如果nums[mid]==target返回mid，如果nums[mid]<target则left=mid+1，否则right=mid-1。最后返回left即为插入位置。',
     '时间：O(logn) | 空间：O(1)'),
    ('65', '搜索二维矩阵', 'Search a 2D Matrix', '中等', '74',
     '给定m×n矩阵，每行从左到右升序，每行第一个数大于上一行最后一个数，判断target是否在矩阵中。',
     '将二维矩阵视为一维有序数组。虚拟索引0到mn-1，映射到二维坐标：row=idx//n,col=idx%n。二分查找虚拟索引。或者从右上角开始搜索。',
     '时间：O(log(mn)) | 空间：O(1)'),
    ('66', '在排序数组中查找元素的第一个和最后一个位置', 'Find First and Last Position of Element in Sorted Array', '中等', '34',
     '给定升序整数数组nums和目标值target，找出给定目标值在数组中的开始位置和结束位置，如果不存在返回[-1,-1]。',
     '两次二分查找。第一次找左边界：当nums[mid]>=target时right=mid-1，否则left=mid+1，最后检查left是否越界和是否等于target。第二次找右边界类似。',
     '时间：O(logn) | 空间：O(1)'),
    ('67', '搜索旋转排序数组', 'Search in Rotated Sorted Array', '中等', '33',
     '给定升序整数数组nums在某个下标处旋转，给定目标值target，如果存在返回索引，否则返回-1。',
     '二分查找。先判断哪一半是有序的，如果target在有序范围内则在有序部分查找，否则在另一半查找。判断有序：如果nums[left]<=nums[mid]则左半有序，否则右半有序。',
     '时间：O(logn) | 空间：O(1)'),
    ('68', '寻找旋转排序数组中的最小值', 'Find Minimum in Rotated Sorted Array', '中等', '153',
     '给定升序数组在某个下标旋转，找出数组中的最小元素。',
     '二分查找。比较nums[mid]和nums[right]，如果nums[mid]<nums[right]则最小值在左半（包括mid），否则在右半。如果nums[mid]>nums[right]则left=mid+1，否则right=mid。',
     '时间：O(logn) | 空间：O(1)'),
    ('69', '寻找两个正序数组的中位数', 'Median of Two Sorted Arrays', '困难', '4',
     '给定两个正序（从小到大排序）数组nums1和nums2，长度分别为m和n，返回两个数组合并后的中位数。要求时间复杂度O(log(m+n))。',
     '二分查找（找第k小数）。将问题转化为找第k小的数。每次比较两个数组的第k/2个元素，排除较小一半的元素，k减去排除的元素数，递归直到k=1或一个数组为空。',
     '时间：O(log(m+n)) | 空间：O(1)'),
]

for p in bs_problems:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

# 十一、栈
add_h(doc, '十一、栈（Stack）', 1)

stack_problems = [
    ('70', '有效的括号', 'Valid Parentheses', '简单', '20',
     '给定只包含()[]{}的字符串s，判断字符串是否有效。有效条件：左括号必须用相同类型右括号闭合，左括号必须以正确顺序闭合。',
     '栈。遍历字符串，遇到左括号入栈，遇到右括号检查栈顶是否匹配，匹配则出栈，不匹配返回false。最后栈为空则有效。',
     '时间：O(n) | 空间：O(n)'),
    ('71', '最小栈', 'Min Stack', '中等', '155',
     '设计支持push、pop、top操作，并能在常数时间内检索最小元素的栈。',
     '辅助栈。使用一个辅助栈存储当前最小值，每次push时如果新元素小于等于辅助栈顶则同时push到辅助栈。pop时如果pop的元素等于辅助栈顶则辅助栈也pop。getMin返回辅助栈顶。',
     '时间：O(1) | 空间：O(n)'),
    ('72', '每日温度', 'Daily Temperatures', '中等', '739',
     '给定整数数组temperatures表示每日温度，返回数组answer，answer[i]表示对于第i天，下一个更高温度出现在几天后，如果不存在则为0。',
     '单调栈。维护一个递减栈，存储温度下标。遍历温度，当当前温度大于栈顶温度时，栈顶出栈，计算天数差为当前下标减栈顶下标。当前索引入栈。',
     '时间：O(n) | 空间：O(n)'),
    ('73', '下一个更大元素I', 'Next Greater Element I', '简单', '496',
     '给定两个没有重复元素的数组nums1和nums2，其中nums1是nums2的子集，找到nums1中每个元素在nums2中的下一个比其大的值。如果不存在返回-1。',
     '单调栈+哈希表。用单调栈找nums2中每个元素的下一个更大元素，用哈希表记录结果。遍历nums1直接查哈希表。',
     '时间：O(n) | 空间：O(n)'),
]

for p in stack_problems:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

print("已添加59-73题（共73题）")
doc.save('/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx')
print("已保存！")
