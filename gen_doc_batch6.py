#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def set_font(run, name='宋体', size=10.5):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_h(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    set_font(r, '黑体', sizes.get(level, 10))
    return h

def add_p(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    set_font(r, '宋体', size)
    return p

doc = Document('/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx')

# 十二、堆
add_h(doc, '十二、堆（Heap）', 1)

heap_problems = [
    ('74', '数组中的第K个最大元素', 'Kth Largest Element in an Array', '中等', '215',
     '给定整数数组nums和整数k，返回数组中第k个最大元素。注意是排序后的第k个最大，不是第k个不同元素。',
     '方法一：排序后取倒数第k个，时间O(nlogn)。方法二：小根堆。维护大小为k的小根堆，遍历数组，堆顶是第k大。方法三：快速选择。基于快排的分区思想，平均时间O(n)。',
     '时间：O(n) | 空间：O(1)'),
    ('75', '前K个高频元素', 'Top K Frequent Elements', '中等', '347',
     '给定整数数组nums和整数k，返回出现频率前k高的元素，可以按任意顺序返回答案。',
     '哈希表+小根堆。先用哈希表统计每个元素频率，然后用大小为k的小根堆，按频率维护前k个元素。或者桶排序：按频率分桶，从高频率开始收集k个元素。',
     '时间：O(nlogk) | 空间：O(n)'),
    ('76', '数据流的中位数', 'Find Median from Data Stream', '困难', '295',
     '设计数据结构支持两种操作：添加整数到数据流，找出当前所有元素的中位数。',
     '双堆。用大根堆存储较小一半，小根堆存储较大一半，保持大小堆大小差不超过1。添加时根据大小判断放入哪个堆，然后平衡两个堆的大小。中位数：如果堆大小相等则两堆顶平均，否则大堆顶。',
     '时间：O(logn) | 空间：O(n)'),
]

for p in heap_problems:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

# 十三、贪心算法
add_h(doc, '十三、贪心算法（Greedy）', 1)

greedy_problems = [
    ('77', '买卖股票的最佳时机', 'Best Time to Buy and Sell Stock', '简单', '121',
     '给定数组prices，prices[i]是某股票第i天价格，只能买卖一次，求最大利润。',
     '贪心。遍历价格，维护最小价格和最大利润。每天计算如果当天卖出利润多少，更新最大利润，然后更新最小价格。',
     '时间：O(n) | 空间：O(1)'),
    ('78', '跳跃游戏', 'Jump Game', '中等', '55',
     '给定非负整数数组nums，最初位于数组第一个下标，数组中每个元素代表在该位置可以跳跃的最大长度，判断能否到达最后一个下标。',
     '贪心。维护最大可达位置maxReach，遍历数组，如果当前位置>maxReach返回false，否则更新maxReach=max(maxReach,i+nums[i])，最后返回true。',
     '时间：O(n) | 空间：O(1)'),
    ('79', '跳跃游戏II', 'Jump Game II', '中等', '45',
     '给定非负整数数组nums，最初位于第一个下标，数组中每个元素代表在该位置可以跳跃的最大长度，返回到达最后一个下标的最小跳跃次数。',
     '贪心。维护当前可达范围end和下一步可达范围farthest，遍历数组，更新farthest，当到达end时跳跃次数加1并更新end为farthest。',
     '时间：O(n) | 空间：O(1)'),
    ('80', '划分字母区间', 'Partition Labels', '中等', '763',
     '给定字符串s，将其划分为尽可能多的片段，同一字母最多出现在一个片段中，返回每个片段的长度。',
     '贪心。先用哈希表记录每个字母最后出现位置。遍历字符串，维护当前片段的最远结束位置end，当遍历到end时形成一个片段。end为当前片段所有字母最后位置的最大值。',
     '时间：O(n) | 空间：O(1)'),
]

for p in greedy_problems:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

# 十四、动态规划
add_h(doc, '十四、动态规划（Dynamic Programming）', 1)

dp_problems = [
    ('81', '爬楼梯', 'Climbing Stairs', '简单', '70',
     '假设你正在爬楼梯，需要n阶才能到达楼顶，每次可以爬1或2个台阶，问有多少种不同的方法可以爬到楼顶。',
     'DP。dp[i]表示爬到第i阶的方法数，dp[i]=dp[i-1]+dp[i-2]，因为可以从i-1爬1步或从i-2爬2步。初始dp[0]=1,dp[1]=1。空间优化：只需维护两个变量。',
     '时间：O(n) | 空间：O(1)'),
    ('82', '斐波那契数', 'Fibonacci Number', '简单', '509',
     '给定n，计算斐波那契数F(n)。F(0)=0,F(1)=1,F(n)=F(n-1)+F(n-2)。',
     'DP。dp[i]=dp[i-1]+dp[i-2]。空间优化：只需维护前两个值。矩阵快速幂可将时间优化到O(logn)。',
     '时间：O(n) | 空间：O(1)'),
    ('83', '杨辉三角', 'Pascal\'s Triangle', '简单', '118',
     '给定非负整数numRows，生成杨辉三角的前numRows行。',
     'DP。每行首尾为1，中间元素等于上一行相邻两元素之和。dp[i][j]=dp[i-1][j-1]+dp[i-1][j]。',
     '时间：O(n²) | 空间：O(n²)'),
    ('84', '打家劫舍', 'House Robber', '中等', '198',
     '给定整数数组nums表示房屋金额，相邻房屋有安全系统不能同时被盗，求一夜能偷窃的最高金额。',
     'DP。dp[i]表示偷窃前i间房屋最高金额，dp[i]=max(dp[i-1],dp[i-2]+nums[i])。不偷i则为dp[i-1]，偷i则为dp[i-2]+nums[i]。空间优化只需维护两个变量。',
     '时间：O(n) | 空间：O(1)'),
    ('85', '完全平方数', 'Perfect Squares', '中等', '279',
     '给定正整数n，找到若干个完全平方数（如1,4,9,...）使它们的和等于n，返回完全平方数最少数量。',
     'DP（完全背包）。dp[i]表示和为i的最少完全平方数个数，dp[i]=min(dp[i-j*j])+1，j从1到sqrt(i)。初始dp[0]=0，其他为正无穷。',
     '时间：O(n√n) | 空间：O(n)'),
    ('86', '零钱兑换', 'Coin Change', '中等', '322',
     '给定硬币面额数组coins和总金额amount，计算可以凑成总金额所需的最少硬币个数，如果无解返回-1。',
     'DP（完全背包）。dp[i]表示凑成金额i的最少硬币数，dp[i]=min(dp[i-coin])+1，coin遍历所有硬币。初始dp[0]=0，其他为amount+1。',
     '时间：O(n×amount) | 空间：O(amount)'),
    ('87', '最长递增子序列', 'Longest Increasing Subsequence', '中等', '300',
     '给定整数数组nums，找到最长严格递增子序列的长度。',
     'DP。dp[i]表示以nums[i]结尾的最长递增子序列长度，dp[i]=max(dp[j])+1，j<i且nums[j]<nums[i]。二分优化：维护一个递增数组，用二分查找替换。',
     '时间：O(n²)或O(nlogn) | 空间：O(n)'),
    ('88', '单词拆分', 'Word Break', '中等', '139',
     '给定字符串s和字符串列表wordDict，判断s是否可以被空格拆分为一个或多个在字典中出现的单词。',
     'DP。dp[i]表示前i个字符能否拆分，dp[i]=OR(dp[j]且s[j:i]在字典中)，j从0到i-1。初始dp[0]=true。',
     '时间：O(n²) | 空间：O(n)'),
    ('89', '最长回文子序列', 'Longest Palindromic Subsequence', '中等', '516',
     '给定字符串s，找到其中最长的回文子序列，并返回该序列的长度。',
     'DP。dp[i][j]表示s[i:j+1]的最长回文子序列长度。如果s[i]==s[j]，dp[i][j]=dp[i+1][j-1]+2；否则dp[i][j]=max(dp[i+1][j],dp[i][j-1])。按区间长度从小到大计算。',
     '时间：O(n²) | 空间：O(n²)'),
]

for p in dp_problems:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

# 十五、多维动态规划
add_h(doc, '十五、多维动态规划（Multi-dimensional DP）', 1)

dp2_problems = [
    ('90', '不同路径', 'Unique Paths', '中等', '62',
     '给定m×n网格，机器人从左上角出发每次只能向下或向右移动一步，到达右下角有多少条不同路径。',
     'DP。dp[i][j]表示到(i,j)的路径数，dp[i][j]=dp[i-1][j]+dp[i][j-1]。第一行和第一列初始化为1。空间优化：只需一维数组。',
     '时间：O(mn) | 空间：O(min(m,n))'),
    ('91', '最小路径和', 'Minimum Path Sum', '中等', '64',
     '给定m×n网格grid包含非负整数，找出从左上角到右下角使路径数字总和最小的路径，返回最小和。',
     'DP。dp[i][j]表示到(i,j)的最小路径和，dp[i][j]=grid[i][j]+min(dp[i-1][j],dp[i][j-1])。第一行和第一列需累加。',
     '时间：O(mn) | 空间：O(min(m,n))'),
    ('92', '最长公共子序列', 'Longest Common Subsequence', '中等', '1143',
     '给定两个字符串text1和text2，返回它们的最长公共子序列的长度。',
     'DP。dp[i][j]表示text1前i个和text2前j个的LCS长度。如果text1[i-1]==text2[j-1]，dp[i][j]=dp[i-1][j-1]+1；否则dp[i][j]=max(dp[i-1][j],dp[i][j-1])。',
     '时间：O(mn) | 空间：O(mn)'),
    ('93', '编辑距离', 'Edit Distance', '困难', '72',
     '给定两个单词word1和word2，返回将word1转换成word2所使用的最少操作数。可插入、删除、替换一个字符。',
     'DP。dp[i][j]表示word1前i个转换到word2前j个的最少操作数。如果word1[i-1]==word2[j-1]，dp[i][j]=dp[i-1][j-1]；否则dp[i][j]=1+min(插入,删除,替换)。',
     '时间：O(mn) | 空间：O(mn)'),
    ('94', '买卖股票的最佳时机含冷冻期', 'Best Time to Buy and Sell Stock with Cooldown', '中等', '309',
     '给定数组prices，prices[i]是第i天价格，卖出股票后有1天冷冻期不能买入，求最大利润。',
     'DP。三种状态：hold持有股票，sold刚卖出，rest不持有且不冷冻。hold[i]=max(hold[i-1],rest[i-1]-prices[i])，sold[i]=hold[i-1]+prices[i]，rest[i]=max(rest[i-1],sold[i-1])。',
     '时间：O(n) | 空间：O(1)'),
    ('95', '正则表达式匹配', 'Regular Expression Matching', '困难', '10',
     '给定字符串s和字符规律p，实现支持.和*的正则表达式匹配。.匹配任意单个字符，*匹配零个或多个前面的元素。',
     'DP。dp[i][j]表示s前i个和p前j个是否匹配。如果p[j-1]是*，dp[i][j]=dp[i][j-2]（匹配0次）或(首字符匹配且dp[i-1][j])（匹配多次）；否则dp[i][j]=首字符匹配且dp[i-1][j-1]。',
     '时间：O(mn) | 空间：O(mn)'),
]

for p in dp2_problems:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

# 十六、技巧
add_h(doc, '十六、技巧（Tricks）', 1)

trick_problems = [
    ('96', '多数元素', 'Majority Element', '简单', '169',
     '给定大小为n的数组nums，找出其中的多数元素。多数元素是出现次数大于⌊n/2⌋的元素。',
     '摩尔投票法。维护候选人和计数，遍历数组，如果计数为0则当前元素作为候选人，如果当前元素等于候选人计数加1否则减1。最后候选人即为多数元素。',
     '时间：O(n) | 空间：O(1)'),
    ('97', '丢失的数字', 'Missing Number', '简单', '268',
     '给定包含[0,n]中n个数的数组nums，找出[0,n]中缺失的那个数。',
     '数学：求和公式n(n+1)/2减去数组和即为缺失数。位运算：异或所有数和索引，成对消去，最后剩下缺失数。',
     '时间：O(n) | 空间：O(1)'),
    ('98', '汉明距离', 'Hamming Distance', '简单', '461',
     '给定两个整数x和y，计算它们二进制表示中对应位不同的位置数。',
     '异或后统计1的个数。n&(n-1)可以消去最低位1，循环计数直到n为0。',
     '时间：O(1) | 空间：O(1)'),
    ('99', '比特位计数', 'Counting Bits', '简单', '338',
     '给定非负整数n，对于0≤i≤n的每个i，计算其二进制表示中1的个数，返回数组。',
     'DP。dp[i]=dp[i>>1]+(i&1)，即i右移一位的1的个数加上最低位。或dp[i]=dp[i&(i-1)]+1，消去最低位1后的1的个数加1。',
     '时间：O(n) | 空间：O(1)'),
    ('100', '2的幂', 'Power of Two', '简单', '231',
     '给定整数n，判断它是否为2的幂次方。',
     '位运算。2的幂二进制表示只有一位为1，如1000。n>0且n&(n-1)==0即为2的幂。n-1会将最低位1变为0，后面全变1，所以n&(n-1)==0说明只有一位1。',
     '时间：O(1) | 空间：O(1)'),
]

for p in trick_problems:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

print("="*50)
print("✅ LeetCode HOT100 完整题解文档生成完成！")
print("="*50)
print("共100道题目，包含：")
print("- 完整题目描述")
print("- 详细解题思路")
print("- 时间/空间复杂度分析")
print("="*50)
doc.save('/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx')
print("文件已保存！")
