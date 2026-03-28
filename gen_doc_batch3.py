#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

def set_font(run, name='宋体', size=10.5):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_h(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    set_font(r, '黑体', sizes.get(level, 10))
    return h

def add_p(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    set_font(r, '宋体', size)
    return p

doc = Document('/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx')

# 继续添加链表题目
list_problems_2 = [
    ('31', '两两交换链表中的节点', 'Swap Nodes in Pairs', '中等', '24',
     '给定链表，两两交换其中相邻节点，并返回交换后链表的头节点。如果节点数是奇数，最后一个节点保持原样。',
     '递归或迭代。递归：交换前两个节点，然后递归处理后续链表。迭代：创建虚拟头节点，用指针遍历，每次处理一对节点，调整指针指向完成交换。',
     '时间：O(n) | 空间：O(1)'),
    ('32', 'K个一组翻转链表', 'Reverse Nodes in k-Group', '困难', '25',
     '给定链表头节点head和整数k，每k个节点一组进行翻转，返回翻转后的链表。如果节点总数不是k的整数倍，最后剩余节点保持原样。',
     '递归或迭代。先检查剩余部分是否有k个节点，如果没有直接返回。如果有k个节点，翻转这k个节点，然后递归处理后续链表。翻转部分：记录第k+1个节点，翻转前k个节点，连接翻转后的尾节点与递归结果。',
     '时间：O(n) | 空间：O(1)'),
    ('33', '复制带随机指针的链表', 'Copy List with Random Pointer', '中等', '138',
     '给定链表，每个节点包含额外随机指针可指向任意节点或null，深拷贝这个链表。',
     '三步法。1)复制每个节点插入到原节点后面；2)根据原节点的random设置复制节点的random；3)拆分链表恢复原链表并取出复制链表。利用原链表结构避免使用哈希表。',
     '时间：O(n) | 空间：O(1)'),
    ('34', '排序链表', 'Sort List', '中等', '148',
     '给定链表头节点head，请将其按升序排列并返回排序后的链表。要求时间O(nlogn)，空间O(1)。',
     '归并排序。快慢指针找中点将链表分成两半，递归排序左右两半，然后合并两个有序链表。合并：创建虚拟头节点，依次选择较小节点接到结果链表。',
     '时间：O(nlogn) | 空间：O(logn)'),
    ('35', '合并K个升序链表', 'Merge k Sorted Lists', '困难', '23',
     '给定链表数组lists，每个链表按升序排列，将所有链表合并到一个升序链表中并返回。',
     '方法一：优先队列。将所有链表头节点入队，每次取出最小节点接到结果，将该节点下一个入队。时间O(nlogk)。方法二：分治。两两合并，递归处理。时间O(nlogk)。',
     '时间：O(nlogk) | 空间：O(k)'),
    ('36', 'LRU缓存', 'LRU Cache', '中等', '146',
     '设计LRU缓存数据结构，支持get和put操作，容量为capacity，超出容量时淘汰最久未使用的数据。要求两个操作时间复杂度O(1)。',
     '哈希表+双向链表。哈希表存储key到节点的映射，双向链表维护访问顺序，头部为最近使用，尾部为最久未使用。get时移动到头部，put时如果存在更新并移动，不存在则创建插入头部，超出容量删除尾部。',
     '时间：O(1) | 空间：O(capacity)'),
]

for p in list_problems_2:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

# 七、二叉树
add_h(doc, '七、二叉树（Binary Tree）', 1)

tree_problems = [
    ('37', '二叉树的中序遍历', 'Binary Tree Inorder Traversal', '简单', '94',
     '给定二叉树根节点root，返回其节点值的中序遍历。',
     '递归：左-根-右。迭代：用栈模拟递归，当前节点入栈后一直向左，到达null时出栈访问，然后转向右子树。Morris遍历： threaded binary tree，时间O(n)空间O(1)。',
     '时间：O(n) | 空间：O(h)'),
    ('38', '二叉树的最大深度', 'Maximum Depth of Binary Tree', '简单', '104',
     '给定二叉树根节点root，返回其最大深度。最大深度是从根节点到最远叶子节点的最长路径上的节点数。',
     '递归：max(左子树深度,右子树深度)+1。BFS：层序遍历，统计层数。DFS：栈模拟递归。',
     '时间：O(n) | 空间：O(h)'),
    ('39', '翻转二叉树', 'Invert Binary Tree', '简单', '226',
     '给定二叉树根节点root，翻转这棵二叉树并返回其根节点。',
     '递归：交换左右子树，然后递归翻转左右子树。BFS：层序遍历，每层交换左右子节点。DFS：栈模拟，处理每个节点时交换左右子节点。',
     '时间：O(n) | 空间：O(h)'),
    ('40', '对称二叉树', 'Symmetric Tree', '简单', '101',
     '给定二叉树根节点root，检查它是否轴对称。',
     '递归：比较左右子树是否镜像对称，即左.left与右.right且左.right与右.left。迭代：用队列或栈，每次取出两个节点比较，然后按对称顺序入队/入栈。',
     '时间：O(n) | 空间：O(h)'),
    ('41', '二叉树的直径', 'Diameter of Binary Tree', '简单', '543',
     '给定二叉树根节点root，计算其直径长度。直径是任意两个节点路径长度中的最大值，可能不经过根节点。',
     '递归。直径=左子树深度+右子树深度。递归计算每个节点的深度，同时更新全局最大直径。一个递归函数返回深度，同时用全局变量或引用参数记录直径。',
     '时间：O(n) | 空间：O(h)'),
    ('42', '二叉树的层序遍历', 'Binary Tree Level Order Traversal', '中等', '102',
     '给定二叉树根节点root，返回其节点值的层序遍历（即逐层从左到右访问）。',
     'BFS。用队列，根节点入队。每次处理一层的所有节点，记录当前层节点值，将下一层节点入队。递归法：按层数递归，将节点值加入对应层的列表。',
     '时间：O(n) | 空间：O(w)'),
    ('43', '将有序数组转换为二叉搜索树', 'Convert Sorted Array to Binary Search Tree', '简单', '108',
     '给定整数数组nums，元素值升序排列，将其转换为高度平衡的二叉搜索树。',
     '递归。选择数组中间元素作为根节点，左半部分构建左子树，右半部分构建右子树。这样构建的树高度平衡。',
     '时间：O(n) | 空间：O(logn)'),
    ('44', '验证二叉搜索树', 'Validate Binary Search Tree', '中等', '98',
     '给定二叉树根节点root，判断其是否为有效的二叉搜索树。BST定义：左子树所有节点值<根节点值<右子树所有节点值，左右子树也是BST。',
     '递归。传递上下界，每个节点值必须在(lower,upper)范围内。初始范围(-inf,+inf)，左子树更新上界为当前值，右子树更新下界为当前值。中序遍历：BST中序遍历结果是有序的。',
     '时间：O(n) | 空间：O(h)'),
]

for p in tree_problems:
    num, title, eng, diff, pid, desc, sol, comp = p
    add_h(doc, f'{num}. {title} ({eng})', 2)
    add_p(doc, f'题号：{pid} | 难度：{diff}', bold=True)
    add_h(doc, '题目描述', 3)
    add_p(doc, desc)
    add_h(doc, '解题思路', 3)
    add_p(doc, sol)
    add_h(doc, '复杂度分析', 3)
    add_p(doc, comp, bold=True)
    doc.add_paragraph()

print("已添加36-44题（共44题）")
doc.save('/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx')
print("已保存！")
