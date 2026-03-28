#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LeetCode HOT100 完整题解 Word文档生成器 - 完整版
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='SimSun', font_size=10.5):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    font_names = ['SimHei', 'Microsoft YaHei', 'SimSun']
    font_size = {1: 18, 2: 16, 3: 14, 4: 12}.get(level, 10)
    set_chinese_font(run, font_names[0], font_size)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_zh(doc, text, bold=False, font_size=10.5):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    set_chinese_font(run, 'SimSun', font_size)
    return p

def add_code_block(doc, code_text):
    """添加代码块"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    set_chinese_font(run, 'Courier New', 9)
    run.font.color.rgb = RGBColor(0, 100, 0)
    return p

def add_problem(doc, number, title, english_title, level, link, description, solution, complexity):
    """添加一道题目"""
    add_heading_zh(doc, f'{number}. {title} ({english_title})', level=2)
    add_paragraph_zh(doc, f'题号：{link} | 难度：{level}', bold=True)
    add_heading_zh(doc, '题目描述', level=3)
    add_paragraph_zh(doc, description)
    add_heading_zh(doc, '解题思路', level=3)
    add_paragraph_zh(doc, solution)
    add_heading_zh(doc, '复杂度分析', level=3)
    add_paragraph_zh(doc, complexity, bold=True)
    doc.add_paragraph()  # 空行

# 创建文档
doc = Document()

# 设置文档标题
title = doc.add_heading('LeetCode 热题 HOT 100 完整题解', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    set_chinese_font(run, 'SimHei', 22)
    run.bold = True

# 添加说明
add_paragraph_zh(doc, '本文档包含LeetCode HOT 100全部100道经典题目的完整题解')
add_paragraph_zh(doc, '每道题包含：题目描述、详细解题思路、核心算法步骤、时间/空间复杂度分析')
add_paragraph_zh(doc, '整理时间：2026年3月')
doc.add_page_break()

# 添加目录
add_heading_zh(doc, '目录', level=1)
toc_items = [
    '一、哈希（4题）',
    '二、双指针（4题）', 
    '三、滑动窗口（3题）',
    '四、普通数组（4题）',
    '五、矩阵（2题）',
    '六、链表（9题）',
    '七、二叉树（15题）',
    '八、图论（8题）',
    '九、回溯（8题）',
    '十、二分查找（6题）',
    '十一、栈（5题）',
    '十二、堆（3题）',
    '十三、贪心算法（6题）',
    '十四、动态规划（1D - 10题）',
    '十五、多维动态规划（6题）',
    '十六、技巧（7题）'
]
for item in toc_items:
    add_paragraph_zh(doc, item)
doc.add_page_break()

# ========== 一、哈希 ==========
add_heading_zh(doc, '一、哈希（Hash）', level=1)

problems_hash = [
    (1, '两数之和', 'Two Sum', '简单', '1', 
     '给定一个整数数组 nums 和一个整数目标值 target，在该数组中找出和为目标值 target 的那两个整数，并返回它们的数组下标。假设每种输入只会对应一个答案，同一个元素不能重复使用。',
     '核心思想是使用哈希表存储已经遍历过的数字及其下标。对于当前数字x，在哈希表中查找target-x是否存在。关键点是先查后存，避免同一个元素使用两次。这种方法只需要遍历一次数组。方法一暴力枚举时间复杂度O(n²)，方法二哈希表时间复杂度O(n)，空间复杂度O(n)。',
     '时间复杂度：O(n) | 空间复杂度：O(n)'),
    
    (2, '字母异位词分组', 'Group Anagrams', '中等', '49',
     '给你一个字符串数组，将字母异位词组合在一起。字母异位词是由重新排列源单词的所有字母得到的一个新单词。',
     '核心思想是字母异位词排序后得到的字符串相同。方法一排序作为键：将每个字符串排序，排序后的字符串作为哈希表的键。时间复杂度O(n × k log k)。方法二字符计数作为键：使用26个长度的数组统计每个字符出现次数，将计数数组转为字符串作为键。时间复杂度O(n × k)，避免排序更优。',
     '时间复杂度：O(n × k) | 空间复杂度：O(n × k)'),
    
    (3, '最长连续序列', 'Longest Consecutive Sequence', '中等', '128',
     '给定一个未排序的整数数组 nums，找出数字连续的最长序列的长度。要求时间复杂度为 O(n)。',
     '核心思想是只从序列的起点开始计数。步骤：1) 将所有数字存入哈希集合，实现O(1)查找；2) 遍历集合，如果num-1不在集合中，说明num是序列起点；3) 从num开始，不断查找num+1, num+2...，直到不连续；4) 记录当前序列长度，更新最大值。关键点是通过判断num-1是否存在，确保每个序列只被计算一次，保证O(n)时间复杂度。',
     '时间复杂度：O(n) | 空间复杂度：O(n)'),
    
    (4, '只出现一次的数字', 'Single Number', '简单', '136',
     '给定一个非空整数数组，除了某个元素只出现一次以外，其余每个元素均出现两次。找出那个只出现了一次的元素。要求线性时间复杂度，不使用额外空间。',
     '核心思想是利用异或运算的性质。a ^ a = 0（任何数与自身异或为0）；a ^ 0 = a（任何数与0异或为其本身）；异或运算满足交换律和结合律。算法：将所有数字进行异或运算，成对出现的数字异或结果为0，最后剩下的就是只出现一次的数字。',
     '时间复杂度：O(n) | 空间复杂度：O(1)'),
]

for p in problems_hash:
    add_problem(doc, *p)

# ========== 二、双指针 ==========
doc.add_page_break()
add_heading_zh(doc, '二、双指针（Two Pointers）', level=1)

problems_tp = [
    (5, '移动零', 'Move Zeroes', '简单', '283',
     '给定一个数组 nums，编写一个函数将所有 0 移动到数组的末尾，同时保持非零元素的相对顺序。必须在不复制数组的情况下原地对数组进行操作。',
     '核心思想是双指针，一个指针遍历，一个指针记录非零元素应该放置的位置。步骤：1) 初始化指针j = 0，表示下一个非零元素应该放的位置；2) 遍历数组，如果nums[i] != 0，将nums[i]赋值给nums[j]，然后j++；3) 将j到数组末尾的所有元素赋值为0。优化版本可以在一次遍历中完成，遍历时直接交换非零元素到前面。',
     '时间复杂度：O(n) | 空间复杂度：O(1)'),
    
    (6, '盛最多水的容器', 'Container With Most Water', '中等', '11',
     '给定一个长度为 n 的整数数组 height。有 n 条垂线，第 i 条线的两个端点是 (i, 0) 和 (i, height[i])。找出其中的两条线，使得它们与 x 轴共同构成的容器可以容纳最多的水。',
     '核心思想是双指针从两端向中间移动，移动较小高度的指针。容器面积 = min(height[left], height[right]) × (right - left)。如果移动较大高度的指针，宽度减小，高度不变或更小，面积一定减小。如果移动较小高度的指针，虽然宽度减小，但高度可能增大，面积可能增大。步骤：初始化left = 0, right = n-1, maxArea = 0；当left < right时，计算当前面积并更新maxArea；移动较小高度的指针。',
     '时间复杂度：O(n) | 空间复杂度：O(1)'),
    
    (7, '三数之和', '3Sum', '中等', '15',
     '给你一个整数数组 nums，判断是否存在三元组 [nums[i], nums[j], nums[k]] 满足 i != j、i != k 且 j != k，同时还满足 nums[i] + nums[j] + nums[k] == 0。返回所有和为 0 且不重复的三元组。',
     '核心思想是排序 + 双指针。步骤：1) 对数组进行排序；2) 遍历数组，固定第一个数nums[i]，如果nums[i] > 0则break（三数之和不可能为0）；3) 去重：如果i > 0且nums[i] == nums[i-1]，跳过；4) 使用双指针left = i+1, right = n-1，计算sum，如果sum == 0记录结果并去重，如果sum < 0则left++，如果sum > 0则right--。去重处理：外层循环跳过重复的第一个数，找到解后跳过左右指针的重复元素。',
     '时间复杂度：O(n²) | 空间复杂度：O(1)'),
    
    (8, '接雨水', 'Trapping Rain Water', '困难', '42',
     '给定 n 个非负整数表示每个宽度为 1 的柱子的高度图，计算按此排列的柱子，下雨之后能接多少雨水。',
     '核心思想是每个位置能接的雨水 = min(左边最大高度, 右边最大高度) - 当前高度。方法一动态规划：预处理每个位置左边的最大高度和右边的最大高度，时间复杂度O(n)，空间复杂度O(n)。方法二双指针：用两个指针从两端向中间移动，维护左边的最大值和右边的最大值，较小的一侧可以计算接水量。步骤：初始化left = 0, right = n-1, leftMax = 0, rightMax = 0, result = 0；当left < right时，处理较小高度的一侧。',
     '时间复杂度：O(n) | 空间复杂度：O(1)'),
]

for p in problems_tp:
    add_problem(doc, *p)

# ========== 三、滑动窗口 ==========
doc.add_page_break()
add_heading_zh(doc, '三、滑动窗口（Sliding Window）', level=1)

problems_sw = [
    (9, '无重复字符的最长子串', 'Longest Substring Without Repeating Characters', '中等', '3',
     '给定一个字符串 s，请你找出其中不含有重复字符的最长子串的长度。',
     '核心思想是滑动窗口，维护一个不包含重复字符的窗口。步骤：1) 初始化left = 0, maxLen = 0；2) 使用哈希集合或哈希表记录窗口中的字符；3) 遍历字符串，right从0到n-1，如果s[right]不在窗口中，将s[right]加入窗口并更新maxLen；如果s[right]已在窗口中，移动left直到s[right]不在窗口中。优化：使用哈希表记录字符的索引，可以直接将left跳到重复字符的下一个位置。',
     '时间复杂度：O(n) | 空间复杂度：O(min(m, n))'),
    
    (10, '找到字符串中所有字母异位词', 'Find All Anagrams in a String', '中等', '438',
     '给定两个字符串 s 和 p，找到 s 中所有 p 的异位词的子串，返回这些子串的起始索引。',
     '核心思想是滑动窗口 + 字符计数。步骤：1) 统计p中每个字符的出现次数；2) 维护一个大小为p.length()的滑动窗口；3) 使用数组统计窗口中每个字符的出现次数；4) 当窗口中的字符计数与p的字符计数相同时，记录起始索引；5) 滑动窗口：右移一位，移除左边字符，添加右边字符。优化：使用一个变量记录当前窗口与目标差异的字符种类数，避免每次比较整个数组。',
     '时间复杂度：O(n) | 空间复杂度：O(1)'),
    
    (11, '和为 K 的子数组', 'Subarray Sum Equals K', '中等', '560',
     '给你一个整数数组 nums 和一个整数 k，请你统计并返回该数组中和为 k 的子数组的个数。子数组是数组中元素的连续非空序列。',
     '核心思想是前缀和 + 哈希表。关键公式：子数组和 = prefixSum[j] - prefixSum[i-1] = k，即 prefixSum[i-1] = prefixSum[j] - k。步骤：1) 使用哈希表记录每个前缀和出现的次数；2) 遍历数组，计算当前前缀和；3) 在哈希表中查找prefixSum - k出现的次数，加到结果中；4) 将当前前缀和加入哈希表。初始化：hashMap[0] = 1，表示前缀和为0出现1次（处理从索引0开始的子数组）。',
     '时间复杂度：O(n) | 空间复杂度：O(n)'),
]

for p in problems_sw:
    add_problem(doc, *p)

# ========== 四、普通数组 ==========
doc.add_page_break()
add_heading_zh(doc, '四、普通数组（Array）', level=1)

problems_arr = [
    (12, '最大子数组和', 'Maximum Subarray', '中等', '53',
     '给你一个整数数组 nums，请你找出一个具有最大和的连续子数组（子数组最少包含一个元素），返回其最大和。',
     '核心思想是动态规划（Kadane算法）。状态定义：dp[i]表示以nums[i]结尾的最大子数组和。状态转移方程：dp[i] = max(nums[i], dp[i-1] + nums[i])，要么从当前元素重新开始，要么延续前面的子数组。优化：不需要存储整个dp数组，只需维护当前值和最大值。步骤：1) 初始化currentSum = nums[0], maxSum = nums[0]；2) 从i=1遍历到n-1，currentSum = max(nums[i], currentSum + nums[i])，maxSum = max(maxSum, currentSum)。',
     '时间复杂度：O(n) | 空间复杂度：O(1)'),
    
    (13, '合并区间', 'Merge Intervals', '中等', '56',
     '以数组 intervals 表示若干个区间的集合，请你合并所有重叠的区间，并返回一个不重叠的区间数组，该数组需恰好覆盖输入中的所有区间。',
     '核心思想是排序 + 贪心。步骤：1) 按区间的左端点排序；2) 初始化结果列表，将第一个区间加入；3) 遍历剩余的区间，如果当前区间的左端点 <= 结果列表最后一个区间的右端点，则有重叠，合并：更新最后一个区间的右端点为max(当前右端点, 最后一个右端点)；否则无重叠，将当前区间加入结果列表。',
     '时间复杂度：O(n log n) | 空间复杂度：O(log n)'),
    
    (14, '轮转数组', 'Rotate Array', '中等', '189',
     '给定一个整数数组 nums，将数组中的元素向右轮转 k 个位置，其中 k 是非负数。',
     '方法一使用额外数组时间复杂度O(n)，空间复杂度O(n)。方法二数组反转（推荐）：思路是先整体反转，再分别反转前后两部分。例如[1,2,3,4,5,6,7], k=3：1) 整体反转：[7,6,5,4,3,2,1]；2) 反转前k个：[5,6,7,4,3,2,1]；3) 反转剩余：[5,6,7,1,2,3,4]。步骤：k = k % n（处理k大于n的情况），反转整个数组，反转[0, k-1]，反转[k, n-1]。',
     '时间复杂度：O(n) | 空间复杂度：O(1)'),
    
    (15, '除自身以外数组的乘积', 'Product of Array Except Self', '中等', '238',
     '给你一个整数数组 nums，返回数组 answer，其中 answer[i] 等于 nums 中除 nums[i] 之外其余各元素的乘积。请不要使用除法，且在 O(n) 时间复杂度内完成此题。',
     '核心思想是前缀积 × 后缀积。步骤：1) 创建结果数组answer；2) 第一次遍历（从左到右）计算前缀积，answer[i] = nums[0] × nums[1] × ... × nums[i-1]；3) 第二次遍历（从右到左）计算后缀积并乘到结果中，维护一个变量suffix表示从右往左的后缀积，answer[i] *= suffix，suffix *= nums[i]。',
     '时间复杂度：O(n) | 空间复杂度：O(1)'),
]

for p in problems_arr:
    add_problem(doc, *p)

print("前15题完成，继续添加更多题目...")

# 保存文档
doc.save('/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx')
print("Word文档已保存！")
print("文件路径：/root/.openclaw/workspace/LeetCode_HOT100_完整题解.docx")
